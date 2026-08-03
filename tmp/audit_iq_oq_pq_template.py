from pathlib import Path
from zipfile import ZipFile
from docx import Document
from docx.oxml.ns import qn

path = Path(r"C:\Users\JONGB\OneDrive - Thermo Fisher Scientific\Documents\Playground\newMBD\output\newMBD_Application_IQ_OQ_PQ_Signoff_Template.docx")
doc = Document(path)

headings = [(p.style.name, p.text) for p in doc.paragraphs if p.style.name.startswith("Heading")]
required = [
    "Installation Qualification (IQ)",
    "Operational Qualification (OQ)",
    "Performance Qualification (PQ)",
    "Data migration and reconciliation",
    "Final validation summary and release authorization",
]
all_text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        all_text += "\n" + " | ".join(c.text for c in row.cells)

missing = [term for term in required if term not in all_text]
fixed_heights = 0
bad_width_rows = 0
header_rows = 0
callout_tables = 0
for table in doc.tables:
    if len(table.rows) == 1 and len(table.columns) == 1:
        callout_tables += 1
    first_trpr = table.rows[0]._tr.trPr
    if first_trpr is not None and first_trpr.find(qn("w:tblHeader")) is not None:
        header_rows += 1
    grid = [int(c.get(qn("w:w"))) for c in table._tbl.tblGrid.gridCol_lst]
    for row in table.rows:
        widths = []
        for cell in row.cells:
            tcw = cell._tc.tcPr.find(qn("w:tcW"))
            widths.append(int(tcw.get(qn("w:w"))))
        if widths != grid:
            bad_width_rows += 1
        trh = row._tr.get_or_add_trPr().find(qn("w:trHeight"))
        if trh is not None and trh.get(qn("w:hRule")) == "exact":
            fixed_heights += 1

with ZipFile(path) as z:
    bad_zip = z.testzip()
    document_xml = z.read("word/document.xml")
    sect_ok = all(token in document_xml for token in [b'w:pgSz w:w="12240" w:h="15840"', b'w:pgMar w:top="1440"'])

assert not missing, f"Missing required sections: {missing}"
assert bad_width_rows == 0, f"Rows with mismatched widths: {bad_width_rows}"
assert fixed_heights == 0, f"Fixed-height rows: {fixed_heights}"
assert bad_zip is None, f"Bad zip member: {bad_zip}"
assert sect_ok, "Page geometry not encoded as expected"
assert header_rows == len(doc.tables) - callout_tables, (header_rows, len(doc.tables), callout_tables)

print({
    "paragraphs": len(doc.paragraphs),
    "headings": len(headings),
    "tables": len(doc.tables),
    "header_rows": header_rows,
    "callout_tables": callout_tables,
    "fixed_height_rows": fixed_heights,
    "bad_width_rows": bad_width_rows,
    "required_sections_missing": missing,
    "zip_integrity": "OK",
    "page_geometry": "Letter, 1-inch margins",
})
