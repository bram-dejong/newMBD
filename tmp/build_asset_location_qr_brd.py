from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "Asset_Location_QR_Audit_BRD.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
WHITE = "FFFFFF"
RISK = "9B1C1C"
GOLD = "7A5A00"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_cell_text(cell, bold=False, color=INK, size=9):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        hdr.cells[idx].text = header
        set_cell_shading(hdr.cells[idx], header_fill)
        style_cell_text(hdr.cells[idx], bold=True, color=DARK_BLUE, size=font_size)

    for data in rows:
        row = table.add_row()
        keep_row_together(row)
        for idx, value in enumerate(data):
            row.cells[idx].text = str(value)
            style_cell_text(row.cells[idx], size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def set_run_font(run, size=None, bold=None, italic=None, color=INK):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, label, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.82)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.167

styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.75)
styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.25)
styles["List Number"].paragraph_format.left_indent = Inches(0.5)
styles["List Number"].paragraph_format.first_line_indent = Inches(-0.25)

heading_tokens = {
    "Heading 1": (16, BLUE, 16, 8),
    "Heading 2": (13, BLUE, 12, 6),
    "Heading 3": (12, DARK_BLUE, 8, 4),
}
for name, (size, color, before, after) in heading_tokens.items():
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.paragraph_format.space_after = Pt(0)
left = hp.add_run("NEW MBD  |  BUSINESS REQUIREMENTS")
set_run_font(left, size=8.5, bold=True, color=MUTED)
right = hp.add_run("\tASSET LOCATION QR AUDIT")
set_run_font(right, size=8.5, color=MUTED)
tabs = hp.paragraph_format.tab_stops
tabs.add_tab_stop(Inches(6.45))

footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)

doc.core_properties.title = "Business Requirements Document — Asset Location QR Audit"
doc.core_properties.subject = "NewMBD QR and camera-assisted asset identification, location anomaly detection, and ERP notification"
doc.core_properties.author = "NewMBD Product and Service Operations"
doc.core_properties.keywords = "NewMBD, Salesforce, Asset, Location, QR, anomaly, ERP, BRD"

# Memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("BUSINESS REQUIREMENTS DOCUMENT")
set_run_font(r, size=10, bold=True, color=BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Asset Location QR Audit")
set_run_font(r, size=25, bold=True, color="000000")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
r = p.add_run("NewMBD service-engineer scanning, reconciliation, anomaly management, and ERP reporting")
set_run_font(r, size=13.5, color="4B5563")

for label, value in [
    ("Document owner", "NewMBD Product and Service Operations"),
    ("Prepared for", "Service Operations, ERP Operations, Salesforce Product Team, and Field Service"),
    ("Version", "0.2 - Camera identification MVP"),
    ("Date", "31 July 2026"),
    ("Target platform", "Salesforce NewMBD"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, bold=True, size=10.5)
    r2 = p.add_run(value)
    set_run_font(r2, size=10.5)

add_callout(
    doc,
    "Decision requested",
    "Approve the business process and minimum requirements for a Salesforce-native location audit in which a service engineer scans asset QR codes or confirms an Asset from its visible serial number, Salesforce identifies location discrepancies, anomalies are retained for resolution, and the responsible ERP resource receives one digest email per audit.",
)

add_heading(doc, "1. Executive Summary", 1)
add_body(
    doc,
    "Service engineers need a reliable way to verify which physical assets are present at a customer location and identify differences from the installed-base location recorded in Salesforce. The proposed NewMBD capability places a QR audit experience on the Salesforce Location record, supports scanning one or more assets, compares the observed inventory with Asset.Location__c, and records each discrepancy as an auditable anomaly."
)
add_body(
    doc,
    "The first release is a reconciliation and notification process. It does not automatically relocate assets. A completed audit creates a durable audit record, stores unknown, wrong-location, and—when the engineer declares a complete inventory—missing-asset anomalies, and sends one summarized notification to the ERP resource configured for the location."
)
add_body(
    doc,
    "Release 1 now includes a provider-ready camera-identification MVP. The engineer photographs the device, is guided to show the nameplate, enters or reviews the extracted serial, selects a matching Salesforce Asset, and explicitly confirms it. In the scratch org, the vision provider is deliberately unconfigured; capture, guidance, serial search, candidate confirmation, and audit reconciliation remain testable without pretending that AI extraction occurred."
)

add_heading(doc, "2. Business Context and Problem Statement", 1)
add_body(
    doc,
    "Installed-base location data can become inaccurate when equipment is relocated, replaced, transferred, incorrectly registered, or physically present without a corresponding Salesforce record. Today, field observations may be communicated through unstructured channels and may not create a consistent, reportable audit trail."
)
add_body(doc, "The business requires a process that:")
for item in [
    "Works in the engineer’s mobile Salesforce experience at the physical location.",
    "Uses a stable QR identifier to identify assets without manual searching.",
    "Provides a guided nameplate-photo path when the QR label is missing or inaccessible.",
    "Distinguishes a targeted spot check from a declared complete inventory.",
    "Preserves discrepancies without silently overwriting the system of record.",
    "Routes actionable evidence to the responsible ERP resource.",
    "Provides reporting and traceability from detection through resolution.",
]:
    add_bullet(doc, item)

add_heading(doc, "3. Business Objectives and Success Measures", 1)
add_table(
    doc,
    ["Objective", "Proposed success measure"],
    [
        ("Improve installed-base accuracy", "Reduction in unresolved wrong-location and unknown-asset exceptions over successive audit cycles."),
        ("Reduce manual reconciliation", "At least 90% of pilot audits completed without offline spreadsheets or manual asset searches."),
        ("Create accountable follow-up", "100% of detected discrepancies stored with location, audit, type, status, and evidence."),
        ("Notify the ERP owner promptly", "ERP digest queued immediately after audit completion; delivery result recorded for every anomaly."),
        ("Support field usability", "Median scan-and-submit time suitable for the agreed pilot inventory size; duplicate scans do not create duplicate evidence."),
        ("Provide operational visibility", "Reports identify open anomalies, notification failures, aging, and anomaly rates by location."),
    ],
    [3400, 5960],
)
add_callout(
    doc,
    "Baseline needed",
    "Final numeric targets for audit duration, anomaly resolution time, email success rate, and data-quality improvement must be confirmed during the pilot.",
    color=GOLD,
)

add_heading(doc, "4. Scope", 1)
add_heading(doc, "4.1 In Scope for Release 1", 2)
for item in [
    "Standard Salesforce Location records related to a parent Account.",
    "Standard Salesforce Assets related to a Location.",
    "A unique QR identifier stored on each scannable Asset.",
    "Continuous mobile QR scanning and manual-entry fallback.",
    "Transient device and nameplate photo capture in Salesforce Mobile and the browser utility.",
    "Serial-number candidate search and explicit engineer confirmation of the Salesforce Asset.",
    "A pluggable vision-provider interface with an honest unconfigured-provider fallback.",
    "Spot Check and Complete Inventory audit modes.",
    "Detection of Unknown Asset, Wrong Location, and Missing at Location anomalies.",
    "One audit record and one child anomaly record per detected discrepancy.",
    "One ERP digest email per completed audit that contains anomalies.",
    "Location-specific recipient configuration with an org-level fallback.",
    "Audit and anomaly reporting, status management, and notification outcome tracking.",
    "Role-based access through a dedicated permission set.",
]:
    add_bullet(doc, item)

add_heading(doc, "4.2 Out of Scope for Release 1", 2)
for item in [
    "Automatic modification of Asset.Location__c based solely on a scan.",
    "Direct ERP API integration or automated ERP master-data changes.",
    "Offline scan synchronization or resumable draft audits.",
    "QR label printing, procurement, and physical deployment.",
    "Automatic asset eligibility logic by lifecycle status or product class until business rules are approved.",
    "Automated anomaly approval, escalation, or remediation workflows beyond notification and status tracking.",
    "Historical migration of prior informal discrepancy reports.",
    "Production AI/OCR activation before an enterprise-approved provider, endpoint, credential, and privacy review are supplied.",
    "Saving nameplate photographs to Salesforce Files or including photographs in ERP email.",
]:
    add_bullet(doc, item)

add_heading(doc, "5. Stakeholders and Responsibilities", 1)
add_table(
    doc,
    ["Stakeholder", "Primary responsibility"],
    [
        ("Service Engineer", "Selects the physical Location, scans assets, chooses the audit mode, reviews the scan list, and submits the audit."),
        ("Service Operations", "Owns the audit process, engineer guidance, anomaly triage, and operational KPIs."),
        ("ERP Resource / ERP Operations", "Receives anomaly digests, validates ERP master data, and coordinates corrections according to system ownership."),
        ("Salesforce Product Team", "Owns NewMBD configuration, business logic, permissions, deployment, and support."),
        ("Installed Base / Data Steward", "Defines asset eligibility, QR governance, data remediation, and duplicate-prevention rules."),
        ("Account / Location Owner", "Confirms customer and site context and supports physical-location validation."),
        ("Information Security / Privacy", "Approves access, email content, retention, and regional data-handling controls."),
    ],
    [2600, 6760],
)

add_heading(doc, "6. Future-State Business Process", 1)
for step in [
    "The engineer opens the Asset Location QR Audit component on the relevant Salesforce Location record.",
    "The engineer selects Spot Check or Complete Inventory. Complete Inventory is used only when the engineer confirms that all in-scope assets at the location will be scanned.",
    "The engineer scans one or more QR codes. The interface normalizes and de-duplicates codes and allows removal or manual entry before submission.",
    "When a QR label cannot be used, the engineer chooses Identify from photo, photographs the whole device, then follows the prompt to show the serial-number nameplate.",
    "The system presents the extracted or manually entered serial number, searches Salesforce, and requires the engineer to select and confirm the matching Asset. Visual appearance alone never confirms a unique Asset.",
    "On submission, Salesforce validates the Location and compares the scanned asset identifiers with the assets assigned to that Location.",
    "Salesforce creates one completed audit record containing the mode, completion time, expected count, scanned count, anomaly count, and outcome.",
    "Salesforce creates anomaly records for each unknown code, asset assigned elsewhere, and—only for Complete Inventory—expected asset not represented in the scan.",
    "If anomalies exist, Salesforce queues one digest email to the Location recipient or configured fallback recipient.",
    "The anomaly owner reviews the evidence and moves each anomaly through Open, Confirmed, Resolved, or Dismissed.",
    "Authorized data owners correct Salesforce and/or ERP according to the approved system-of-record process; the audit evidence remains retained.",
]:
    add_numbered(doc, step)

add_heading(doc, "7. Business Requirements", 1)
add_heading(doc, "7.1 Location and Asset Foundation", 2)
add_table(
    doc,
    ["ID", "Requirement", "Priority"],
    [
        ("DAT-001", "The system shall relate each service Location to one parent Account.", "Must"),
        ("DAT-002", "The system shall allow each Asset to reference its registered service Location.", "Must"),
        ("DAT-003", "The system shall store a stable, unique, case-insensitive QR identifier for each scannable Asset.", "Must"),
        ("DAT-004", "The system shall retain Account context separately from physical Location and shall not use Asset.AccountId as the physical-location comparison key.", "Must"),
        ("DAT-005", "The business shall define which Asset statuses, types, and ownership conditions are eligible for Complete Inventory comparison.", "Must"),
        ("DAT-006", "The system shall prevent duplicate QR identifiers from being assigned to multiple Assets.", "Must"),
    ],
    [1050, 6900, 1410],
)

add_heading(doc, "7.2 Scanning Experience", 2)
add_table(
    doc,
    ["ID", "Requirement", "Priority"],
    [
        ("SCN-001", "The system shall provide an Asset QR scanning experience in the context of a Salesforce Location.", "Must"),
        ("SCN-002", "The mobile experience shall support continuous scanning of multiple QR codes.", "Must"),
        ("SCN-003", "The system shall provide manual code entry as a fallback for desktop use, camera limitations, and testing.", "Should"),
        ("SCN-004", "The client and server shall trim, normalize, and de-duplicate submitted codes case-insensitively.", "Must"),
        ("SCN-005", "The engineer shall be able to review and remove captured codes before final submission.", "Must"),
        ("SCN-006", "The system shall prevent submission when no valid code is present unless the approved Complete Inventory process explicitly permits an empty audit.", "Must"),
        ("SCN-007", "The system shall support a defined maximum number of unique codes per audit and provide a clear limit message.", "Must"),
        ("SCN-008", "The interface shall provide accessible success, progress, cancellation, and error feedback.", "Should"),
        ("SCN-009", "The system shall offer a guided Identify from photo workflow in Salesforce Mobile and the browser utility.", "Must"),
        ("SCN-010", "The workflow shall first capture the device context and then request a close, readable image of the serial-number label.", "Must"),
        ("SCN-011", "An AI or OCR result shall be advisory and shall not identify a unique Asset from appearance alone.", "Must"),
        ("SCN-012", "The engineer shall be able to edit the extracted serial number and search Salesforce for exact serial matches.", "Must"),
        ("SCN-013", "The system shall require explicit engineer selection and confirmation when one or more Asset candidates are returned.", "Must"),
        ("SCN-014", "A camera-confirmed Asset shall be submitted by Asset ID and shall remain auditable even when it has no QR value.", "Must"),
        ("SCN-015", "Images shall be resized, validated, processed transiently, and discarded without being stored or emailed.", "Must"),
        ("SCN-016", "When AI is unavailable or unconfigured, the interface shall say so and provide manual serial entry without simulating a result.", "Must"),
    ],
    [1050, 6900, 1410],
)

add_heading(doc, "7.3 Reconciliation and Anomaly Detection", 2)
add_table(
    doc,
    ["ID", "Requirement", "Priority"],
    [
        ("REC-001", "The engineer shall choose Spot Check or Complete Inventory before submitting an audit.", "Must"),
        ("REC-002", "Spot Check shall evaluate only submitted QR codes and shall not create Missing at Location anomalies.", "Must"),
        ("REC-003", "Complete Inventory shall compare submitted codes with all eligible Assets assigned to the selected Location.", "Must"),
        ("REC-004", "A scanned code that does not match an Asset shall create an Unknown Asset anomaly.", "Must"),
        ("REC-005", "A matched Asset whose registered Location differs from the selected Location shall create a Wrong Location anomaly.", "Must"),
        ("REC-006", "An eligible Asset assigned to the selected Location but not represented in a Complete Inventory scan shall create a Missing at Location anomaly.", "Must"),
        ("REC-007", "The system shall not update Asset location automatically as a result of anomaly detection.", "Must"),
        ("REC-008", "Reconciliation shall be bulk-safe and shall not perform a query or email operation for each scanned code.", "Must"),
    ],
    [1050, 6900, 1410],
)

add_heading(doc, "7.4 Audit and Anomaly Records", 2)
add_table(
    doc,
    ["ID", "Requirement", "Priority"],
    [
        ("AUD-001", "Every finalized submission shall create one Asset Location Audit record.", "Must"),
        ("AUD-002", "The audit shall record Location, mode, completed timestamp, scanned count, expected count, anomaly count, and outcome.", "Must"),
        ("AUD-003", "Each detected discrepancy shall create one anomaly linked to its audit.", "Must"),
        ("AUD-004", "An anomaly shall retain type, scanned Location, Salesforce Location when known, Asset when known, scanned code, details, status, and notification outcome.", "Must"),
        ("AUD-005", "Anomaly status shall support Open, Confirmed, Resolved, and Dismissed.", "Must"),
        ("AUD-006", "Completed audit evidence shall remain available after an Asset or Location is corrected.", "Must"),
        ("AUD-007", "The system shall preserve creator, timestamps, and field history required for auditability.", "Must"),
        ("AUD-008", "Production hardening shall prevent replayed submissions from creating duplicate audits or duplicate emails.", "Should"),
    ],
    [1050, 6900, 1410],
)

add_heading(doc, "7.5 ERP Notification", 2)
add_table(
    doc,
    ["ID", "Requirement", "Priority"],
    [
        ("NTF-001", "When an audit contains anomalies, the system shall queue one digest email after the audit transaction succeeds.", "Must"),
        ("NTF-002", "The Location-specific ERP recipient shall take precedence over the org-level fallback recipient.", "Must"),
        ("NTF-003", "The digest shall identify the audit, scanned Location, anomaly type, Asset or scanned code, and recorded Salesforce Location.", "Must"),
        ("NTF-004", "The system shall record notification success and timestamp on the anomaly records only after a successful send.", "Must"),
        ("NTF-005", "If no recipient is configured or email fails, the system shall retain the anomalies and record the notification error.", "Must"),
        ("NTF-006", "An audit with no anomalies shall not generate an ERP anomaly email.", "Must"),
        ("NTF-007", "Non-production environments shall suppress or reroute external emails according to environment controls.", "Must"),
        ("NTF-008", "If ERP automation parses the email, the message format shall be versioned and structured; an API or middleware interface remains the preferred future state.", "Should"),
    ],
    [1050, 6900, 1410],
)

add_heading(doc, "8. Anomaly Rules", 1)
add_table(
    doc,
    ["Anomaly", "Trigger", "Applicable mode", "Required evidence"],
    [
        ("Unknown Asset", "Normalized code matches no Asset QR identifier.", "Spot Check and Complete Inventory", "Audit, scanned Location, submitted code, explanatory detail."),
        ("Wrong Location", "Matched Asset.Location__c differs from the scanned Location, including a blank registered Location.", "Spot Check and Complete Inventory", "Audit, Asset, scanned Location, Salesforce Location when present, code."),
        ("Missing at Location", "Eligible Asset assigned to the scanned Location is absent from the finalized scan.", "Complete Inventory only", "Audit, Asset, scanned Location, Asset QR code when available."),
    ],
    [1700, 3100, 1900, 2660],
)
add_callout(
    doc,
    "Control",
    "Complete Inventory is an explicit engineer declaration. It must not be inferred from the number of scans, because doing so could create false missing-asset anomalies.",
)

add_heading(doc, "9. Data Model and Data Requirements", 1)
add_table(
    doc,
    ["Record", "Relationship / key data", "Business purpose"],
    [
        ("Account", "Parent of Location", "Customer context and ownership hierarchy."),
        ("Location", "Account__c; ERP_Anomaly_Email__c", "Physical audit point and notification routing."),
        ("Asset", "Location__c; Asset_QR_Code__c; SerialNumber; Product2Id; AccountId", "Installed-base record, serial match, and canonical location comparison."),
        ("Asset Location Audit", "Location; mode; timestamps; counts; status", "One record per finalized scan session."),
        ("Asset Location Anomaly", "Master-detail to Audit; optional Asset; observed and recorded Locations; evidence; resolution status", "Durable discrepancy and follow-up record."),
        ("Asset Audit Setting", "Default ERP recipient", "Org-level fallback configuration."),
    ],
    [1900, 3250, 4210],
)
add_heading(doc, "9.1 Data Quality Preconditions", 2)
for item in [
    "Every Location in pilot scope has the correct parent Account.",
    "Every eligible Asset has the correct Location and customer Account context.",
    "Every scannable Asset has one unique QR identifier that matches its physical label.",
    "Duplicate identifiers are resolved before data load.",
    "Location recipients are valid, approved, and tested in a non-production environment.",
    "Assets lacking a QR code are identified before Complete Inventory use and handled according to the approved eligibility rule.",
]:
    add_bullet(doc, item)

add_heading(doc, "10. Security, Privacy, and Controls", 1)
add_table(
    doc,
    ["ID", "Requirement", "Priority"],
    [
        ("SEC-001", "Only authorized users shall access the scanner, eligible Locations, Assets, audits, and anomalies.", "Must"),
        ("SEC-002", "Service engineers shall receive read access to Location and Asset data and create/read access needed for audits and anomalies.", "Must"),
        ("SEC-003", "The server-side process shall enforce record sharing and field/object permissions.", "Must"),
        ("SEC-004", "QR values shall be treated as untrusted input and shall not contain credentials or sensitive personal information.", "Must"),
        ("SEC-005", "Email content shall contain only information necessary for ERP reconciliation.", "Must"),
        ("SEC-006", "Audit visibility and organization-wide defaults shall be reviewed before production; Private is the recommended default.", "Must"),
        ("SEC-007", "Edits to completed evidence and notification-control fields shall be restricted to approved support or integration roles.", "Must"),
        ("SEC-008", "Retention, regional privacy, and works-council requirements shall be confirmed before broad rollout.", "Must"),
        ("SEC-009", "Photo bytes shall not be persisted, logged, attached to records, or included in ERP notification for the MVP.", "Must"),
        ("SEC-010", "A production vision service shall use an approved Named Credential and External Credential; secrets shall not be stored in LWC, Apex, or custom metadata.", "Must"),
        ("SEC-011", "The server shall re-query every camera-confirmed Asset under sharing and field-level security before finalization.", "Must"),
    ],
    [1050, 6900, 1410],
)

add_heading(doc, "11. Reporting Requirements", 1)
add_table(
    doc,
    ["ID", "Requirement", "Priority"],
    [
        ("RPT-001", "Users shall report audits by date, engineer, Account, Location, mode, counts, and outcome.", "Should"),
        ("RPT-002", "Users shall report anomalies by type, status, age, Asset, Location, and responsible operating unit.", "Should"),
        ("RPT-003", "Users shall identify notification failures and anomalies that have not been sent successfully.", "Must"),
        ("RPT-004", "Operational dashboards should show anomaly rates, resolution aging, audit volume, and repeat anomalies by Location.", "Should"),
        ("RPT-005", "Data-quality reporting should identify eligible Assets with blank Location, missing QR code, or duplicate/invalid identifiers.", "Should"),
    ],
    [1050, 6900, 1410],
)

add_heading(doc, "12. Non-Functional Requirements", 1)
add_table(
    doc,
    ["Area", "Requirement"],
    [
        ("Performance", "The submit operation shall use a bounded number of database queries and DML operations and support the agreed maximum codes per audit."),
        ("Reliability", "Email failure shall not roll back the completed audit or its anomalies."),
        ("Usability", "The mobile workflow shall minimize typing, show captured codes clearly, and distinguish normal scanner cancellation from errors."),
        ("Accessibility", "Controls, status messages, and error messages shall be keyboard- and assistive-technology friendly where supported."),
        ("Compatibility", "QR scanning shall support Salesforce Mobile; photo capture, guided serial entry, and confirmation shall support Salesforce Mobile and the desktop utility."),
        ("Privacy", "Nameplate images shall be compressed, transmitted once for analysis, and discarded without persistence in the MVP."),
        ("Auditability", "CreatedBy, timestamps, status history, and notification results shall remain reportable."),
        ("Supportability", "Configuration shall avoid hard-coded recipient addresses and shall expose actionable failure details."),
        ("Scalability", "The design shall support extension to middleware/API notification without replacing the core audit and anomaly model."),
    ],
    [2100, 7260],
)

add_heading(doc, "13. Acceptance Criteria", 1)
criteria = [
    ("AC-001", "A user can open the audit component on a Location and scan or manually enter multiple codes."),
    ("AC-002", "Duplicate codes differing only by whitespace or letter case are counted once."),
    ("AC-003", "A clean Spot Check creates one completed audit, zero anomalies, and no email."),
    ("AC-004", "An unmatched code creates an Unknown Asset anomaly."),
    ("AC-005", "An Asset registered at another or blank Location creates a Wrong Location anomaly."),
    ("AC-006", "A Spot Check never creates a Missing at Location anomaly."),
    ("AC-007", "A Complete Inventory creates a Missing at Location anomaly for each eligible expected Asset not scanned."),
    ("AC-008", "The audit counts equal the normalized scan set, expected eligible population, and created anomalies."),
    ("AC-009", "An anomaly digest uses the Location recipient before the fallback recipient."),
    ("AC-010", "Missing recipient and email failure preserve the audit/anomalies and record a notification error."),
    ("AC-011", "Authorized users can report and update anomaly resolution status; unauthorized users cannot run the process."),
    ("AC-012", "No Asset location is automatically changed by submitting an audit."),
    ("AC-013", "The agreed automated test suite passes at or above the Salesforce production coverage threshold."),
    ("AC-014", "A user can start Identify from photo, capture a device image, and receive guidance to show the serial-number label."),
    ("AC-015", "When the vision provider is unconfigured, the UI states that clearly and allows manual serial entry."),
    ("AC-016", "An exact serial search returns only Assets visible to the user and requires explicit candidate confirmation."),
    ("AC-017", "A confirmed Asset without a QR value is included in the audit by Asset ID and reconciled correctly."),
    ("AC-018", "Duplicate QR and camera observations of the same Asset are reconciled once by the server."),
    ("AC-019", "Unsupported, unreadable, or oversized images produce recoverable errors and do not persist photo data."),
    ("AC-020", "No captured photo is stored in Salesforce Files, audit/anomaly records, logs, or ERP email."),
]
add_table(doc, ["ID", "Acceptance criterion"], criteria, [1200, 8160])

add_heading(doc, "14. Rollout and Change Management", 1)
for step in [
    "Approve system-of-record ownership, asset eligibility, QR governance, email content, and anomaly ownership.",
    "Profile pilot data for missing/duplicate QR codes, blank or invalid Locations, and inactive Assets.",
    "Backfill Location-to-Account and Asset-to-Location relationships and validate a small set against known physical inventory.",
    "Populate non-production ERP recipients and confirm Salesforce email deliverability and rerouting controls.",
    "Configure the Location Lightning page, permission assignments, anomaly layouts, list views, reports, and dashboards.",
    "Pilot with a small number of verified Locations and train participating service engineers.",
    "Measure scan duration, anomaly accuracy, notification delivery, and resolution aging; correct false-positive rules.",
    "Roll out by region or business unit with documented support, retry, and escalation procedures.",
]:
    add_numbered(doc, step)

add_heading(doc, "15. Risks and Mitigations", 1)
add_table(
    doc,
    ["Risk", "Business impact", "Mitigation"],
    [
        ("Incorrect or incomplete Asset.Location__c data", "High initial anomaly volume and reduced confidence.", "Profile and remediate pilot data; treat early results as a controlled baseline."),
        ("Assets without valid QR labels", "Complete Inventory produces avoidable missing exceptions.", "Define eligibility, label assets, and run a data-quality report before enabling complete audits."),
        ("Engineer selects Complete Inventory prematurely", "False missing-asset reports and unnecessary ERP work.", "Use explicit confirmation, training, and clear mode descriptions."),
        ("Incorrect ERP recipient", "Customer/asset information sent to the wrong party or no action taken.", "Validate recipients, use environment rerouting, and record send outcomes."),
        ("Email treated as machine interface", "Fragile parsing and integration failures.", "Use a versioned structured format temporarily and plan middleware/API integration."),
        ("Duplicate submission", "Duplicate anomalies and repeated emails.", "Add a client request key and server-side idempotency before production scale."),
        ("Broad audit sharing", "Installed-base information exposed beyond need-to-know.", "Review OWD and sharing rules; adopt Private by default."),
        ("Legacy NewMBD metadata dependencies", "Scratch-org and CI rebuilds fail outside configured environments.", "Maintain a dependency-safe deployment manifest and remove or isolate managed-package metadata."),
        ("AI provider not approved or configured", "Automated serial extraction cannot run in production.", "Keep the provider adapter disabled, disclose the fallback, and complete security/privacy approval before credential activation."),
        ("Incorrect OCR or visual match", "The wrong installed-base Asset could be submitted.", "Require a readable serial, exact server-side candidate lookup, and explicit engineer confirmation."),
        ("Customer/nameplate photo exposure", "Sensitive site or equipment information could be retained or transmitted unnecessarily.", "Resize and process transiently; do not persist, log, or email images in the MVP."),
    ],
    [2700, 3000, 3660],
    font_size=8.5,
)

add_heading(doc, "16. Assumptions and Dependencies", 1)
add_heading(doc, "16.1 Assumptions", 2)
for item in [
    "Standard Salesforce Location is the NewMBD physical-location entity.",
    "Location.Account__c represents the customer hierarchy and Asset.Location__c is the canonical Salesforce location used for comparison.",
    "The QR value is an opaque, stable Asset identifier and is unique after case-insensitive normalization.",
    "The engineer submits the audit while online.",
    "Email is mandatory for Release 1, even if an API or middleware interface is preferred later.",
    "A scan is evidence of physical presence but not sufficient authority to relocate an Asset automatically.",
    "A visual classification is advisory; the visible serial number and engineer confirmation are the authoritative camera-identification controls.",
]:
    add_bullet(doc, item)

add_heading(doc, "16.2 Dependencies", 2)
for item in [
    "Salesforce mobile access to the barcode scanner and device camera permissions.",
    "NewMBD Location, Asset, audit, anomaly, custom-metadata, Apex, Lightning component, and permission-set metadata.",
    "A Location Lightning record page containing the Asset Location QR Audit component.",
    "Asset and Location master-data preparation and QR label governance.",
    "Salesforce email deliverability, approved sender configuration, and valid ERP recipient addresses.",
    "An agreed anomaly owner, resolution SLA, and ERP/Salesforce data-correction procedure.",
    "A repeatable dependency-safe NewMBD deployment process for scratch, validation, and production environments.",
    "An enterprise-approved vision/OCR provider, Named Credential, External Credential, and data-processing review before AI extraction is activated.",
]:
    add_bullet(doc, item)

add_heading(doc, "17. Open Business Decisions", 1)
add_table(
    doc,
    ["Decision", "Recommended default", "Owner"],
    [
        ("Which Assets are eligible for Complete Inventory?", "Active, installed, independently auditable Assets only.", "Service Operations / Data Steward"),
        ("How are expected Assets without QR codes handled?", "Exclude from missing comparison and report as data-quality backlog after implementation alignment.", "Service Operations"),
        ("Who owns Asset location master data?", "ERP when the installed-base integration is ERP-led; otherwise explicitly name Salesforce ownership.", "ERP and Salesforce Product Owners"),
        ("May a confirmed anomaly trigger relocation?", "Not automatically in Release 1; use a reviewed resolution action.", "Service Operations"),
        ("What is the anomaly response SLA?", "Define by severity and region during pilot.", "Service Operations"),
        ("Is email human-readable or machine-parsed?", "Human-readable digest for Release 1; API/middleware for automation.", "ERP Operations"),
        ("What is the audit retention period?", "Apply the enterprise service-data retention standard.", "Privacy / Records Management"),
        ("What visibility model applies?", "Private audits with role-based sharing unless reporting requires broader access.", "Salesforce Security Owner"),
        ("Which vision/OCR provider may process nameplate images?", "Use only the enterprise-approved provider through Named/External Credential; keep the adapter disabled until approved.", "Information Security / Salesforce Product Owner"),
        ("May nameplate photos be retained?", "No for the MVP; revisit only with explicit business purpose, consent, retention, encryption, and legal-hold controls.", "Privacy / Records Management"),
    ],
    [3100, 3900, 2360],
    font_size=8.5,
)

add_heading(doc, "18. Approval", 1)
add_body(doc, "Approval confirms the business process, scope, priorities, and acceptance criteria. It does not authorize production deployment until security, data readiness, email controls, and the open business decisions above are completed.")
add_table(
    doc,
    ["Role", "Name", "Decision / signature", "Date"],
    [
        ("Service Operations Owner", "", "", ""),
        ("ERP Operations Owner", "", "", ""),
        ("Salesforce Product Owner", "", "", ""),
        ("Information Security / Privacy", "", "", ""),
    ],
    [2400, 2200, 2960, 1800],
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
