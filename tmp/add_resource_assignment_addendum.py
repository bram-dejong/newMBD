from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "Salesforce_Org_Resources_and_Skills_Guide.docx"
OUTPUT = ROOT / "output" / "Salesforce_Org_Resources_and_Skills_Guide_with_Resource_Addendum.docx"

NAVY = "17365D"
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
TEXT = "222222"
WHITE = "FFFFFF"


def set_font(run, size=11, color=TEXT, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is None:
        existing = OxmlElement("w:tblHeader")
        tr_pr.append(existing)
    existing.set(qn("w:val"), "true")


def set_cell_text(cell, text, *, size=9.2, color=TEXT, bold=False, align=None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), size=size, color=color, bold=bold)


roles = [
    ("Leadership & governance", "Executive sponsor"),
    ("Leadership & governance", "Business product owner"),
    ("Leadership & governance", "Program / project manager"),
    ("Architecture & delivery", "Salesforce solution architect"),
    ("Architecture & delivery", "Business analyst / process designer"),
    ("Architecture & delivery", "Salesforce administrator"),
    ("Architecture & delivery", "Salesforce developer"),
    ("Architecture & delivery", "Integration architect / developer"),
    ("Architecture & delivery", "Data architect / migration specialist"),
    ("Architecture & delivery", "QA / test lead"),
    ("Architecture & delivery", "DevOps / release engineer"),
    ("Business workstreams", "Sales Cloud business lead / SME"),
    ("Business workstreams", "Service Cloud business lead / SME"),
    ("Business workstreams", "Revenue Cloud / Revenue Management lead"),
    ("Business workstreams", "Product catalog and pricing owner"),
    ("Business workstreams", "Order, contract and renewals owner"),
    ("Finance & controls", "Finance / billing owner"),
    ("Finance & controls", "Tax and payments specialist"),
    ("Finance & controls", "Legal / contracting representative"),
    ("Security & adoption", "Security / privacy specialist"),
    ("Security & adoption", "Change and adoption lead"),
    ("Security & adoption", "Training lead / business change champion"),
]


doc = Document(SOURCE)

# Add a clean page break without changing the source document's existing sections.
page_break = doc.add_paragraph()
page_break.add_run().add_break(WD_BREAK.PAGE)

title = doc.add_paragraph(style="Heading 1")
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(6)
title.paragraph_format.keep_with_next = True
set_font(title.add_run("Addendum A - Resource Assignment Register"), size=16, color=BLUE, bold=True)

intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(10)
intro.paragraph_format.line_spacing = 1.2
set_font(intro.add_run(
    "Use this register to record the named resources assigned to the Salesforce implementation. "
    "Add multiple names in a cell where a role is shared, and use the final column for allocation, availability, or assignment notes."
))

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
widths = [2350, 2900, 2100, 2010]  # 9360 DXA total
headers = ["Workstream / role", "Assigned resource name(s)", "Business area / supplier", "Allocation / notes"]

for idx, header in enumerate(headers):
    cell = table.rows[0].cells[idx]
    shade_cell(cell, NAVY)
    set_cell_text(cell, header, size=9.2, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
repeat_header(table.rows[0])

last_category = None
for index, (category, role) in enumerate(roles):
    cells = table.add_row().cells
    if index % 2:
        for cell in cells:
            shade_cell(cell, LIGHT_GRAY)
    if category != last_category:
        role_text = f"{category}\n{role}"
        last_category = category
    else:
        role_text = role
    set_cell_text(cells[0], role_text, size=9.0, bold=True)
    set_cell_text(cells[1], "")
    set_cell_text(cells[2], "")
    set_cell_text(cells[3], "")
    cells[1].paragraphs[0].paragraph_format.space_after = Pt(5)
    cells[2].paragraphs[0].paragraph_format.space_after = Pt(5)
    cells[3].paragraphs[0].paragraph_format.space_after = Pt(5)
    table.rows[-1].height = Inches(0.36)
    table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

set_table_geometry(table, widths)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(8)
note.paragraph_format.space_after = Pt(6)
note.paragraph_format.line_spacing = 1.15
set_font(note.add_run("Register owner: "), size=9.5, color=MID_GRAY, bold=True)
set_font(note.add_run("____________________________    "), size=9.5, color=MID_GRAY)
set_font(note.add_run("Last updated: "), size=9.5, color=MID_GRAY, bold=True)
set_font(note.add_run("____________________________"), size=9.5, color=MID_GRAY)

doc.core_properties.title = "Resources & Skills for a New Salesforce Org - with Resource Assignment Addendum"
doc.save(OUTPUT)
print(OUTPUT)
