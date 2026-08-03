from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "newMBD_SAP2CASE_Only_Business_Requirements_Document_v0.2.docx"

# standard_business_brief preset + named memo-masthead overrides
FONT = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20262E"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF8E8"
GOLD = "7A5A00"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=11, color=INK, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_text(cell, text, bold=False, color=INK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE, size=font_size)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value, size=font_size)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_numbering_definition(doc, num_id, abstract_id, fmt, text, left=720, hanging=360):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_bullet(doc, text, num_id=50):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run_font(p.add_run(text))
    return p


def add_step(doc, text, num_id=51):
    return add_bullet(doc, text, num_id=num_id)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(TABLE_INDENT / 1440)
    p.paragraph_format.right_indent = Inches(TABLE_INDENT / 1440)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), MID_GRAY)
        borders.append(border)
    p_pr.append(borders)
    set_run_font(p.add_run(label + " "), size=10.5, color=accent, bold=True)
    set_run_font(p.add_run(text), size=10.5, color=INK)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    settings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("newMBD | Business Requirements Document"), size=9, color=MUTED, bold=True)
    set_run_font(p.add_run("    SAP2CASE inbound email automation"), size=9, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("Internal working draft  |  Page "), size=9, color=MUTED)
    add_page_field(p)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_numbering_definition(doc, 50, 50, "bullet", "•")
    add_numbering_definition(doc, 51, 51, "decimal", "%1.")

    # Memo masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("BUSINESS REQUIREMENTS DOCUMENT"), size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run("newMBD SAP2CASE Inbound Email Automation"), size=24, color="000000", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    set_run_font(p.add_run("Optimized creation of MBD service Cases from structured SAP notification emails"), size=13, color=MUTED)

    metadata = [
        ("Document status", "Draft for business and technical review"),
        ("Version", "0.2 - SAP2CASE-only revision"),
        ("Date", "3 August 2026"),
        ("Prepared for", "newMBD / SDG Service Applications"),
        ("Environment status", "Implemented and tested in newMBD scratch org"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(f"{label}: "), size=10.5, bold=True)
        set_run_font(p.add_run(value), size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_callout(
        doc,
        "Decision summary.",
        "Create a newMBD-specific SAP2CASE automation that processes MBD inbound notifications, preserves the service issue narrative, and resolves the appropriate newMBD Account, Asset, Location, Contact, and owner records.",
    )

    add_heading(doc, "1. Purpose and business context", 1)
    add_body(doc, "This Business Requirements Document defines the intended business behavior, scope, controls, data mappings, acceptance criteria, and rollout decisions for the newMBD SAP2CASE inbound email process. It translates a structured SAP service notification email into a consistently populated Salesforce Case.")
    add_body(doc, "A dedicated newMBD automation is required because the target data model uses Salesforce Account, Asset, Location, Contact, User, and Case records. The process must preserve the diagnostic narrative, resolve trusted SAP identifiers, and apply MBD-specific Case defaults in a supportable and testable way.")

    add_heading(doc, "2. Business objectives and outcomes", 1)
    objectives = [
        "Create an MBD service Case automatically from a valid structured SAP2CASE email.",
        "Preserve the full human-readable issue description while separating SAP metadata from the Case description.",
        "Associate the Case with the correct customer Account, Asset, Location, Contact, and active owner when trusted identifiers produce an unambiguous match.",
        "Prevent non-MBD SAP notifications and ordinary inbound emails from being altered by this automation.",
        "Provide a dedicated and maintainable SAP2CASE intake process for newMBD.",
        "Provide a bulk-safe and testable automation that can be supported independently.",
    ]
    for item in objectives:
        add_bullet(doc, item)

    add_heading(doc, "3. Scope", 1)
    add_heading(doc, "3.1 In scope", 2)
    for item in [
        "Inbound Salesforce Case creation initiated by an email whose Case subject contains SAP2CASE.",
        "Recognition of the SAP payload beginning at the Notification line, including forwarded Outlook headers.",
        "Processing only when the payload contains Division: MBD.",
        "Parsing and mapping Notification, Subject, Body, AccountId, Owner, Equipment, Type, Division, ContactId, Material, Symptom, and Status labels.",
        "Bulk resolution of Account, Asset, Location, Contact, and User records.",
        "Defaulting Case origin, priority, source name, division, and auto-entitlement flag.",
        "Automated tests for the supplied example, scope guards, and a 200-record bulk scenario.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2 Out of scope", 2)
    for item in [
        "Changes to SAP email generation or the upstream Outlook forwarding rule.",
        "Production deployment, production data migration, and production email routing activation.",
        "Business mapping for Material and SAP Status until target fields and ownership are approved.",
        "Changes to downstream entitlement processing beyond setting the approved auto-entitlement flag.",
        "Operational exception dashboards or notifications for unmatched SAP identifiers.",
        "Inbound service processes for divisions other than MBD.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Stakeholders and ownership", 1)
    add_table(
        doc,
        ["Stakeholder", "Business interest / responsibility"],
        [
            ("MBD Service Operations", "Owns the service intake outcome, validates Case usability, and approves business rules."),
            ("SAP Service Process Owner", "Owns the notification format and the reliability of identifiers supplied in the email."),
            ("newMBD Product Owner", "Approves scope, prioritization, release readiness, and open mapping decisions."),
            ("Salesforce / Service Application Team", "Owns implementation, deployment, monitoring, and technical support."),
            ("Data Migration / Master Data Team", "Populates Account, Asset, Contact, and Location identifiers used for matching."),
            ("Service Desk / Case Owners", "Review Cases where a lookup cannot be resolved and complete missing assignments."),
        ],
        [2600, 6760],
    )

    add_heading(doc, "5. Future-state business process", 1)
    process_steps = [
        "SAP produces a structured service notification email and routes it to the approved newMBD inbound address.",
        "Salesforce Email-to-Case creates a preliminary Case using the inbound email subject and body.",
        "The newMBD trigger evaluates the Case before insert. It continues only when the Case subject contains SAP2CASE and a description is present.",
        "The handler ignores forwarded mail headers until it reaches the Notification line, which marks the start of the SAP payload.",
        "The handler parses the embedded subject, multiline issue body, and SAP metadata. It continues only when Division is MBD.",
        "Account, Asset, Contact, and active User records are resolved in bulk using SAP identifiers or the supplied owner name.",
        "The Case is populated with business fields and defaults. If an optional lookup is not resolved unambiguously, the Case remains available for operational review rather than being linked to an uncertain record.",
        "Downstream newMBD Case automation, including entitlement processing, continues after the Case is inserted.",
    ]
    for step in process_steps:
        add_step(doc, step)

    add_heading(doc, "6. Business requirements", 1)
    requirements = [
        ("BR-001", "The solution shall process only inbound Cases whose subject contains SAP2CASE and whose description contains a recognizable SAP payload."),
        ("BR-002", "The solution shall use the Notification line as the boundary between forwarded email headers and the structured SAP payload."),
        ("BR-003", "The solution shall process only payloads identified as Division MBD."),
        ("BR-004", "The embedded SAP subject shall replace the transport-level email subject on the resulting Case."),
        ("BR-005", "The complete multiline business issue body shall be retained in Case Description, including meaningful colons and line breaks."),
        ("BR-006", "SAP metadata appearing after the Body section shall not remain mixed into the Case Description."),
        ("BR-007", "The Account shall be resolved using the supplied AccountId and Account.SVMX_ERP_Account_ID__c when a match exists."),
        ("BR-008", "The Asset shall be resolved using the supplied Equipment value and Asset.ERP_Equipment_ID__c when a match exists."),
        ("BR-009", "A resolved Asset shall populate both Case.AssetId and Case.Component__c and shall supply Case.Site__c from the Asset Location."),
        ("BR-010", "When the SAP AccountId does not resolve but the Asset resolves, the Asset Account shall be used as the Case Account fallback."),
        ("BR-011", "The Contact shall be resolved using Contact.SAPExternalId__c; ambiguous matches shall be accepted only when exactly one candidate belongs to the resolved Account."),
        ("BR-012", "The owner shall be changed only when the supplied owner name resolves to exactly one active Salesforce User."),
        ("BR-013", "When an owner is resolved, the Case currency shall follow the owner's currency."),
        ("BR-014", "The solution shall set Supplied Name to SAP2CASE, Origin to Email, Priority to Medium, Division to MBD, and Perform Auto Entitlement to true."),
        ("BR-015", "The SAP Type and Symptom values shall populate the corresponding newMBD Case fields within their supported lengths."),
        ("BR-016", "Non-MBD SAP2CASE emails and normal inbound emails shall remain unchanged by this automation."),
        ("BR-017", "The solution shall support a full Salesforce trigger batch without SOQL queries inside record-processing loops."),
        ("BR-018", "A failed optional lookup shall not cause an incorrect Account, Asset, Contact, Location, or User association."),
    ]
    add_table(doc, ["ID", "Business requirement"], requirements, [1200, 8160], font_size=8.8)

    add_heading(doc, "7. Data mapping", 1)
    mapping_rows = [
        ("Notification", "Case.CAPA_ID__c", "Direct text mapping; maximum 25 characters."),
        ("Embedded Subject", "Case.Subject", "Replaces forwarding/transport subject; maximum 255 characters."),
        ("Body", "Case.Description", "Multiline content retained; SAP trailer fields removed; maximum 32,000 characters."),
        ("AccountId", "Case.AccountId", "Lookup by Account.SVMX_ERP_Account_ID__c."),
        ("Equipment", "Case.AssetId and Case.Component__c", "Lookup by Asset.ERP_Equipment_ID__c."),
        ("Resolved Asset Location", "Case.Site__c", "Copied from Asset.Location__c."),
        ("ContactId", "Case.ContactId", "Lookup by Contact.SAPExternalId__c with Account-aware ambiguity handling."),
        ("Owner", "Case.OwnerId", "Exact active User name; applied only for one unambiguous match."),
        ("Resolved owner currency", "Case.CurrencyIsoCode", "Copied from the resolved User."),
        ("Type", "Case.Type", "Direct text/picklist value; maximum 40 characters."),
        ("Division", "Case.Division__c", "Validated as MBD and stored as MBD."),
        ("Symptom", "Case.Symptom__c", "Direct text mapping; maximum 18 characters."),
        ("Material", "Not mapped", "Parsed for future use; target field requires business approval."),
        ("Status", "Not mapped", "Parsed for future use; SAP status-to-Case status translation requires approval."),
    ]
    add_table(doc, ["Inbound value", "newMBD target", "Rule"], mapping_rows, [2100, 2800, 4460], font_size=8.4)

    add_heading(doc, "8. Business rules and exception handling", 1)
    rules = [
        "A Case is eligible only when its transport subject contains SAP2CASE and a nonblank description is present.",
        "The first Notification line starts the SAP payload. A Subject line before Notification is treated as a forwarded email header and ignored.",
        "Body content continues until a recognized trailer label is encountered. Unknown labels and colons inside the issue narrative remain part of the body.",
        "Only Division MBD is transformed. Other divisions remain untouched for their owning process to handle.",
        "Lookup failures are nonfatal. Existing Email-to-Case ownership or blank optional lookups remain in place for manual review.",
        "Multiple Contact or active User matches are not resolved unless the Contact can be uniquely selected for the resolved Account.",
        "Material and SAP Status are not written to Salesforce until approved target fields and translations are defined.",
    ]
    for item in rules:
        add_bullet(doc, item)

    add_heading(doc, "9. Non-functional requirements", 1)
    nfr_rows = [
        ("NFR-001", "Performance", "The handler shall use no more than four bulk lookup queries for a trigger batch."),
        ("NFR-002", "Scalability", "The handler shall support 200 inbound Case records in one invocation."),
        ("NFR-003", "Data integrity", "Ambiguous lookup results shall not create a potentially incorrect relationship."),
        ("NFR-004", "Maintainability", "SAP parsing and record resolution shall be implemented as a dedicated newMBD handler."),
        ("NFR-005", "Testability", "Automated tests shall cover sample mapping, scope guards, and bulk processing."),
        ("NFR-006", "Observability", "Deployment and test outcomes shall be retained as release evidence; operational exception reporting is a rollout decision."),
        ("NFR-007", "Security", "The trigger shall operate within Salesforce platform controls; administrative field access shall be governed through approved profiles or permission sets."),
    ]
    add_table(doc, ["ID", "Category", "Requirement"], nfr_rows, [1200, 1700, 6460], font_size=8.6)

    add_heading(doc, "10. Acceptance criteria and UAT", 1)
    uat_rows = [
        ("UAT-01", "Supplied MBD sample email", "Case is created with embedded subject, clean multiline description, notification, Account, Contact, Asset, Location, Type, Division, Symptom, defaults, and auto-entitlement flag."),
        ("UAT-02", "Forwarded Outlook headers present", "Outer Subject is ignored; embedded Subject after Notification is used."),
        ("UAT-03", "Body contains 'Summary of issues: 0'", "The line remains in Case Description and is not treated as SAP metadata."),
        ("UAT-04", "Division is CDD or IDD", "Case is not transformed by the newMBD automation."),
        ("UAT-05", "Normal email without SAP2CASE marker", "Case is not transformed by the automation."),
        ("UAT-06", "AccountId is unknown but Equipment resolves", "Case uses the Asset Account as fallback and links the resolved Asset and Location."),
        ("UAT-07", "ContactId has multiple candidates", "Contact is populated only when exactly one candidate belongs to the resolved Account; otherwise it remains unassigned."),
        ("UAT-08", "Owner name is missing or ambiguous", "Existing Email-to-Case owner remains unchanged."),
        ("UAT-09", "200 eligible Cases processed together", "All Cases are transformed and the handler uses at most four queries."),
        ("UAT-10", "Material and SAP Status supplied", "Values are parsed without changing an unapproved Salesforce target field."),
    ]
    add_table(doc, ["Test", "Scenario", "Expected result"], uat_rows, [1100, 2600, 5660], font_size=8.3)

    add_heading(doc, "11. Risks and controls", 1)
    risk_rows = [
        ("Identifier population gaps", "Account, Asset, or Contact may not resolve after migration.", "Validate and load external identifiers before enabling production routing."),
        ("Owner name variation", "Names may differ between SAP text and Salesforce User.Name.", "Confirm naming governance or approve a stable owner identifier enhancement."),
        ("Email format drift", "Upstream label or ordering changes could reduce mapping accuracy.", "Treat the SAP email format as a governed interface and regression-test changes."),
        ("Unmapped Material/Status", "Users may expect values to drive Product or Case Status.", "Approve target fields and translation rules through change control."),
        ("Hidden operational failures", "Optional lookups can remain unresolved without a queue signal.", "Define a report or exception flag for unmatched identifiers before production."),
    ]
    add_table(doc, ["Risk", "Impact", "Control / mitigation"], risk_rows, [2300, 2900, 4160], font_size=8.3)

    add_heading(doc, "12. Deployment and operational readiness", 1)
    add_body(doc, "The implementation has been deployed to the newMBD scratch org. Deployment 0AfSv00000L7jp7KAB completed successfully with 13 components. Post-deployment test run 707Sv00001u3a84 passed all three tests, with 100% trigger coverage and 92% test-run coverage for the solution.")
    readiness = [
        "Complete business UAT using a real inbound email routed through the scratch-org email service or an equivalent controlled test path.",
        "Populate representative Account.SVMX_ERP_Account_ID__c, Asset.ERP_Equipment_ID__c, Contact.SAPExternalId__c, Asset Account, and Asset Location data.",
        "Confirm the approved newMBD inbound email address and Email-to-Case routing configuration.",
        "Confirm field-level access for operational support and data-load roles, including Asset.ERP_Equipment_ID__c.",
        "Approve Material and SAP Status mapping decisions or explicitly defer them for the first release.",
        "Define monitoring for unmatched identifiers and parsing exceptions.",
        "Validate that downstream entitlement and Case automation behaves correctly for SAP-created MBD Cases.",
    ]
    for item in readiness:
        add_bullet(doc, item)

    add_heading(doc, "13. Open decisions", 1)
    decisions = [
        ("OD-01", "Should Material resolve a Product, populate a text field, or remain informational?", "MBD Product Owner / Master Data"),
        ("OD-02", "How should SAP Status codes translate to newMBD Case Status or status category?", "MBD Service Operations"),
        ("OD-03", "Should unmatched identifiers create an exception flag, report entry, or support Task?", "Service Application Team"),
        ("OD-04", "Should owner matching move from display name to a stable integration identifier?", "SAP Process Owner / Salesforce Team"),
        ("OD-05", "Which permission set will govern access to the new Asset equipment identifier?", "Salesforce Security / Product Owner"),
    ]
    add_table(doc, ["ID", "Decision required", "Decision owner"], decisions, [1100, 5660, 2600], font_size=8.5)

    add_heading(doc, "14. Approval", 1)
    add_body(doc, "Approval confirms that the stated scope, business rules, mappings, acceptance criteria, and deferred decisions are suitable for controlled progression beyond the scratch org.")
    approval_rows = [
        ("MBD Business Owner", "", "", ""),
        ("newMBD Product Owner", "", "", ""),
        ("SAP Service Process Owner", "", "", ""),
        ("Salesforce Technical Owner", "", "", ""),
    ]
    add_table(doc, ["Role", "Name", "Decision", "Date"], approval_rows, [2900, 2300, 2300, 1860], font_size=8.8)

    add_heading(doc, "Appendix A. Sample outcome", 1)
    add_body(doc, "For the supplied notification 000004041703, the expected Case subject is 'ST Neph Creeping/Flashing LEDs (FSE)'. The description retains the troubleshooting narrative, including 'Summary of issues: 0', while AccountId, Owner, Equipment, Type, Division, ContactId, Material, Symptom, and Status trailer lines are removed from the description and evaluated as structured metadata.")
    add_callout(
        doc,
        "Current implementation status.",
        "Scratch-org deployment and automated verification are complete. Production readiness still depends on business UAT, identifier data quality, email routing, permissions, monitoring, and approval of the open decisions in this BRD.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    doc.core_properties.title = "newMBD SAP2CASE Inbound Email Automation - Business Requirements Document"
    doc.core_properties.subject = "Business requirements for optimized MBD SAP-to-Case processing"
    doc.core_properties.author = "SDG Service Applications"
    doc.core_properties.keywords = "newMBD, SAP2CASE, Salesforce, Email-to-Case, BRD"
    doc.core_properties.comments = "Draft for business and technical review"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
