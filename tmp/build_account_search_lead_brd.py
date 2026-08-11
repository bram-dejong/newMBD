from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_guided_account_creation_brd import (
    ROOT, BLUE, DARK_BLUE, NAVY, MUTED, PALE_BLUE, PALE_GOLD, GOLD,
    configure_styles, configure_section, add_numbering_definition, apply_num,
    add_heading, add_body, add_bullet, add_numbered, add_callout,
    add_table, set_paragraph_bottom_border, set_run_font, page_break,
)

OUTPUT = ROOT / "output" / "Business_Requirements_Document_Account_Search_and_Lead_Creation.docx"


def build_document():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_numbering_definition(doc, 71, 71, "bullet")
    add_numbering_definition(doc, 72, 72, "decimal")
    add_numbering_definition(doc, 73, 73, "decimal")
    doc.core_properties.title = "Business Requirements Document - Account Search and Lead Creation"
    doc.core_properties.subject = "Salesforce Account search with D&B-assisted Lead creation"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "Salesforce, LWC, Account search, Lead, D&B, DB_Data, fuzzy matching"

    p = doc.add_paragraph("BUSINESS REQUIREMENTS DOCUMENT", style="Title")
    p.paragraph_format.space_before = Pt(18)
    doc.add_paragraph("Account Search and D&B-Assisted Lead Creation", style="Subtitle")
    metadata = [
        ("Document status", "Draft for business review"),
        ("Version", "2.0"),
        ("Date", "10 August 2026"),
        ("Business sponsor", "TBD"),
        ("Business owner", "Salesforce Product Owner / Lead Management Owner (TBD)"),
        ("Delivery platform", "Salesforce Lightning Experience, Salesforce Mobile, and Experience Cloud-capable LWC"),
        ("Implementation target", "newMBD-scratch (revised flow deployed and validated)"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + ": ")
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    set_paragraph_bottom_border(rule, color=BLUE, size="14", space="5")
    add_callout(
        doc,
        "Decision requested:",
        "Approve Account Search as the controlled intake channel. When no Account exists, users will create a Lead enriched by a selected D&B organization rather than create an Account directly.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc, "Users need a fast, controlled way to determine whether a company already exists as a Salesforce Account. Creating a new Account during early qualification introduces duplicate risk and bypasses the Lead lifecycle. The revised solution searches accessible Accounts using fuzzy company-name logic and directs the user to the existing Account whenever a match is found.")
    add_body(doc, "When no Account match is found, the process searches DB_Data__c for the corresponding D&B organization. The user selects the correct D&B result when candidates exist, reviews a pre-populated Lead form, enters the contact details, and creates a Lead linked through Lead.D_B_Record__c. If neither an Account nor a D&B organization matches, the user may still create an unlinked Lead for qualification.")
    add_body(doc, "The solution never creates an Account. Account conversion remains governed by the existing Lead conversion process, preserving qualification, ownership, and downstream automation. The LWC is responsive for desktop and mobile, while Apex user-mode security and a dedicated permission set enforce access.")

    add_heading(doc, "2. Business Context and Problem Statement", 1)
    add_body(doc, "Company lookup is often the first step in a sales or service intake process. Exact-text searches miss variations caused by punctuation, corporate suffixes, spacing, and typing errors. A user who cannot immediately find an Account may create another Account, fragmenting commercial activity, service history, reporting, ownership, and customer hierarchy.")
    add_body(doc, "The business requires an Account-first search experience. A confirmed existing Account ends the process. A company that is not yet an Account should enter Salesforce as a Lead, using D&B reference data wherever possible to improve company identity and address quality.")

    add_heading(doc, "3. Objectives and Success Measures", 1)
    for item in [
        "Make Account search the mandatory first step for new-company intake.",
        "Prevent the guided process from creating Accounts directly.",
        "Route users to an existing Account when a fuzzy match is found.",
        "Create a Lead when no Account exists, using D&B data for company and address enrichment.",
        "Preserve a traceable relationship between the Lead and the selected DB_Data__c record.",
        "Provide a consistent, accessible browser and mobile experience.",
    ]:
        add_bullet(doc, item)
    add_table(
        doc,
        ["Measure", "Proposed target", "Evidence"],
        [
            ("Accounts created by the guided process", "0", "Transaction audit and code review"),
            ("Lead blocked when an Account match exists", "100%", "Automated tests and UAT"),
            ("D&B selection when blocking candidates exist", "100%", "Wizard behavior and UAT"),
            ("Desktop and mobile completion", "Successful", "Cross-form-factor UAT"),
            ("Search response under representative volume", "95th percentile under 3 seconds", "Performance test"),
            ("Focused automated test execution", "100% pass", "Current scratch result: 8/8; 93% feature coverage"),
        ],
        [3400, 1950, 4010],
        font_size=9.0,
    )

    page_break(doc)
    add_heading(doc, "4. Scope", 1)
    add_heading(doc, "4.1 In Scope", 2)
    for item in [
        "A responsive Account Search LWC exposed as a custom tab and available to Lightning App Pages, Home Pages, and Experience Cloud pages.",
        "Automatic and explicit search after at least three company-name characters are entered.",
        "Fuzzy comparison against accessible Account.Name and DB_Data__c.Duns_Name__c values.",
        "Blocking Lead creation and offering an Open Account action when an Account match exists.",
        "D&B result presentation with company, DUNS number, match score, and address context.",
        "Mandatory D&B selection when candidates exist and no Account exists.",
        "Pre-population of Lead company and address fields from the selected D&B organization.",
        "Creation of a Lead with Lead.D_B_Record__c when D&B was selected.",
        "Creation of an unlinked Lead when neither Account nor D&B matches.",
        "Server-side revalidation immediately before Lead insert.",
        "Permission-set access, user-mode queries/DML, responsive layout, keyboard operation, and accessible status messages.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.2 Out of Scope", 2)
    for item in [
        "Direct Account creation from the Account Search component.",
        "Automated Account creation or Lead conversion.",
        "Automatic merging or deletion of existing Accounts or Leads.",
        "Real-time D&B API integration, DUNS purchase, or creation of DB_Data__c records.",
        "Bulk Lead import, Lead deduplication, or mass Account matching.",
        "Automatic selection among multiple plausible D&B candidates.",
        "Changes to Account.D_B_Record__c or DB_Data__c.Account__c relationship behavior.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Stakeholders and Responsibilities", 1)
    add_table(
        doc,
        ["Stakeholder", "Primary responsibility", "Decision / input required"],
        [
            ("Business Sponsor", "Own business outcome and adoption", "Approve scope and rollout"),
            ("Salesforce Product Owner", "Prioritize requirements and accept release", "Approve thresholds, fields, and navigation"),
            ("Lead Management Owner", "Own qualification and conversion policy", "Approve Lead fields, status, and routing"),
            ("Customer Data Steward", "Own D&B and Account identity quality", "Validate match rules and escalation path"),
            ("Sales / Service Users", "Search companies and create Leads", "Participate in desktop/mobile UAT"),
            ("Salesforce Administrator", "Assign access and configure apps", "Assign permission set and mobile navigation"),
            ("Development Team", "Maintain LWC, Apex, metadata, and tests", "Support tuning and deployment"),
        ],
        [2200, 3500, 3660],
        font_size=9.0,
    )

    add_heading(doc, "6. Future-State Process", 1)
    for step in [
        "The user opens Account Search from a desktop browser or Salesforce Mobile navigation.",
        "The user enters at least three characters of the company name. Search starts after typing pauses or when the user selects Search Accounts.",
        "The solution searches accessible Accounts and D&B organizations and scores normalized name similarity.",
        "If an Account match exists, Lead creation is blocked and the user opens the existing Account.",
        "If no Account exists but D&B candidates are found, the user selects the correct D&B organization.",
        "The Lead form is displayed with company and available address fields pre-populated. The user enters at least the contact last name and any available contact details.",
        "If neither Account nor D&B matches, the Lead form is displayed without a D&B relationship.",
        "On save, Apex repeats Account and D&B checks, creates the Lead in user mode, and navigates to the Lead record.",
        "Any subsequent Account creation occurs through the approved Lead qualification and conversion process.",
    ]:
        add_numbered(doc, step)

    page_break(doc)
    add_heading(doc, "7. Business Rules", 1)
    rules = [
        ("BR-01", "Company Name is mandatory and must contain 3 to 255 characters."),
        ("BR-02", "Account search is mandatory before the Lead form becomes available."),
        ("BR-03", "A Salesforce Account match at or above the fuzzy threshold always blocks Lead creation."),
        ("BR-04", "When D&B candidates exist and no Account matches, the user must select the correct D&B record before creating the Lead."),
        ("BR-05", "When neither Account nor D&B matches, the user may create an unlinked Lead."),
        ("BR-06", "The process shall not insert, update, or reserve an Account."),
        ("BR-07", "A Lead created from D&B shall reference the selected parent through Lead.D_B_Record__c."),
        ("BR-08", "One D&B record may support many Leads and many Accounts; selection does not reserve the D&B row."),
        ("BR-09", "Lead Last Name is mandatory. First Name, Email, Phone, and Website are optional unless org-specific rules require them."),
        ("BR-10", "Company-name normalization ignores case, punctuation, whitespace, and common trailing corporate suffixes for fuzzy comparison."),
        ("BR-11", "The initial blocking threshold is 74% similarity and must be validated with representative business data."),
        ("BR-12", "Only the ten highest-scoring candidates in each result category are displayed; candidate retrieval is bounded."),
        ("BR-13", "Server-side Account and D&B checks are authoritative and run again immediately before Lead insert."),
        ("BR-14", "Lead Status shall use the org's active default value; if no default is configured, the first active status is used."),
    ]
    add_table(doc, ["Rule", "Business rule"], rules, [1150, 8210], font_size=9.3)

    add_heading(doc, "8. Functional Requirements", 1)
    functional = [
        ("FR-001", "Expose Account Search through a responsive Lightning custom tab and supported Lightning/Experience page targets.", "Must"),
        ("FR-002", "Accept a company name of 3 to 255 characters and reject invalid search attempts.", "Must"),
        ("FR-003", "Start search after approximately 500 ms of inactivity and provide an explicit Search Accounts action.", "Should"),
        ("FR-004", "Retrieve accessible Account and DB_Data__c candidates using bounded containment, prefix, and suffix patterns.", "Must"),
        ("FR-005", "Normalize names and calculate Levenshtein similarity before returning matches at or above 74%.", "Must"),
        ("FR-006", "Display existing Account name, match score, and an Open Account action.", "Must"),
        ("FR-007", "Suppress the D&B selection and Lead form whenever an Account match exists.", "Must"),
        ("FR-008", "Display D&B company, DUNS number, match score, and available address context when no Account exists.", "Must"),
        ("FR-009", "Require selection when one or more D&B candidates are returned.", "Must"),
        ("FR-010", "Populate Lead Company, Street, City, State/Province, Postal Code, and Country from selected D&B data.", "Must"),
        ("FR-011", "Allow users to review and edit all pre-populated Lead fields before save.", "Must"),
        ("FR-012", "Capture First Name, Last Name, Email, Phone, Website, and address fields; require Last Name.", "Must"),
        ("FR-013", "Validate that a supplied D&B identifier is an accessible DB_Data__c record.", "Must"),
        ("FR-014", "Repeat Account and D&B matching before Lead insert and block if the state changed.", "Must"),
        ("FR-015", "Insert Lead.D_B_Record__c when D&B was selected and leave it null when no D&B match exists.", "Must"),
        ("FR-016", "Permit multiple Leads to reference the same D&B organization.", "Must"),
        ("FR-017", "Determine an active Lead Status from org configuration rather than hard-coding a status label.", "Must"),
        ("FR-018", "Navigate to the newly created Lead after successful insert.", "Should"),
        ("FR-019", "Announce progress and errors in an accessible live region using plain-language messages.", "Must"),
        ("FR-020", "Never perform Account DML from the Account Search controller.", "Must"),
    ]
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        functional,
        [1050, 7110, 1200],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.8,
    )

    page_break(doc)
    add_heading(doc, "9. Non-Functional Requirements", 1)
    nonfunctional = [
        ("NFR-001", "Responsive design", "The flow shall reflow to one column on small screens without horizontal scrolling."),
        ("NFR-002", "Accessibility", "Use labelled Salesforce base components, keyboard controls, visible validation, and live status/error announcements."),
        ("NFR-003", "Security", "Queries and Lead insert shall honor sharing, CRUD, field-level security, and Apex access."),
        ("NFR-004", "Least privilege", "The permission set shall grant Account read, DB_Data__c read, Lead create/read/edit, the Lead D&B lookup, Apex class, and tab."),
        ("NFR-005", "Performance", "Search shall be debounced and bounded to 200 candidates per object and ten displayed matches per category."),
        ("NFR-006", "Reliability", "Server checks shall be authoritative; browser state alone shall never authorize Lead creation."),
        ("NFR-007", "Maintainability", "Threshold, normalization, limits, mapped fields, and status selection shall be documented and tested."),
        ("NFR-008", "Compatibility", "Support Lightning Experience and supported Salesforce Mobile form factors."),
        ("NFR-009", "Observability", "Standard Lead audit fields and deployment/test evidence shall support operational review."),
    ]
    add_table(doc, ["ID", "Category", "Requirement"], nonfunctional, [1050, 1900, 6410], font_size=9.0)

    add_heading(doc, "10. Data Requirements and Model", 1)
    add_callout(doc, "Relationship:", "DB_Data__c (one) -> Lead (many), implemented by Lead.D_B_Record__c. The D&B row is reference data and is not updated or reserved by this process.")
    add_heading(doc, "10.1 D&B-to-Lead Mapping", 2)
    add_table(
        doc,
        ["D&B source", "Lead target", "Rule"],
        [
            ("DB_Data__c.Duns_Name__c", "Lead.Company", "Populate on selection; user may review/edit"),
            ("DB_Data__c.Geocoded_Address__c", "Lead.Street", "Populate when present"),
            ("DB_Data__c.Geocoded_City__c", "Lead.City", "Populate when present"),
            ("DB_Data__c.Geocoded_State__c", "Lead.State", "Populate when present; validate org controls"),
            ("DB_Data__c.Geocoded_Zip__c", "Lead.PostalCode", "Populate when present"),
            ("DB_Data__c.Geocoded_Country__c", "Lead.Country", "Populate when present; validate org controls"),
            ("DB_Data__c.Id", "Lead.D_B_Record__c", "Persist selected D&B reference"),
        ],
        [2950, 2500, 3910],
        font_size=9.1,
    )
    add_heading(doc, "10.2 Lead Contact Data", 2)
    add_body(doc, "The user supplies Lead contact details. LastName is required by this process and Salesforce. FirstName, Email, Phone, and Website are optional unless target-org validation or routing rules impose additional requirements. Lead Status is derived from the target org's active picklist configuration.")

    add_heading(doc, "11. Security and Access", 1)
    for item in [
        "Users must be assigned the Account Search and Lead Creation permission set (API name Guided_Account_Creation).",
        "The permission set provides Account read access only; it does not grant Account create or edit access.",
        "The permission set grants Lead create/read/edit, Lead.D_B_Record__c edit, DB_Data__c read, controller access, and tab visibility.",
        "AccountCreationController runs with sharing. Candidate queries use USER_MODE and Lead insert uses user-mode DML.",
        "Users only match Accounts and D&B records they can access. Business owners must decide whether private Account visibility requires a data-steward or centralized-search model.",
    ]:
        add_bullet(doc, item)

    page_break(doc)
    add_heading(doc, "12. Acceptance Criteria", 1)
    acceptance = [
        ("AC-01", "Existing Account", "A fuzzy Account match blocks Lead creation and provides Open Account."),
        ("AC-02", "No Account, D&B match", "D&B candidates are shown and Lead creation remains unavailable until one is selected."),
        ("AC-03", "D&B pre-population", "Selected D&B company and available address values populate the Lead form."),
        ("AC-04", "No Account or D&B", "The user can create an unlinked Lead with Company and Last Name."),
        ("AC-05", "Server recheck", "An Account created after initial search but before save causes Lead creation to be blocked."),
        ("AC-06", "D&B validation", "An invalid or inaccessible D&B identifier is rejected."),
        ("AC-07", "One-to-many", "Multiple Leads can reference the same DB_Data__c record."),
        ("AC-08", "No Account DML", "No Account is inserted or updated in any wizard outcome."),
        ("AC-09", "Status", "The new Lead receives an active status from org configuration."),
        ("AC-10", "Mobile", "Search, Account open, D&B selection, Lead entry, and save work without horizontal scrolling."),
        ("AC-11", "Browser", "The component is accessible from the Account Search tab in a supported browser."),
        ("AC-12", "Accessibility", "Controls are labelled and keyboard-operable; status and errors are announced."),
        ("AC-13", "Security", "Unauthorized users cannot access the tab/controller; authorized users remain subject to record and field permissions."),
        ("AC-14", "Automated validation", "Focused Apex tests pass with at least 75% coverage; current result is 8/8 and 93%."),
    ]
    add_table(doc, ["ID", "Scenario", "Expected result"], acceptance, [1050, 1900, 6410], font_size=8.9)

    add_heading(doc, "13. Assumptions and Dependencies", 1)
    add_heading(doc, "13.1 Assumptions", 2)
    for item in [
        "DB_Data__c contains usable Duns_Name__c values and sufficient address context.",
        "Users can identify the correct D&B organization from company, DUNS number, score, and address.",
        "The 74% threshold will be tested using representative regional, multilingual, abbreviated, and misspelled names.",
        "Target-org Lead validation rules, flows, assignment rules, duplicate rules, and address controls accept the intended values or will be addressed during UAT.",
        "Lead conversion remains the approved path to create a new Account after qualification.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "13.2 Dependencies", 2)
    for item in [
        "DB_Data__c and required D&B fields are deployed and populated.",
        "Lead.D_B_Record__c is deployed and available to authorized users.",
        "The permission set is assigned and Account Search is added to desktop/mobile navigation.",
        "The active Lead Status configuration is valid.",
        "Business owners supply UAT examples and define the no-match escalation process.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "14. Risks, Controls, and Open Decisions", 1)
    risks = [
        ("Threshold too low", "Legitimate new Leads are blocked by Account false positives", "Tune with UAT corpus; steward escalation", "Open"),
        ("Threshold too high", "Existing Accounts are missed and duplicate Leads are created", "Monitor outcomes and retest threshold", "Open"),
        ("Private Account sharing", "User may not see an Account that already exists", "Centralized visibility or steward process", "Decision"),
        ("Large D&B dataset", "Search may slow or omit candidates", "Performance test; consider search-key/SOSL optimization", "Open"),
        ("Address controls", "D&B state/country value may fail Lead validation", "Canonicalize values in UAT", "Open"),
        ("Lead automation", "Assignment rules, flows, or duplicate rules may reject insert", "Target-org regression and routing tests", "Open"),
        ("No D&B match", "Lead may have weaker company identity", "Permit qualification; steward enrichment process", "Decision"),
        ("Hard-coded threshold", "Tuning requires deployment", "Evaluate Custom Metadata configuration", "Future"),
    ]
    add_table(doc, ["Risk / decision", "Impact", "Control / action", "Status"], risks, [2250, 2600, 3410, 1100], font_size=8.5)

    page_break(doc)
    add_heading(doc, "15. Deployment and Rollout", 1)
    for step in [
        "Confirm target-org Lead required fields, record types, statuses, assignment rules, validation rules, duplicate rules, and automation.",
        "Deploy Lead.D_B_Record__c, Apex, tests, LWC, custom tab, and permission-set updates through the approved release process.",
        "Assign the permission set to pilot users and add Account Search to desktop and mobile app navigation.",
        "Run focused and regression Apex tests and confirm required deployment coverage.",
        "Execute UAT for existing Account matches, D&B matches, no-match Leads, invalid D&B identifiers, mobile layout, and Lead routing.",
        "Train users to open existing Accounts and to use Lead qualification/conversion for new companies.",
        "Monitor search latency, Account false positives/negatives, D&B selection quality, Lead insert failures, and downstream conversion results.",
    ]:
        p = doc.add_paragraph(style="Normal")
        apply_num(p, 73)
        r = p.add_run(step)
        set_run_font(r)

    add_heading(doc, "16. Test Evidence", 1)
    add_body(doc, "The revised solution was deployed to newMBD-scratch on 10 August 2026. The AccountCreationControllerTest run completed with 8 of 8 executions passing and 93% feature test-run coverage. Covered scenarios include Account blocking, fuzzy D&B selection, D&B-linked Lead creation, no-match Lead creation, multiple Leads per D&B record, and deterministic fuzzy normalization.")
    add_callout(doc, "Release gate:", "Production approval also requires target-org regression testing, Account visibility decisions, Lead routing validation, representative-name UAT, and desktop/mobile acceptance.")

    add_heading(doc, "17. Approvals", 1)
    add_body(doc, "Approval confirms acceptance of the requirements, business rules, assumptions, risks, and rollout conditions in this document.")
    add_table(
        doc,
        ["Role", "Name", "Decision", "Date"],
        [
            ("Business Sponsor", "", "Approve / Reject", ""),
            ("Salesforce Product Owner", "", "Approve / Reject", ""),
            ("Lead Management Owner", "", "Approve / Reject", ""),
            ("Customer Data Steward", "", "Approve / Reject", ""),
            ("Technical Lead", "", "Approve / Reject", ""),
        ],
        [2400, 2600, 2400, 1960],
        font_size=9.3,
    )

    add_heading(doc, "Appendix A. Implementation Traceability", 1)
    add_table(
        doc,
        ["Component", "Implementation", "Requirement coverage"],
        [
            ("Account Search LWC", "force-app/main/default/lwc/accountCreationWizard", "FR-001 to FR-012, FR-018, FR-019"),
            ("AccountCreationController", "Search and createLead Apex actions", "FR-004 to FR-020"),
            ("Lead D&B lookup", "Lead.D_B_Record__c", "BR-07, FR-015, FR-016"),
            ("Permission set", "Guided_Account_Creation", "NFR-003, NFR-004, AC-13"),
            ("Custom tab", "Guided_Account_Creation labelled Account Search", "FR-001, AC-11"),
            ("Focused tests", "AccountCreationControllerTest", "AC-01 to AC-09, AC-14"),
        ],
        [2600, 3800, 2960],
        font_size=8.8,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
