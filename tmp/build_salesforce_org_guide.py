from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "Salesforce_Org_Resources_and_Skills_Guide.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
TEXT = "222222"
WHITE = "FFFFFF"
GOLD = "B7791F"
GREEN = "2F6B4F"


def set_font(run, name="Calibri", size=11, color=TEXT, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    repeat_header(row)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_repeatable_numbering_paragraph(p, level=0):
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_repeatable_numbering_paragraph(p, level)
    set_font(p.add_run(text))
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_repeatable_numbering_paragraph(p)
    set_font(p.add_run(text))
    return p


def body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size={1:16, 2:13, 3:12}[level], color={1:BLUE, 2:BLUE, 3:DARK_BLUE}[level], bold=True)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.04)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(label + "  "), color=accent, bold=True)
    set_font(p.add_run(text), color=TEXT)


def add_role_table(doc):
    data = [
        ("Executive sponsor", "Funding, priorities and cross-functional decisions", "Part-time"),
        ("Business product owner", "Scope, backlog, value and acceptance", "50-100%"),
        ("Program / project manager", "Plan, dependencies, risk, budget and governance", "Full-time"),
        ("Salesforce solution architect", "End-to-end data, security, integration and platform design", "Full-time initially"),
        ("Business analyst / process designer", "Requirements, process maps and user stories", "1-3"),
        ("Salesforce administrator", "Configuration, permissions, Flow, reports and supportability", "1-2"),
        ("Salesforce developer", "Apex, Lightning and complex extensions", "1-3"),
        ("Integration architect / developer", "ERP, identity, telephony, tax, payments and middleware", "1-3"),
        ("Data architect / migration specialist", "Data model, cleansing, migration and reconciliation", "1-2"),
        ("QA / test lead", "Regression, integration, performance and UAT", "1-3"),
        ("DevOps / release engineer", "Git, CI/CD, environments and deployments", "0.5-1"),
        ("Change and adoption lead", "Training, communications and adoption measurement", "1"),
        ("Security / privacy specialist", "Identity, access, audit, retention and compliance", "Part-time"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2400, 5520, 1440])
    headers = ["Resource", "Primary responsibility", "Typical need"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        set_font(p.add_run(h), size=9.5, color=WHITE, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 2 else WD_ALIGN_PARAGRAPH.LEFT
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(data):
        cells = table.add_row().cells
        if row_idx % 2:
            for c in cells:
                shade_cell(c, LIGHT_GRAY)
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(text), size=9.2, bold=(i == 0))
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, [2400, 5520, 1440])


def add_footer_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    set_font(p.add_run("Salesforce Org Resources & Skills Guide   |   "), size=8.5, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MID_GRAY)
    r_pr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "17")
    r_pr.append(sz)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Compact reference guide preset.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for level, size, color, before, after in [
    (1, 16, BLUE, 18, 10),
    (2, 13, BLUE, 14, 7),
    (3, 12, DARK_BLUE, 10, 5),
]:
    s = styles[f"Heading {level}"]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Bullet 2", "List Number"):
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.paragraph_format.left_indent = Inches(0.375)
    s.paragraph_format.first_line_indent = Inches(-0.188)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.25

# Running header/footer, deliberately quiet.
hp = section.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("IMPLEMENTATION REFERENCE GUIDE"), size=8.5, color=MID_GRAY, bold=True)
add_footer_page_number(section)

# First-page opening block (customer-pack inspired, without decorative rules).
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("SALESFORCE IMPLEMENTATION GUIDE"), size=10, color=GOLD, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(7)
set_font(p.add_run("Resources & Skills for a New Salesforce Org"), size=27, color=NAVY, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
set_font(p.add_run("Sales Cloud, Service Cloud and Revenue Cloud / Revenue Management"), size=14, color=DARK_BLUE, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
p.paragraph_format.line_spacing = 1.2
set_font(p.add_run("A practical planning guide for assembling the implementation team, platform resources, environments, skills and governance needed to deliver an integrated customer and product-to-cash platform."), size=11.5, color=TEXT)

add_callout(
    doc,
    "Important product decision",
    "Salesforce documentation increasingly refers to the newer Revenue Cloud platform as Revenue Management. Confirm whether the program will use the newer platform or legacy Salesforce CPQ and Billing before estimating, staffing or designing migration paths.",
    fill="FFF8E8",
    accent=GOLD,
)

heading(doc, "1. Executive summary", 1)
body(doc, "This program should be treated as a cross-functional transformation, not simply a CRM configuration. Sales, service and product-to-cash processes share customer, product, price, contract, order and asset data. Decisions made in one workstream can therefore create downstream constraints in the others.")
body(doc, "For a moderately complex implementation, plan for a core delivery team of approximately 10-18 people plus business subject-matter experts. A small, low-integration deployment may operate with 6-8; a global or highly regulated product-to-cash program may require materially more.")

heading(doc, "2. Core implementation team", 1)
body(doc, "The following roles form the minimum multidisciplinary capability. Several roles can be combined in a small program, but accountability should remain explicit.")
add_role_table(doc)

heading(doc, "3. Business workstreams and subject-matter experts", 1)

heading(doc, "Sales Cloud", 2)
for item in [
    "Lead generation, qualification and conversion",
    "Account, contact and customer hierarchy management",
    "Opportunity stages, sales methodology and pipeline governance",
    "Forecasting, quotas, territories and sales performance measures",
    "Partner or channel sales, where applicable",
    "Approval and discount authority",
    "Sales reporting, dashboards and adoption metrics",
]:
    bullet(doc, item)

heading(doc, "Service Cloud", 2)
for item in [
    "Case intake, classification, routing, queues and escalation",
    "Service levels, entitlements and milestones",
    "Knowledge creation, approval, publication and retirement",
    "Omni-Channel, email, chat, messaging and voice",
    "Contact-center operations and quality management",
    "Complaints, returns, warranties and field service, if applicable",
    "Customer portals, self-service and service analytics",
]:
    bullet(doc, item)

heading(doc, "Revenue Cloud / Revenue Management", 2)
body(doc, "This is normally the most specialized workstream. It needs SMEs who understand the complete product-to-cash lifecycle, including finance and operational controls.")
for item in [
    "Product ownership, catalog structure, bundles, attributes and selling models",
    "Pricing methods, price books, discount policies and approval thresholds",
    "Quote configuration and generation",
    "Contracting, legal clauses and commercial governance",
    "Order capture, orchestration and fulfillment handoffs",
    "Subscriptions, amendments, renewals, cancellations and asset lifecycle",
    "Usage or consumption rating, if applicable",
    "Billing schedules, invoicing, credits, taxation, payments and collections",
    "ERP, finance, tax and revenue-recognition integration boundaries",
]:
    bullet(doc, item)

add_callout(doc, "Staffing principle", "A Salesforce configurator without product-to-cash domain expertise is not sufficient for Revenue Cloud. Pair platform specialists with RevOps, product, finance, tax, legal and order-management owners.")

heading(doc, "4. Required technical skills", 1)
for item in [
    "Salesforce data modeling and the shared Sales, Service and Revenue object model",
    "Sales Cloud configuration, forecasting, territories and pipeline design",
    "Service Cloud, Knowledge, Omni-Channel, entitlements and contact-center patterns",
    "Revenue Cloud catalog, configuration, pricing, quoting, contracts, orders, assets and billing",
    "Flow, approval automation and error-handling patterns",
    "Apex and Lightning Web Components for justified extensions",
    "REST, SOAP, Bulk API, Platform Events and event-driven integration",
    "Identity, SSO, MFA, connected apps, named credentials and integration-user design",
    "Role hierarchy, sharing, permission sets, field-level security and auditability",
    "Data migration, cleansing, deduplication, reconciliation and archival",
    "Reports, dashboards, analytics and KPI design",
    "Salesforce CLI, Git, CI/CD, test automation and release governance",
    "Platform limits, performance, large data volumes and resilient solution design",
    "Training, adoption, operating model and post-go-live support",
]:
    bullet(doc, item)

heading(doc, "Relevant credentials", 2)
body(doc, "Certifications are useful evidence of foundational knowledge, but they do not replace implementation experience. Relevant credentials include:")
for item in [
    "Salesforce Administrator and Platform App Builder",
    "Sales Cloud Consultant and Service Cloud Consultant",
    "Revenue Cloud Consultant",
    "Platform Developer I",
    "Data Architect, Sharing and Visibility Architect, and Integration Architect",
    "Application Architect or System Architect for senior design leadership",
]:
    bullet(doc, item)

heading(doc, "5. Platform, environment and delivery resources", 1)
body(doc, "Establish the delivery lifecycle before feature build begins. Production should never be the primary development, integration-test or training environment.")
for item in [
    "Production org",
    "Individual Developer sandboxes and/or scratch orgs",
    "Shared integration environment",
    "QA and system-test environment",
    "User acceptance testing environment",
    "Training environment",
    "Partial Copy or Full Copy sandbox for realistic end-to-end testing",
    "Git repository as the metadata source of truth",
    "Salesforce CLI and automated CI/CD pipeline",
    "Automated test framework and regression suite",
    "Test-data generation, masking and refresh procedures",
    "Architecture, requirements, decision and operating-model repository",
    "Release calendar, change control and incident-management process",
]:
    bullet(doc, item)

add_callout(doc, "Environment sizing", "Sandbox types and quantities depend on the Salesforce edition and purchased entitlements. Include refresh cadence, data masking, release-preview strategy and integration endpoint management in the environment plan.")

heading(doc, "6. Licensing and connected-system checklist", 1)

heading(doc, "Licensing decisions", 2)
for item in [
    "Salesforce edition and geographic / Hyperforce requirements",
    "Sales Cloud, Service Cloud and Revenue Cloud user counts and personas",
    "Revenue Cloud Advanced and/or Billing functional entitlements",
    "Digital Engagement, Voice, telephony or contact-center capabilities",
    "Experience Cloud external-user licensing",
    "Knowledge, Field Service and analytics requirements",
    "Shield encryption, event monitoring, audit and compliance requirements",
    "Sandbox types, storage, API capacity and integration-user licences",
]:
    bullet(doc, item)

heading(doc, "Likely integrations", 2)
for item in [
    "ERP, general ledger and accounts receivable",
    "Master data management and product lifecycle management",
    "Identity provider and user provisioning",
    "Marketing automation and customer data platforms",
    "Contact-center and telephony platforms",
    "E-commerce and partner channels",
    "Tax engines, payment gateways and electronic signature",
    "Contract lifecycle management",
    "Data warehouse, lakehouse and enterprise reporting",
]:
    bullet(doc, item)

heading(doc, "7. Essential design deliverables", 1)
for item in [
    "Business capability map and agreed scope boundaries",
    "Current-state and target-state process maps",
    "Phased roadmap, dependencies and measurable outcomes",
    "System context, integration and data-flow diagrams",
    "Canonical customer, product, pricing, contract, order and asset models",
    "Product catalog ownership and governance model",
    "Pricing, discount and approval decision tables",
    "Quote-to-order-to-invoice process and exception design",
    "Security personas, access matrix and data-classification model",
    "Data retention, archival, migration and reconciliation plans",
    "Environment, source-control and deployment strategy",
    "Test strategy, requirements traceability and acceptance criteria",
    "Reporting and KPI catalogue",
    "Support model, release governance and adoption plan",
]:
    bullet(doc, item)

heading(doc, "8. Recommended implementation sequence", 1)
for item in [
    "Establish program governance, architecture principles, identity, security, DevOps and the shared customer/product data model.",
    "Implement foundational Sales Cloud processes and reporting.",
    "Implement Service Cloud and customer-service integrations.",
    "Build the Revenue Cloud catalog, pricing, quoting and commercial approvals.",
    "Add contracts, orders, assets, amendments, cancellations and renewals.",
    "Add billing, tax, usage and ERP/finance integrations where in scope.",
    "Complete end-to-end product-to-cash testing, data migration, training and phased rollout.",
]:
    numbered(doc, item)

add_callout(doc, "Architecture recommendation", "Design Revenue Cloud from the beginning even if it goes live later. Product, price book, account, contract, order and asset decisions made during the Sales Cloud phase can otherwise create expensive rework.", fill="EDF7F2", accent=GREEN)

heading(doc, "9. Readiness gate before build", 1)
body(doc, "Configuration should begin only after the program can answer the following questions with named owners and documented decisions:")
for item in [
    "Which Revenue Cloud product generation and licence set is being implemented?",
    "What is the authoritative source for customer, product, price, contract, order and billing data?",
    "Which processes will remain in ERP, finance, tax, CLM, telephony or other platforms?",
    "Which user and integration personas require access to which data and operations?",
    "What are the migration volumes, quality issues, retention rules and reconciliation tolerances?",
    "Which end-to-end scenarios constitute go-live acceptance?",
    "Who owns the product catalog, pricing rules, service taxonomy and release process after go-live?",
]:
    bullet(doc, item)

heading(doc, "10. Salesforce reference resources", 1)
sources = [
    ("Prepare for a Revenue Management implementation", "https://help.salesforce.com/s/articleView?id=ind.setup_revenue_cloud.htm&language=en_US&type=5"),
    ("Revenue Cloud / Revenue Management product documentation", "https://help.salesforce.com/s/articleView?id=sf.revenue_lifecycle_management.htm&language=en_US&type=5"),
    ("Salesforce Certified Revenue Cloud Consultant", "https://trailhead.salesforce.com/en/credentials/revenuecloudconsultant"),
    ("Salesforce sandbox types and management practices", "https://help.salesforce.com/s/articleView?id=000212583&language=en_US&type=1"),
    ("When to use a Salesforce sandbox", "https://help.salesforce.com/s/articleView?id=deploy_sandboxes_intro.htm&language=en_US"),
    ("Salesforce Well-Architected framework", "https://architect.salesforce.com/docs/architect/well-architected/guide/overview.html"),
    ("Salesforce Well-Architected security guidance", "https://architect.salesforce.com/docs/architect/well-architected/guide/secure.html"),
    ("Salesforce Well-Architected resilience guidance", "https://architect.salesforce.com/docs/architect/well-architected/guide/resilient.html"),
]
for title, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    set_repeatable_numbering_paragraph(p)
    add_hyperlink(p, title, url)

body(doc, "Licensing, naming and product capabilities change over time. Validate the final SKU, edition, regional, limit and coexistence assumptions with Salesforce and the implementation partner before contract signature and solution baseline approval.")

doc.core_properties.title = "Resources & Skills for a New Salesforce Org"
doc.core_properties.subject = "Sales Cloud, Service Cloud and Revenue Cloud implementation planning"
doc.core_properties.author = ""
doc.core_properties.keywords = "Salesforce, Sales Cloud, Service Cloud, Revenue Cloud, Revenue Management, implementation"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
