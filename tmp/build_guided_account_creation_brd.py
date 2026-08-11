from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_ROOT = Path(r"C:\Users\JONGB\.codex\plugins\cache\openai-primary-runtime\documents\26.805.11740\skills\documents")
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry

ROOT = Path(r"C:\Users\JONGB\OneDrive - Thermo Fisher Scientific\Documents\Playground\newMBD")
OUTPUT = ROOT / "output" / "Business_Requirements_Document_Guided_Account_Creation.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "202124"
MUTED = "5F6368"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF6D6"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
GRID = "B8C2CC"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_paragraph_bottom_border(paragraph, color=BLUE, size="12", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc, num_id, abstract_id, kind="bullet"):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, 71)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, 72)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_body(doc, text, bold_lead=None, italic=False):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.10)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_table(doc, headers, rows, widths, alignments=None, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for i, label in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = (alignments or [WD_ALIGN_PARAGRAPH.LEFT] * len(headers))[i]
        r = p.add_run(label)
        set_run_font(r, size=9.3, color=NAVY, bold=True)
    set_repeat_table_header(header)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = (alignments or [WD_ALIGN_PARAGRAPH.LEFT] * len(headers))[i]
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    apply_table_geometry(table, widths, indent_dxa=120)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True
    for level, size, color, before, after in [
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    title = styles["Title"]
    title.font.name = "Calibri Light"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri Light")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri Light")
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_after = Pt(16)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("BUSINESS REQUIREMENTS DOCUMENT  |  ACCOUNT SEARCH AND LEAD CREATION")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    set_paragraph_bottom_border(p, color="D7DBE2", size="6", space="3")
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("Internal working document")
    set_run_font(r, size=9, color=MUTED)

    even_header = section.even_page_header
    ep = even_header.paragraphs[0]
    ep.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ep.paragraph_format.space_after = Pt(0)
    r = ep.add_run("BUSINESS REQUIREMENTS DOCUMENT  |  ACCOUNT SEARCH AND LEAD CREATION")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    set_paragraph_bottom_border(ep, color="D7DBE2", size="6", space="3")
    even_footer = section.even_page_footer
    efp = even_footer.paragraphs[0]
    efp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    efp.paragraph_format.space_before = Pt(0)
    r = efp.add_run("Internal working document")
    set_run_font(r, size=9, color=MUTED)


def page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_numbering_definition(doc, 71, 71, "bullet")
    add_numbering_definition(doc, 72, 72, "decimal")
    doc.core_properties.title = "Business Requirements Document - Guided Account Creation"
    doc.core_properties.subject = "Salesforce Account creation with D&B fuzzy matching and duplicate prevention"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "Salesforce, LWC, Account, D&B, DB_Data, fuzzy matching, duplicate prevention, BRD"

    # Memo masthead.
    p = doc.add_paragraph("BUSINESS REQUIREMENTS DOCUMENT", style="Title")
    p.paragraph_format.space_before = Pt(18)
    p = doc.add_paragraph("Guided Account Creation with D&B Matching", style="Subtitle")
    metadata = [
        ("Document status", "Draft for business review"),
        ("Version", "1.0"),
        ("Date", "10 August 2026"),
        ("Business sponsor", "TBD"),
        ("Business owner", "Salesforce Product Owner / Customer Data Steward (TBD)"),
        ("Delivery platform", "Salesforce Lightning Experience, Salesforce Mobile, and Experience Cloud-capable LWC"),
        ("Implementation target", "newMBD-scratch (feature deployed and validated)"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + ": ")
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    set_paragraph_bottom_border(rule, color=BLUE, size="14", space="5")
    add_callout(
        doc,
        "Decision requested:",
        "Approve the controlled Account-creation process and authorize production-readiness activities, including legacy duplicate assessment, normalized-name backfill, permission assignment, and desktop/mobile navigation rollout.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc, "Salesforce users need a controlled way to create Accounts without introducing duplicate customer records or bypassing the existing D&B reference dataset. The proposed solution is a responsive Lightning Web Component (LWC) that searches both DB_Data__c and Account as the user enters a company name, applies fuzzy name matching, and determines whether the user must select an existing D&B record, open an existing Account, or may continue with a genuinely new Account.")
    add_body(doc, "The solution treats DB_Data__c as a reference parent. One D&B record may relate to many distinct Accounts through Account.D_B_Record__c. Selecting a D&B record therefore does not reserve that record or prevent other legitimate child Accounts. Duplicate protection applies to Account identity, not to the number of Accounts related to one D&B parent.")
    add_body(doc, "The implementation adds defense in depth: guided fuzzy checks in the LWC, a server-side recheck immediately before insert, and a normalized unique Account-name key maintained by a trigger to protect against exact duplicates and concurrent inserts. A dedicated permission set limits access to the required Apex class, objects, fields, and navigation tab.")

    add_heading(doc, "2. Business Context and Problem Statement", 1)
    add_body(doc, "Account creation is a high-impact data-governance event. Duplicate Accounts fragment commercial activity, ownership, service history, reporting, and customer hierarchy. Name variation, punctuation, corporate suffixes, and typing errors make exact-text lookups insufficient. At the same time, D&B represents an organization-level reference and may legitimately support multiple Salesforce Accounts, such as sites, divisions, or operating entities.")
    add_body(doc, "The business requires a single entry experience that is usable in a desktop browser and on mobile, makes the D&B search unavoidable, prevents creation when an Account already exists, and still permits a new Account when no reasonable match is found.")

    add_heading(doc, "3. Objectives and Success Measures", 1)
    objectives = [
        "Prevent duplicate Account creation through the guided channel.",
        "Make D&B matching a mandatory decision point before Account creation.",
        "Reduce manual re-keying by pre-populating available company and billing-address data.",
        "Preserve the one-to-many relationship from DB_Data__c to Account.",
        "Provide a consistent, accessible experience in Lightning desktop, Salesforce Mobile, and supported browser contexts.",
        "Apply Salesforce sharing, object, field, and Apex-access controls throughout the transaction.",
    ]
    for item in objectives:
        add_bullet(doc, item)
    add_table(
        doc,
        ["Measure", "Proposed target", "Evidence"],
        [
            ("Duplicate Accounts created through the wizard", "0", "Duplicate-control test results and production monitoring"),
            ("D&B selection when a blocking D&B match is returned", "100%", "Wizard behavior and UAT scenarios"),
            ("Desktop and mobile completion", "Successful", "Cross-form-factor UAT"),
            ("Search response under representative volume", "95th percentile under 3 seconds", "Performance test in a production-like dataset"),
            ("Focused automated test execution", "100% pass", "Apex test run; current scratch result is 6/6"),
        ],
        [3600, 1900, 3860],
        font_size=9.1,
    )

    page_break(doc)
    add_heading(doc, "4. Scope", 1)
    add_heading(doc, "4.1 In Scope", 2)
    for item in [
        "A guided Account-creation LWC exposed as a Salesforce custom tab, App Page component, Home Page component, and Experience Cloud page component.",
        "Automatic and on-demand name search after at least three characters are entered.",
        "Fuzzy matching against DB_Data__c.Duns_Name__c and existing Account.Name values.",
        "Blocking creation when a matching Salesforce Account is found.",
        "Mandatory D&B selection when one or more blocking D&B matches are found.",
        "Pre-population of Account name and available billing-address fields from the selected D&B record.",
        "Creation of Account.D_B_Record__c as the child-side link to the selected D&B parent.",
        "Normalized exact-name uniqueness across new and updated Accounts.",
        "Server-side duplicate revalidation immediately before Account insert.",
        "Permission-set-based access and responsive, keyboard-usable interface behavior.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.2 Out of Scope", 2)
    for item in [
        "Automatic merging or deletion of existing duplicate Accounts.",
        "Creation or enrichment of DB_Data__c records from external D&B services.",
        "Real-time D&B API integration or DUNS-number purchase workflows.",
        "A replacement for Enterprise Territory Management, customer hierarchy governance, or Account ownership assignment.",
        "Bulk Account conversion, Lead conversion, or mass-import user interfaces.",
        "Automated selection among multiple plausible D&B matches.",
        "Changes to the existing DB_Data__c.Account__c lookup; it is not used as the authoritative relationship because it cannot represent the required one-to-many model.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Stakeholders and Responsibilities", 1)
    add_table(
        doc,
        ["Stakeholder", "Primary responsibility", "Decision / input required"],
        [
            ("Business sponsor", "Own business outcome and adoption", "Approve scope, funding, and rollout"),
            ("Salesforce Product Owner", "Prioritize requirements and accept release", "Approve threshold, fields, and navigation placement"),
            ("Customer Data Steward", "Own duplicate and D&B governance", "Validate match rules and remediate legacy collisions"),
            ("Sales / Service users", "Create and consume Account records", "Participate in desktop and mobile UAT"),
            ("Salesforce Administrator", "Assign access and configure apps", "Assign permission set and add navigation item"),
            ("Salesforce Development Team", "Maintain LWC, Apex, trigger, tests", "Support tuning and production deployment"),
            ("Information Security", "Review least-privilege access", "Confirm access model and audit expectations"),
        ],
        [2200, 3560, 3600],
        font_size=9.0,
    )

    add_heading(doc, "6. Future-State Business Process", 1)
    add_numbered(doc, "The user opens Guided Account Creation from a Lightning browser tab or from an app navigation menu in Salesforce Mobile.")
    add_numbered(doc, "The user enters at least three characters of the Account name. The solution waits briefly for typing to pause or the user selects Search D&B.")
    add_numbered(doc, "The solution searches accessible D&B reference data and accessible Salesforce Accounts, normalizes company names, calculates fuzzy similarity, and displays the strongest matches.")
    add_numbered(doc, "If a matching Account exists, creation is blocked and the user is offered a direct Open Account action.")
    add_numbered(doc, "If D&B matches exist and no matching Account exists, the user selects the correct D&B parent. The form is pre-populated with available name and billing-address data.")
    add_numbered(doc, "If no D&B or Account match exists, the user may complete a new Account without a D&B parent.")
    add_numbered(doc, "On save, the server repeats duplicate checks, applies user-mode security, creates the Account, maintains the normalized unique key, and navigates the user to the new record.")

    page_break(doc)
    add_heading(doc, "7. Business Rules", 1)
    rules = [
        ("BR-01", "Account Name is mandatory and must contain 3 to 255 characters."),
        ("BR-02", "The user may not bypass matching. Creation is enabled only after the current name has been evaluated or a D&B result has been selected."),
        ("BR-03", "A matching Salesforce Account always blocks creation, regardless of whether a D&B record has been selected."),
        ("BR-04", "When one or more D&B matches are returned and no Account match exists, the user must select the correct D&B record before creation."),
        ("BR-05", "When neither a D&B match nor an Account match exists, the Account may be created without a D&B relationship."),
        ("BR-06", "DB_Data__c is the parent and Account is the child. A D&B record may have many distinct child Accounts; an Account may reference no more than one D&B parent through Account.D_B_Record__c."),
        ("BR-07", "Selecting a D&B record must not update or reserve DB_Data__c.Account__c and must not prevent other distinct Accounts from linking to the same D&B parent."),
        ("BR-08", "Company-name normalization ignores case, punctuation, whitespace, and common trailing corporate suffixes for fuzzy comparison."),
        ("BR-09", "The initial blocking fuzzy threshold is 74% similarity. Threshold changes require controlled testing because false positives prevent creation and false negatives allow duplicates."),
        ("BR-10", "Only the ten highest-scoring matches are presented to the user; candidate retrieval is bounded to protect interactive performance."),
        ("BR-11", "The server must repeat matching immediately before insert so a stale or manipulated browser state cannot bypass controls."),
        ("BR-12", "A normalized unique Account-name key must be maintained on insert and update to prevent exact normalized duplicates and concurrent inserts."),
        ("BR-13", "Existing Accounts must be assessed and backfilled before production enforcement is considered complete."),
    ]
    add_table(doc, ["Rule", "Business rule"], rules, [1150, 8210], font_size=9.4)

    add_heading(doc, "8. Functional Requirements", 1)
    functional = [
        ("FR-001", "The solution shall expose a responsive Guided Account Creation LWC through a custom navigation tab.", "Must"),
        ("FR-002", "The component shall be available for Lightning App Pages, Home Pages, and Experience Cloud pages.", "Must"),
        ("FR-003", "The solution shall accept an Account name of 3 to 255 characters and prevent invalid search or save attempts.", "Must"),
        ("FR-004", "The solution shall automatically initiate search after approximately 500 ms of typing inactivity and shall provide an explicit search action.", "Should"),
        ("FR-005", "The solution shall retrieve accessible candidate records using containment, prefix, and suffix patterns and shall score candidates using normalized Levenshtein similarity.", "Must"),
        ("FR-006", "The solution shall remove common trailing corporate suffixes from the fuzzy comparison value without altering the saved Account name.", "Must"),
        ("FR-007", "D&B results shall display company name, DUNS number, match score, and available address context.", "Must"),
        ("FR-008", "Existing Account results shall display Account name, match score, and an action to open the Account.", "Must"),
        ("FR-009", "The solution shall block the Account form when an existing Account match is present.", "Must"),
        ("FR-010", "The solution shall require a D&B selection when D&B matches exist and no Account match exists.", "Must"),
        ("FR-011", "Selecting D&B shall populate Account Name, Billing Street, Billing City, Billing State/Province, Billing Postal Code, and Billing Country when those values are available.", "Must"),
        ("FR-012", "The user shall be able to review and edit pre-populated values before save.", "Must"),
        ("FR-013", "The form shall also accept Phone and Website values.", "Should"),
        ("FR-014", "The server shall verify that a supplied D&B identifier is a valid, accessible DB_Data__c record.", "Must"),
        ("FR-015", "The server shall repeat Account and D&B matching before insert and return a clear blocking message when the state has changed.", "Must"),
        ("FR-016", "The Account insert shall populate Account.D_B_Record__c when a D&B parent was selected.", "Must"),
        ("FR-017", "The solution shall permit multiple distinctly named Accounts to reference the same D&B parent.", "Must"),
        ("FR-018", "After successful insert, the solution shall navigate to the new Account record.", "Should"),
        ("FR-019", "Errors and progress states shall be announced in an accessible live region and presented in plain language.", "Must"),
        ("FR-020", "The duplicate trigger shall populate the normalized Account-name key on both insert and update.", "Must"),
    ]
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        functional,
        [1050, 7110, 1200],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.9,
    )

    page_break(doc)
    add_heading(doc, "9. Non-Functional Requirements", 1)
    nonfunctional = [
        ("NFR-001", "Responsive design", "The form and match cards shall reflow to a single-column layout on small screens without horizontal scrolling."),
        ("NFR-002", "Accessibility", "Use Salesforce base components, associated labels, keyboard-operable controls, visible validation, and live status/error announcements."),
        ("NFR-003", "Security", "Apex queries and DML shall honor record sharing, CRUD, field-level security, and the assigned permission set."),
        ("NFR-004", "Least privilege", "Access shall be delivered by the Guided Account Creation permission set and limited to the required Apex class, tab, objects, and fields."),
        ("NFR-005", "Performance", "Interactive searches shall be debounced, candidate retrieval shall be bounded, and no more than ten scored results per category shall be returned to the UI."),
        ("NFR-006", "Concurrency", "The unique normalized-name field shall reject concurrent exact normalized duplicates at the database layer."),
        ("NFR-007", "Maintainability", "Match threshold, normalization behavior, candidate limits, and mapped fields shall be documented and covered by Apex tests."),
        ("NFR-008", "Reliability", "Server-side validation shall be authoritative; client state shall never be sufficient to authorize creation."),
        ("NFR-009", "Compatibility", "The solution shall support Salesforce Lightning Experience and supported Salesforce Mobile form factors."),
        ("NFR-010", "Observability", "Standard Salesforce Created By, Created Date, field history where configured, and deployment/test evidence shall support operational review."),
    ]
    add_table(doc, ["ID", "Category", "Requirement"], nonfunctional, [1050, 1900, 6410], font_size=9.0)

    add_heading(doc, "10. Data Requirements and Model", 1)
    add_heading(doc, "10.1 Authoritative Relationship", 2)
    add_callout(doc, "Data model:", "DB_Data__c (one) -> Account (many), implemented through Account.D_B_Record__c. The existing DB_Data__c.Account__c lookup is not written by this solution and is not authoritative for this process.")
    add_body(doc, "Each Account may relate to zero or one D&B parent. Multiple Accounts may legitimately share a D&B parent when they represent distinct Salesforce Account records. Duplicate evaluation remains Account-name-based and is applied even when the same D&B parent is reused.")

    add_heading(doc, "10.2 Pre-Population Mapping", 2)
    add_table(
        doc,
        ["D&B source", "Account target", "Rule"],
        [
            ("DB_Data__c.Duns_Name__c", "Account.Name", "Populate on selection; user may review/edit before save"),
            ("DB_Data__c.Geocoded_Address__c", "Account.BillingStreet", "Populate when present"),
            ("DB_Data__c.Geocoded_City__c", "Account.BillingCity", "Populate when present"),
            ("DB_Data__c.Geocoded_State__c", "Account.BillingState", "Populate when present; validate against state configuration"),
            ("DB_Data__c.Geocoded_Zip__c", "Account.BillingPostalCode", "Populate when present"),
            ("DB_Data__c.Geocoded_Country__c", "Account.BillingCountry", "Populate when present; validate against country configuration"),
            ("DB_Data__c.Id", "Account.D_B_Record__c", "Persist selected parent lookup"),
        ],
        [2950, 2700, 3710],
        font_size=9.1,
    )

    add_heading(doc, "10.3 Duplicate-Key Data", 2)
    add_body(doc, "Account.Normalized_Account_Name__c is a unique text field populated by the AccountDuplicateGuard trigger. Its exact-key normalization lowercases Account Name and removes non-alphanumeric characters. This protects exact normalized duplicates such as 'ACME, Inc.' and 'acme inc' after the key has been populated.")
    add_callout(doc, "Production prerequisite:", "Before rollout, calculate the normalized value for existing Accounts, identify collisions, resolve or formally exempt those collisions, and backfill the field. Without this step, an old Account with a null normalized key cannot participate fully in database-level uniqueness, although the guided fuzzy search will still identify it.", fill=PALE_GOLD, accent=GOLD)

    add_heading(doc, "11. Security and Access", 1)
    for item in [
        "Users must be assigned the Guided Account Creation permission set.",
        "The permission set grants access to AccountCreationController, the Guided Account Creation tab, Account create/read/edit access, DB_Data__c read access, and the fields used by the solution.",
        "AccountCreationController runs with sharing. Candidate queries use USER_MODE and the insert uses user-mode DML.",
        "The system-maintained normalized field remains trigger-controlled; any user-provided value is overwritten during Account insert or update.",
        "Users only see and match records they are permitted to access. Business owners must decide whether incomplete visibility could allow duplicates across private Account sharing boundaries and, if needed, provide a data-steward operating model.",
    ]:
        add_bullet(doc, item)

    page_break(doc)
    add_heading(doc, "12. Acceptance Criteria", 1)
    acceptance = [
        ("AC-01", "No match", "Given a valid company name with no D&B or Account match, the form becomes available and a new Account can be created without D_B_Record__c."),
        ("AC-02", "D&B match", "Given one or more D&B matches and no Account match, creation remains unavailable until the user selects a D&B result."),
        ("AC-03", "Pre-population", "When a D&B result is selected, company name and available billing-address values are copied into the form for review."),
        ("AC-04", "Existing Account", "Given a fuzzy Account match at or above the threshold, creation is blocked and the user can open the existing Account."),
        ("AC-05", "Server recheck", "If a duplicate appears after the initial search but before save, server-side processing blocks the insert."),
        ("AC-06", "Exact duplicate", "A second Account with the same normalized exact name is rejected by the unique key."),
        ("AC-07", "One-to-many", "Two distinctly named Accounts can be created with the same DB_Data__c parent and both retain Account.D_B_Record__c."),
        ("AC-08", "Relationship integrity", "Selecting a D&B record does not update DB_Data__c.Account__c and does not reserve the D&B row."),
        ("AC-09", "Mobile", "The complete search, selection, form, save, and navigation flow is usable in Salesforce Mobile without horizontal scrolling."),
        ("AC-10", "Browser", "The component is accessible from the Guided Account Creation Lightning tab and operates in a supported desktop browser."),
        ("AC-11", "Accessibility", "All interactive controls have visible labels, keyboard operation is supported, and errors/status changes are announced."),
        ("AC-12", "Security", "A user without the permission set cannot access the Apex action/tab; an authorized user remains subject to sharing and field permissions."),
        ("AC-13", "Validation", "The focused Apex suite passes with at least 75% coverage for deployed classes; current scratch validation is 6/6 executions and 90% test-run coverage."),
    ]
    add_table(doc, ["ID", "Scenario", "Expected result"], acceptance, [1050, 1900, 6410], font_size=8.9)

    add_heading(doc, "13. Assumptions and Dependencies", 1)
    add_heading(doc, "13.1 Assumptions", 2)
    for item in [
        "DB_Data__c contains sufficiently complete Duns_Name__c values to support name matching.",
        "Users can distinguish the correct D&B organization using company name, DUNS number, match score, and address context.",
        "The 74% threshold is an initial business setting and will be validated with representative regional and multilingual names during UAT.",
        "Account record types, validation rules, state/country picklists, flows, and triggers in the target org accept the mapped values or will be addressed during deployment testing.",
        "Guided Account Creation will be the preferred user channel for manual Account creation where duplicate controls are required.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "13.2 Dependencies", 2)
    for item in [
        "DB_Data__c and the required D&B fields are deployed and populated.",
        "Account.D_B_Record__c exists and represents the child-side reference to D&B.",
        "The Guided Account Creation permission set is assigned to approved users.",
        "The Guided Account Creation custom tab is added to the relevant desktop and mobile app navigation menus.",
        "Legacy Account normalization and collision remediation are completed before production cutover.",
        "Business data stewards provide UAT examples for punctuation, suffixes, abbreviations, transliterations, and common misspellings.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "14. Risks, Controls, and Open Decisions", 1)
    risks = [
        ("Legacy normalized keys are null", "Exact API duplicate may not collide with a legacy row", "Assess, remediate, and backfill before production", "Open"),
        ("Threshold too low", "False positives block legitimate creation", "Tune with UAT corpus; provide data-steward escalation", "Open"),
        ("Threshold too high", "Near-duplicates are missed", "Monitor duplicate outcomes and retest threshold", "Open"),
        ("Private sharing hides Accounts", "User may not see an existing duplicate", "Define steward visibility or centralized creation model", "Decision"),
        ("Large D&B dataset", "Prefix/contains search may slow or omit candidates", "Performance test; consider a precomputed search key or SOSL strategy", "Open"),
        ("State/country picklists", "D&B address value may fail Account validation", "Map/canonicalize values in UAT and production readiness", "Open"),
        ("Direct standard Account creation", "Users can bypass fuzzy wizard logic", "Restrict manual entry points or extend enterprise Duplicate Rules", "Decision"),
        ("Hard-coded threshold", "Business tuning requires code deployment", "Evaluate Custom Metadata configuration as an enhancement", "Future"),
    ]
    add_table(doc, ["Risk / decision", "Impact", "Control / action", "Status"], risks, [2250, 2600, 3410, 1100], font_size=8.5)

    page_break(doc)
    add_heading(doc, "15. Deployment and Rollout Requirements", 1)
    rollout = [
        "Confirm the target production data model and retrieve any org-specific Account validations, automation, and address controls.",
        "Profile existing Accounts using the normalized-name algorithm; identify duplicate-key collisions and fuzzy near-duplicates.",
        "Resolve or formally govern collisions, then backfill Account.Normalized_Account_Name__c in a controlled batch.",
        "Deploy the field, trigger, Apex controller and tests, LWC, custom tab, and permission set through the approved release process.",
        "Run all local tests plus the focused AccountCreationControllerTest suite and confirm required coverage.",
        "Assign the permission set to pilot users and add the custom tab to desktop and mobile app navigation.",
        "Execute UAT using representative company names, D&B parents with multiple Accounts, address variations, existing duplicates, and mobile devices.",
        "Train users on match selection, blocked Account handling, and escalation to the Customer Data Steward.",
        "Monitor the first release window for blocked saves, search latency, false positives, false negatives, and bypass through other Account-creation channels.",
    ]
    for item in rollout:
        add_numbered(doc, item)

    add_heading(doc, "16. Test Evidence", 1)
    add_body(doc, "The implementation was deployed to newMBD-scratch. The focused AccountCreationControllerTest execution completed successfully on 10 August 2026 with 6 of 6 executions passing and 90% test-run coverage. Test scenarios cover fuzzy D&B matching, mandatory D&B selection, creation with no match, existing Account blocking, normalized uniqueness, and multiple distinct Accounts related to one D&B parent.")
    add_callout(doc, "Release gate:", "Production approval must also include org-specific validation, regression testing, legacy-key backfill evidence, desktop/mobile UAT, and approval of the open decisions in Section 14.")

    add_heading(doc, "17. Approvals", 1)
    add_body(doc, "Approval confirms that the requirements, business rules, assumptions, and production prerequisites in this document are accepted for implementation and rollout planning.")
    add_table(
        doc,
        ["Role", "Name", "Decision", "Date"],
        [
            ("Business Sponsor", "", "Approve / Reject", ""),
            ("Salesforce Product Owner", "", "Approve / Reject", ""),
            ("Customer Data Steward", "", "Approve / Reject", ""),
            ("Information Security", "", "Approve / Reject", ""),
            ("Technical Lead", "", "Approve / Reject", ""),
        ],
        [2400, 2600, 2400, 1960],
        font_size=9.3,
    )

    add_heading(doc, "Appendix A. Implementation Traceability", 1)
    trace = [
        ("Guided Account Creation LWC", "force-app/main/default/lwc/accountCreationWizard", "FR-001 to FR-013, FR-018, FR-019"),
        ("AccountCreationController", "force-app/main/default/classes/AccountCreationController.cls", "FR-005 to FR-017"),
        ("AccountDuplicateGuard", "Apex class and Account before-insert/before-update trigger", "BR-12, FR-020, NFR-006"),
        ("Normalized Account Name", "Account.Normalized_Account_Name__c", "BR-12, NFR-006"),
        ("Permission set", "Guided_Account_Creation", "NFR-003, NFR-004, AC-12"),
        ("Custom tab", "Guided_Account_Creation", "FR-001, AC-10"),
        ("Focused tests", "AccountCreationControllerTest", "AC-01 to AC-08, AC-13"),
    ]
    add_table(doc, ["Component", "Implementation", "Requirement coverage"], trace, [2600, 3800, 2960], font_size=8.8)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
