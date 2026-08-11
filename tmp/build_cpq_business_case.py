from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUT = Path(r"C:\Users\JONGB\OneDrive - Thermo Fisher Scientific\Documents\Playground\newMBD\output\newMBD_CPQ_Business_Case.docx")
LIFECYCLE_IMG = Path(r"C:\Users\JONGB\Downloads\E2E Commercial Execution Lifecycle.png")
OUT.parent.mkdir(parents=True, exist_ok=True)
if not LIFECYCLE_IMG.exists():
    raise FileNotFoundError(f"Lifecycle image not found: {LIFECYCLE_IMG}")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MID_BLUE = "DCE8F3"
LIGHT_BLUE = "EEF4F9"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE5"
TEXT = "1F2933"
MUTED = "667085"
WHITE = "FFFFFF"
GREEN = "1F6D4C"
AMBER = "7A5A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.insert_element_before(
            shd, "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"
        )
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.insert_element_before(tc_mar, "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, italic=None, color=TEXT, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_geometry(table, widths_dxa, indent_dxa=120, borders=True):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.insert_element_before(layout, "w:tblCellMar", "w:tblLook")
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.insert_element_before(tbl_ind, "w:tblBorders", "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.insert_element_before(tbl_borders, "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tbl_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tbl_borders.append(node)
        node.set(qn("w:val"), "single" if borders else "nil")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), MID_GRAY)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_numbering(doc, num_fmt, lvl_text):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), lvl_text)
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    if num_fmt == "bullet":
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        rpr.append(fonts)
        lvl.append(rpr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        first_num.addprevious(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_el = OxmlElement("w:numId")
    num_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_el)
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_para(doc, text="", bold=False, italic=False, color=TEXT, size=11, align=None, after=6, before=0, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_mixed_para(doc, parts, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    for text, bold, italic, color in parts:
        r = p.add_run(text)
        set_run_font(r, bold=bold, italic=italic, color=color)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.10
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), accent)
    left.set(qn("w:space"), "8")
    borders.append(left)
    p_pr.insert_element_before(borders, "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:spacing", "w:ind")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.insert_element_before(shd, "w:tabs", "w:suppressAutoHyphens", "w:spacing", "w:ind")
    r1 = p.add_run(label + "  ")
    set_run_font(r1, bold=True, color=NAVY)
    r2 = p.add_run(text)
    set_run_font(r2, color=TEXT)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(text)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            if row_idx % 2 == 1:
                set_cell_shading(cells[idx], "FAFBFC")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            set_run_font(r, size=font_size, color=TEXT)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
doc.settings.odd_and_even_pages_header_footer = True
section = doc.sections[0]
section.different_first_page_header_footer = False
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Standard business brief style tokens.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Title", 28, NAVY, 0, 8),
    ("Subtitle", 14, MUTED, 0, 12),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.10
    style.paragraph_format.keep_with_next = True

if "Source Text" not in [s.name for s in doc.styles]:
    source_style = doc.styles.add_style("Source Text", WD_STYLE_TYPE.PARAGRAPH)
else:
    source_style = doc.styles["Source Text"]
source_style.font.name = "Calibri"
source_style.font.size = Pt(9)
source_style.font.color.rgb = RGBColor.from_string(MUTED)
source_style.paragraph_format.space_before = Pt(4)
source_style.paragraph_format.space_after = Pt(4)

bullet_num_id = add_numbering(doc, "bullet", "\uf0b7")
decimal_num_id = add_numbering(doc, "decimal", "%1.")
target_decimal_num_id = add_numbering(doc, "decimal", "%1.")

# Running page furniture.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
hr = hp.add_run("newMBD  |  Commercial Platform Business Case")
set_run_font(hr, size=8.5, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.paragraph_format.space_before = Pt(0)
fr = fp.add_run("August 2026  |  Page ")
set_run_font(fr, size=8.5, color=MUTED)
add_page_field(fp)

# Explicit even-page furniture prevents LibreOffice from applying an implicit
# zero-margin even-page style while preserving the same design on every page.
even_header = section.even_page_header
ehp = even_header.paragraphs[0]
ehp.alignment = WD_ALIGN_PARAGRAPH.LEFT
ehp.paragraph_format.space_after = Pt(0)
ehr = ehp.add_run("newMBD  |  Commercial Platform Business Case")
set_run_font(ehr, size=8.5, bold=True, color=MUTED)

even_footer = section.even_page_footer
efp = even_footer.paragraphs[0]
efp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
efp.paragraph_format.space_before = Pt(0)
efr = efp.add_run("August 2026  |  Page ")
set_run_font(efr, size=8.5, color=MUTED)
add_page_field(efp)

# Cover - proposal_centerpiece pattern.
add_para(doc, "newMBD", bold=True, color=BLUE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=12, before=34)
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Global Commercial Configuration,\nPricing and Renewal Platform")
p.paragraph_format.space_after = Pt(8)
sub = doc.add_paragraph(style="Subtitle")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Business case for replacing P360 and Impact and enabling eCommerce")

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(8)
rule.paragraph_format.space_after = Pt(18)
p_pr = rule._p.get_or_add_pPr()
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:color"), BLUE)
p_bdr.append(bottom)
p_pr.insert_element_before(p_bdr, "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:spacing", "w:ind")

meta = doc.add_table(rows=3, cols=4)
meta_rows = [
    ("Prepared for", "Leadership team and process stakeholders", "Decision horizon", "12-month TSA"),
    ("TSA start", "12 August 2026", "Required cutover", "12 August 2027"),
    ("Evaluation", "Three years, USD", "Scope", "Global / approximately 250 users"),
]
for i, row in enumerate(meta_rows):
    for j, value in enumerate(row):
        p = meta.rows[i].cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=9.5, bold=(j % 2 == 0), color=(BLUE if j % 2 == 0 else TEXT))
        if j % 2 == 0:
            set_cell_shading(meta.rows[i].cells[j], LIGHT_GRAY)
set_table_geometry(meta, [1350, 3330, 1350, 3330], borders=False)
set_repeat_table_header(meta.rows[0])

add_callout(doc, "Decision required", "Approve immediate competitive validation and implementation funding for a global replacement capability. Do not preselect a product until it proves SAP pricing, automated high-volume renewals and reagent-rental capability.")
add_para(doc, "Version 1.1  |  5 August 2026", color=MUTED, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, before=20)
doc.add_page_break()

# Executive summary.
doc.add_heading("Executive recommendation", level=1)
add_para(doc, "Approve an immediate programme to select and implement a global commercial configuration, pricing and renewal platform before the P360/Impact exit on 12 August 2027.")
add_callout(doc, "Core conclusion", "The capability replacement is mandatory for business continuity. Salesforce Revenue Management is a strong candidate because of CRM and commerce alignment, but it should be selected only after a proof of capability using newMBD data and deal scenarios.")

for item in [
    "P360 and Impact remain with Thermo Fisher and cannot transfer to newMBD.",
    "The supported TSA lasts 12 months and will not be extended.",
    "Both interactive quoting and automated pricing-renewal capabilities must operate at cutover.",
    "The target is global, serves approximately 250 users, retains Salesforce as CRM and SAP ECC6 as the pricing and availability authority.",
    "eCommerce should follow shortly after cutover and reuse the same product, package and pricing services.",
]:
    add_list_item(doc, item, bullet_num_id)

add_para(doc, "Leadership should not approve the legacy Salesforce CPQ managed package as a predetermined product. Salesforce states that it is end-of-sale to new customers and that strategic investment has shifted to Revenue Cloud Advanced, also described in current documentation as Revenue Management. Licensing entitlement after divestiture remains to be established.")

doc.add_heading("Approval requested", level=2)
for item in [
    "Immediate funding for a competitive fit-gap and proof-of-capability phase.",
    "A global replacement programme targeting readiness ahead of 12 August 2027.",
    "Selection based on demonstrated recommended pricing, customer/SKU sales history, cost-based margin control, automated SAP renewal loading and reagent-rental outcomes.",
    "A commerce-ready design without committing yet to an eCommerce vendor.",
    "A three-year USD investment case combining least-cost continuity, margin protection, capital utilization, productivity and growth.",
]:
    add_list_item(doc, item, target_decimal_num_id)

# Context.
doc.add_heading("1. Business context and case for change", level=1)
doc.add_heading("1.1 Divestiture-driven business continuity", level=2)
add_para(doc, "P360 is Thermo Fisher's bespoke CPQ platform. Impact supports pricing agreements and renewals. Both remain with Thermo Fisher after the divestiture. Both tools currently provide recommended prices and customer/SKU sales history, while product cost is used to calculate margins. Impact automatically loads approved renewals to SAP ECC6. newMBD can extract product, pricing, customer, quote and contract data from P360, together with pricing and configuration rules. From Impact, current and historical renewal and pricing-agreement data can be extracted; the underlying calculation, workflow and automated SAP-load logic are not confirmed as transferable.")
add_para(doc, "The resulting investment is not an optional transformation. Without a replacement at TSA exit, newMBD would lose essential capabilities to generate valid quotes and renew customer pricing. The financial case must therefore distinguish mandatory revenue and margin protection from incremental efficiency or growth benefits.")

doc.add_heading("1.2 Current-state complexity", level=2)
add_para(doc, "The still-current functionality-gap workbook records significant concerns across global pricing structures, renewals, SAP integration, master data, approvals, reporting and operational support. All items currently marked unacceptable are treated as Day-1 requirements, while the business process behind each item must be challenged and simplified rather than copied screen-for-screen.")
for item in [
    "Correct SAP-derived list prices and current/future validity dates.",
    "Recommended prices and sales history at customer/SKU level.",
    "Product cost inputs, margin calculation and margin-based controls.",
    "Sold-to, ship-to, customer hierarchy, CPG, end-user, GPO/template and VSP pricing.",
    "US rebates, Fisher net pricing and product restrictions by sales organization.",
    "Bulk pricing renewals, exception management and approval routing.",
    "Error monitoring, reporting, auditability and support service levels.",
]:
    add_list_item(doc, item, bullet_num_id)

doc.add_heading("1.3 Design principle", level=2)
add_callout(doc, "Simplify, do not clone", "The target should not create Salesforce replicas of two bespoke tools. It should establish a single commercial operating model, while allowing a specialized renewal component where bulk processing is better served outside an interactive CPQ workflow.")

# Target model.
doc.add_heading("2. Proposed target operating model", level=1)
for item in [
    "Salesforce provides customer, opportunity and user context.",
    "SAP ECC6 remains authoritative for final pricing, product availability and customer-specific condition records.",
    "The selected CPQ capability manages configuration, commercial controls, approvals and quotations.",
    "A dedicated bulk-renewal capability may be used where it provides a simpler and more reliable outcome than interactive CPQ.",
    "Approved renewals must continue to load automatically to SAP ECC6, with reconciliation, locking, failure handling and operational monitoring.",
    "Recommended prices, customer/SKU sales history and product cost must be available in the selling workflow so margin is calculated consistently.",
    "The product, package and pricing services are designed for reuse by eCommerce shortly after cutover.",
]:
    add_list_item(doc, item, decimal_num_id)

add_callout(doc, "Architecture guardrail", "A separate renewal component is acceptable, but automated SAP loading and the current pricing intelligence must not regress. The selection should optimize the end-to-end business outcome rather than force all functionality into one product.")

# Lifecycle alignment.
doc.add_heading("3. Commercial execution lifecycle alignment", level=1)
add_para(doc, "The target capability should be designed as a commercial lifecycle platform, not only a quote generator. It must connect opportunity qualification; solution, pricing and proposal development; evaluation and contracting; operational handover; and commercial performance, amendments and renewals. Performance and customer insights should feed the next account-planning and opportunity cycle.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
r = p.add_run()
picture = r.add_picture(str(LIFECYCLE_IMG), width=Inches(6.5))
picture._inline.docPr.set("title", "Commercial execution lifecycle")
picture._inline.docPr.set("descr", "Six-stage commercial execution lifecycle from territory and account strategy through opportunity qualification, solution pricing and proposal, contracting, implementation, and commercial performance and lifecycle management, with continuous insights and governance across all stages.")
add_para(doc, "Figure 1. End-to-end commercial execution lifecycle (provided by newMBD).", italic=True, color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
for item in [
    "Stages 2-4: qualify opportunities, apply recommended prices and customer/SKU sales history, configure reagent rentals or packages, calculate cost-based margin, route approvals and produce compliant offers.",
    "Stage 5: hand over an awarded deal to SAP and operations with complete customer, product, price, asset, commitment, invoicing and fulfillment data.",
    "Stage 6: monitor obligations, reagent purchases, revenue, margin, pricing agreements, amendments and renewals; automatically load approved renewal outcomes to SAP.",
    "Across all stages: use governance gates, approvals, auditable data, reporting, compliance controls and a continuous insights loop to improve future account and pricing decisions.",
]:
    add_list_item(doc, item, bullet_num_id)
add_callout(doc, "Design implication", "CPQ is one component of the lifecycle. The business outcome depends equally on pricing intelligence, contract and asset obligations, automated SAP execution, performance monitoring and renewal feedback.")

doc.add_heading("4. Strategic benefits and differentiators", level=1)
doc.add_heading("4.1 Reagent-rental contracts", level=2)
add_para(doc, "newMBD retains ownership of capital assets placed at no charge in return for a one-year customer commitment to purchase products. Commitments may cover spend, product quantity, test volume, product mix and periodic minimums. Asset economics are recovered through reduced reagent discounting, but compliance with purchase commitments is not currently monitored.")
add_para(doc, "A modern commercial platform can connect the asset, commitment, term and applicable products; calculate deal economics; route exceptions; monitor actual reagent purchases; identify shortfalls before expiry; and support renewal, recovery or renegotiation. The result is stronger margin discipline and more productive use of retained capital.")

doc.add_heading("4.2 Reagent package deals", level=2)
add_para(doc, "Reagent packages are a new capability intended to increase sales volume. The package structure and eCommerce transaction model remain open. The chosen platform should support both fixed and configurable bundles without bespoke development for every new offer.")

doc.add_heading("4.3 Pricing intelligence, margin and renewals", level=2)
add_para(doc, "P360 and Impact provide recommended prices and customer/SKU sales history, and product cost information is used to calculate margins. These are baseline commercial controls—not optional analytics—and must remain available in the future workflow for pricing decisions, approvals and performance reporting.")
add_para(doc, "Impact automatically loads approved renewals to SAP ECC6. Direct EU alone processes almost 7,000 renewals in the run-up to year-end, and sales representatives spend three to four months completing the year-end renewal cycle. The replacement must preserve automated SAP execution while reducing elapsed time, rep workload and exception handling. The 187-250 VSP agreements identified previously should be treated as one agreement subset, not as the total renewal volume.")

doc.add_heading("4.4 eCommerce enablement", level=2)
add_para(doc, "eCommerce is planned shortly after the CPQ cutover. Designing the commercial platform and product model once avoids creating separate configuration and pricing logic for assisted sales and digital channels. The initial online journey may support direct purchases, configurable packages, negotiated quotes or a combination; no channel model is assumed in this case.")

# Financial case.
doc.add_heading("5. Financial case", level=1)
add_para(doc, "No unsupported values have been introduced. The approval model should populate the following calculations with validated stakeholder data and keep hard savings, capacity benefits, margin protection and growth benefits separate.")

benefit_rows = [
    ("Revenue continuity", "Margin dependent on quoting and renewals x expected disruption avoided", "Mandatory risk protection"),
    ("User productivity", "250 users x annual hours saved x loaded hourly cost x 3 years", "Capacity; hard only if actioned"),
    ("Renewal capacity", "Renewals x handling time reduction x loaded hourly cost", "Productivity / capacity"),
    ("Renewal automation preservation", "Renewals x manualization avoided x loaded hourly cost", "Continuity / productivity protection"),
    ("Renewal timeliness", "Revenue renewed earlier x protected uplift or leakage avoided", "Margin / revenue protection"),
    ("Recommended-price adoption", "Eligible revenue x adoption improvement x realized price improvement", "Gross-margin improvement"),
    ("Sales-history insight", "Eligible customer/SKU base x cross-sell or package uplift x gross margin", "Incremental profit"),
    ("Cost-based margin control", "Revenue exposed to low-margin exceptions x avoidable margin leakage", "Gross-margin protection"),
    ("Pricing-error reduction", "Current annual error cost x expected reduction", "Hard saving"),
    ("Rental commitment recovery", "Identified shortfalls x achievable recovery margin", "Margin protection"),
    ("Rental discount improvement", "Rental-linked revenue x avoidable discount reduction", "Gross-margin improvement"),
    ("Capital utilization", "Placed-asset value x improvement in productive utilization", "Capital efficiency"),
    ("Package growth", "Eligible reagent revenue x adoption x volume uplift x gross margin", "Incremental profit"),
    ("eCommerce cost-to-serve", "Digital orders x assisted-to-digital handling-cost difference", "Hard/capacity benefit"),
    ("eCommerce growth", "Digital incremental revenue x gross margin", "Incremental profit"),
    ("Bespoke-build avoidance", "Bespoke rebuild and run cost less selected-platform cost", "Cost avoidance"),
]
add_table(doc, ["Benefit", "Three-year calculation", "Classification"], benefit_rows, [2200, 4600, 2560], font_size=8.8)

add_callout(doc, "Avoid double counting", "Employee time saved is not both a cash saving and additional selling capacity unless leadership commits to a corresponding capacity action. Revenue growth and margin improvement should also be isolated from avoided disruption.", fill="FFF7E6", accent=AMBER)

doc.add_heading("5.1 Three-year cost framework", level=2)
for item in [
    "Product licenses for users who genuinely require full functionality.",
    "Implementation, configuration and programme management.",
    "SAP ECC6 and Salesforce integration.",
    "P360 and Impact data migration and reconciliation.",
    "Reconstruction and validation of Impact renewal logic.",
    "Global master-data remediation, testing and deployment.",
    "Training, support, administration and enhancement capacity.",
    "eCommerce enablement and programme contingency.",
]:
    add_list_item(doc, item, bullet_num_id)

doc.add_heading("5.2 Licensing sensitivity", level=2)
add_para(doc, "Salesforce currently publishes list pricing of USD 150 per user/month for Revenue Cloud Growth and USD 200 per user/month for Revenue Cloud Advanced. If all 250 users required full licenses, the purely illustrative three-year list-license range would be USD 1.35 million to USD 1.80 million. This excludes discounts, implementation, support, Commerce licenses, taxes and any existing entitlements; it is a sensitivity boundary, not a newMBD estimate.")
add_para(doc, "Salesforce also states that B2B product configuration requires applicable Commerce licensing plus Revenue Cloud Advanced. Commerce licensing and architecture must therefore be considered during selection rather than added after the CPQ design is fixed.")

# Options.
doc.add_heading("6. Options to evaluate", level=1)
option_rows = [
    ("Salesforce Revenue Management + renewal component", "Strong CRM and commerce alignment; bulk renewal remains optimized", "Multiple components; SAP/global structures require proof"),
    ("Salesforce Revenue Management for quoting and renewals", "More unified user and data model", "Risk of forcing mass renewal into an interactive CPQ pattern"),
    ("Third-party CPQ / price-management platform", "May offer stronger SAP pricing or agreement support", "Additional platform and commerce integration complexity"),
    ("Bespoke rebuild", "Maximum functional fidelity", "Highest delivery, maintenance and future-commerce risk"),
    ("No replacement", "No implementation spend", "Not viable after 12 August 2027"),
]
add_table(doc, ["Option", "Strategic strengths", "Principal concerns"], option_rows, [2500, 3430, 3430], font_size=9.0)

add_callout(doc, "Selection position", "No product decision has been made. Salesforce Revenue Management should be treated as a leading candidate, not as an approved solution.")

# Proof.
doc.add_heading("7. Mandatory proof of capability", level=1)
add_para(doc, "Selection must use executable demonstrations with migrated sample data and representative deal scenarios. Written claims that a requirement is configurable are insufficient.")
proof_items = [
    "Correct SAP-derived list price and current/future validity.",
    "Recommended prices and current sales history by customer/SKU.",
    "SAP product cost, margin calculation, margin warnings and approval thresholds.",
    "Sold-to, ship-to, customer hierarchy, CPG and end-user pricing.",
    "US GPO/template pricing, rebates, Fisher net pricing and VSP.",
    "Excluded or withdrawn products by sales organization.",
    "Approval routing, exception management and coverage delegation.",
    "SAP corrections, locking, failures, reconciliation and error monitoring.",
    "Performance test for almost 7,000 direct-EU year-end renewals, including customer-specific exceptions and automated SAP load.",
    "A complete reagent-rental deal and purchase-commitment monitoring.",
    "A configurable reagent package with commerce-compatible pricing.",
    "Operational reporting, audit trail, support and enhancement model.",
]
for item in proof_items:
    add_list_item(doc, item, bullet_num_id)

# Delivery.
doc.add_heading("8. Delivery approach", level=1)
roadmap_rows = [
    ("Aug-Oct 2026", "Challenge processes; profile data; fit-gap; vendor proofs", "Validated requirements and shortlist"),
    ("Oct-Dec 2026", "Select solution; confirm architecture and contract", "Approved design and mobilized delivery"),
    ("Dec 2026-Apr 2027", "Configure; build renewal capability; migrate and integrate", "Feature-complete solution"),
    ("Mar-Jun 2027", "System testing; global validation; migration rehearsals", "Proven end-to-end operation"),
    ("May-Jul 2027", "User acceptance; training; controlled pilot; cutover rehearsal", "Operational readiness"),
    ("Jul-Aug 2027", "Final migration; production cutover; stabilization", "TSA exit by 12 Aug 2027"),
]
add_table(doc, ["Timing", "Primary activity", "Exit outcome"], roadmap_rows, [1800, 4700, 2860], font_size=9.2)
add_para(doc, "Workstreams must overlap. Optional document generation, electronic signature and unnecessary recreation of P360 screens should not consume capacity needed for pricing accuracy, renewals, data readiness and cutover.")

# Risks.
doc.add_heading("9. Principal risks and controls", level=1)
risk_rows = [
    ("Hard 12-month deadline", "No extension and global scope", "Early selection; parallel workstreams; rehearsed cutover"),
    ("Limited incumbent expertise", "P360/Impact SMEs assumed unavailable", "Data-led discovery; stakeholder validation; reverse-engineering controls"),
    ("Undocumented renewal logic", "Only Impact data/history confirmed transferable", "Dedicated renewal discovery and reconciliation workstream"),
    ("Renewal scale or performance regression", "Almost 7,000 direct-EU year-end renewals and a three-to-four-month rep cycle", "Volume testing; automated SAP load; exception dashboards; cutover rehearsal"),
    ("Functional regression", "Recommended prices, customer/SKU history and cost-based margin are current controls", "Baseline inventory; sample-data proof; reconciled acceptance tests"),
    ("Requirement overload", "All current unacceptable gaps are Day 1", "Challenge outcomes; defer optional presentation features"),
    ("SAP/master-data complexity", "Account levels, validity, duplicates and restrictions", "Data remediation; proof with production-like scenarios"),
    ("Premature product selection", "Legacy CPQ terminology may obscure current products", "Competitive proof and licensing confirmation before award"),
    ("Commerce rework", "eCommerce follows shortly after cutover", "Common catalog, package and pricing services from the outset"),
]
add_table(doc, ["Risk", "Why it matters", "Required control"], risk_rows, [2250, 3200, 3910], font_size=9.0)

# Measures and conclusion.
doc.add_heading("10. Success measures", level=1)
for item in [
    "Production cutover completed before 12 August 2027 without loss of quoting or renewal capability.",
    "All Day-1 pricing and account-level scenarios pass agreed acceptance tests.",
    "Almost 7,000 direct-EU year-end renewals can be processed at current or better scale without loss of control.",
    "Approved renewals load automatically to SAP ECC6 with reconciliation and visible exception ownership.",
    "Recommended prices and sales history are available by customer/SKU, and product-cost-driven margins reconcile to the agreed source.",
    "SAP pricing outputs reconcile accurately and operational errors are visible and owned.",
    "Reagent-rental commitments, actual purchases and exceptions are visible to responsible teams.",
    "Reagent packages can be introduced without bespoke development for every offer.",
    "The same commercial services can support eCommerce shortly after cutover.",
]:
    add_list_item(doc, item, bullet_num_id)

doc.add_heading("11. Conclusion", level=1)
add_para(doc, "newMBD must replace the P360 and Impact capabilities to preserve its ability to trade after the TSA. The investment also creates a strategic opportunity to establish one governed commercial lifecycle: preserve recommended pricing, sales-history insight, cost-based margin control and automated renewals; protect reagent-rental economics; introduce reagent packages; and create the commercial foundation for eCommerce.")
add_para(doc, "The appropriate decision is not to approve a named CPQ product immediately. Leadership should approve a time-bound selection and implementation programme in which Salesforce Revenue Management and credible alternatives must demonstrate the required outcomes against SAP ECC6, global pricing structures, bulk renewals and reagent-rental contracts.")
add_callout(doc, "Recommended decision", "Fund the replacement programme now; select the product only after proof. Optimize for continuity by August 2027 and reuse by eCommerce thereafter.", fill="EAF5EF", accent=GREEN)

# Sources.
doc.add_heading("Sources", level=1)
sources = [
    ("Salesforce - Navigating the Future of Salesforce CPQ: Product End of Sale (Not End of Life)", "https://www.salesforce.com/sales/cpq/end-of-life/"),
    ("Salesforce Help - Configure Your Products in Revenue Management", "https://help.salesforce.com/s/articleView?id=ind.product_configurator_introduction.htm&language=en_US&type=5"),
    ("Salesforce Help - Configurable Products and Bundles for B2B Stores", "https://help.salesforce.com/s/articleView?id=commerce.comm_product_configurator.htm&language=en_US&type=5"),
    ("Salesforce - Revenue Cloud Pricing", "https://www.salesforce.com/sales/revenue-lifecycle-management/revenue-optimization-pricing/"),
    ("Salesforce Help - Set Up Product Configurator for B2B Stores", "https://help.salesforce.com/s/articleView?id=commerce.comm_product_configurator_setup.htm&language=en_US&type=5"),
]
for title, url in sources:
    p = doc.add_paragraph(style="Source Text")
    add_hyperlink(p, title, url)
    p.add_run("\n" + url)
add_para(doc, "Internal sources: Copy of CPQ functionality gaps.xlsx, Sheet1, rows 2-36; E2E Commercial Execution Lifecycle.png; and August 2026 stakeholder comments supplied by Paul. Requirements remain subject to challenge and proof.", size=9, color=MUTED, after=4)

doc.core_properties.title = "newMBD Global Commercial Configuration, Pricing and Renewal Platform - Business Case"
doc.core_properties.subject = "Business case for replacing P360 and Impact and enabling eCommerce"
doc.core_properties.author = "newMBD"
doc.core_properties.keywords = "newMBD, CPQ, Revenue Management, SAP ECC6, Salesforce, P360, Impact, eCommerce"

doc.save(OUT)
print(OUT)
