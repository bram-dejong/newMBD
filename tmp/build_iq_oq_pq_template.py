from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path(r"C:\Users\JONGB\OneDrive - Thermo Fisher Scientific\Documents\Playground\newMBD\output\newMBD_Application_IQ_OQ_PQ_Signoff_Template.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "5E6A75"
WHITE = "FFFFFF"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"
BORDER = "AAB4BE"
PLACEHOLDER = "[Enter response]"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths_dxa[min(i, len(widths_dxa)-1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=10, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, size=9.1, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], header_fill)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(values):
            set_cell_text(row.cells[i], str(value), size=font_size)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_key_value_table(doc, pairs, widths=(2500, 6860)):
    rows = []
    for label, value in pairs:
        rows.append((label, value))
    table = add_table(doc, ("Field", "Response"), rows, list(widths), header_fill=LIGHT_GRAY, font_size=9.5)
    for row in table.rows[1:]:
        set_cell_shading(row.cells[0], CALLOUT)
        for r in row.cells[0].paragraphs[0].runs:
            r.bold = True
    return table


def add_heading(doc, text, level=1, keep=True):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = keep
    return p


def add_body(doc, text, italic=False, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=10.5, italic=italic, color=color)
    return p


def add_callout(doc, label, text, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color=fill, size=2)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_text(cell, "", size=10)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    set_font(r, size=10, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_font(run, size=8.5, color=MUTED)


def add_signature_table(doc, roles):
    rows = []
    for role in roles:
        rows.append((role, "[Name]", "[Approved / Approved with conditions / Rejected]", "[Signature or e-approval reference]", "[Date]"))
    return add_table(doc, ("Approver role", "Name", "Decision", "Signature / approval reference", "Date"), rows,
                     [1900, 1350, 2350, 2500, 1260], header_fill=LIGHT_BLUE, font_size=8.6)


def add_test_case_table(doc, phase, count=5):
    rows = []
    for i in range(1, count + 1):
        rows.append((f"{phase}-{i:02d}", "[Requirement ID(s)]", "[Test objective / scenario]", "[Expected result]", "[Pass / Fail / N/A]", "[Evidence ID / deviation ID]"))
    return add_table(doc, ("Test ID", "Requirement", "Test objective / scenario", "Expected result", "Result", "Evidence / deviation"), rows,
                     [900, 1250, 2400, 2100, 1050, 1660], header_fill=LIGHT_BLUE, font_size=8.3)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

# Compact reference guide token map.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

title_style = styles["Title"]
title_style.font.name = "Calibri"
title_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
title_style.font.size = Pt(24)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor.from_string(INK)
title_style.paragraph_format.space_after = Pt(5)

subtitle = styles.add_style("Document Subtitle", WD_STYLE_TYPE.PARAGRAPH)
subtitle.font.name = "Calibri"
subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
subtitle.font.size = Pt(13)
subtitle.font.color.rgb = RGBColor.from_string(MUTED)
subtitle.paragraph_format.space_after = Pt(14)

# Running header/footer (memo masthead pattern, restrained for a reusable form).
header = sec.header
hp = header.paragraphs[0]
hp.text = "newMBD DIVESTITURE PROGRAM  |  CONTROLLED TEMPLATE"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.paragraph_format.space_after = Pt(0)
for run in hp.runs:
    set_font(run, size=8.5, bold=True, color=MUTED)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("newMBD | Application IQ/OQ/PQ Validation & Signoff  |  Page ")
set_font(run, size=8.5, color=MUTED)
add_page_field(fp)

# First-page masthead.
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("APPLICATION VALIDATION & RELEASE RECORD")
set_font(r, size=9.5, bold=True, color=BLUE)
p = doc.add_paragraph("IQ / OQ / PQ Signoff Template", style="Title")
p = doc.add_paragraph("Application migration or build-back for the establishment of newMBD following the MBD divestiture from Thermo Fisher Scientific", style="Document Subtitle")

add_key_value_table(doc, [
    ("Application / service name", "[Enter official name]"),
    ("newMBD application ID", "[Enter portfolio / CMDB ID]"),
    ("Migration disposition", "[ ] Migrate  [ ] Build back  [ ] Replace  [ ] Reconfigure  [ ] Retire / archive"),
    ("Business process / capability", PLACEHOLDER),
    ("Business owner", "[Name, title, function]"),
    ("IT service owner", "[Name, title, function]"),
    ("Validation lead", "[Name, title, function]"),
    ("Document ID / version", "[Document ID] / [Version]"),
    ("Target cutover / go-live", "[YYYY-MM-DD, time zone]"),
    ("Confidentiality", "Internal - newMBD transition program"),
])

add_callout(doc, "Purpose", "Use one completed record per application or clearly bounded application service. The record captures risk-based validation evidence and formal IT/business acceptance for installation qualification (IQ), operational qualification (OQ), performance qualification (PQ), and production release.")
add_callout(doc, "Tailoring note", "The validation lead and Quality/Compliance representative should tailor this template to the application's regulatory, GxP, privacy, cybersecurity, records-retention, and legal obligations. Mark a section N/A only with documented rationale and approval.", fill=AMBER)

add_heading(doc, "Approval gates at a glance", 1)
add_table(doc, ("Gate", "Minimum approval intent", "Required approvers"), [
    ("Protocol / plan approval", "Scope, risk, requirements, test strategy and acceptance criteria are suitable before execution.", "IT owner; business owner; Quality/Compliance when applicable"),
    ("IQ completion", "Target environment and installed/configured components are verified and controlled.", "IT owner; validation lead; Quality/Compliance when applicable"),
    ("OQ completion", "Functions, controls, interfaces and failure handling operate as specified.", "IT owner; business process owner; Quality/Compliance when applicable"),
    ("PQ completion", "End-to-end business processes perform acceptably with representative users/data.", "Business process owner; IT owner; Quality/Compliance when applicable"),
    ("Release authorization", "Residual risk is accepted and the application may enter newMBD production use.", "Business owner; IT owner; release authority; Quality/Compliance when applicable"),
], [1600, 4860, 2900], font_size=8.8)

doc.add_page_break()
add_heading(doc, "1. Document control and execution rules", 1)
add_heading(doc, "1.1 Revision history", 2)
add_table(doc, ("Version", "Date", "Author", "Change summary", "Status"), [
    ("0.1", "[YYYY-MM-DD]", "[Name]", "Initial draft", "Draft"),
    ("[x.x]", "[YYYY-MM-DD]", "[Name]", PLACEHOLDER, "[Draft / Approved / Superseded]"),
], [900, 1300, 1500, 3900, 1760])
add_heading(doc, "1.2 Execution conventions", 2)
add_table(doc, ("Rule", "Required practice"), [
    ("Contemporaneous records", "Record actual results, tester identity, execution date/time, environment and evidence reference when the test is performed."),
    ("Evidence integrity", "Store screenshots, logs, exports, approvals and reconciliations in the approved repository. Use immutable or version-controlled references where required."),
    ("Corrections", "Do not obscure the original record. Correct according to applicable good documentation practice and electronic-signature procedures."),
    ("Failures / deviations", "A failed step must reference a deviation or defect. Retesting requires documented impact assessment, approved resolution and traceable new evidence."),
    ("N/A use", "Enter N/A only with a rationale. The validation lead confirms that exclusion does not weaken the intended control."),
    ("Electronic approvals", "If signed electronically, record the system, workflow or approval reference and retain the audit trail."),
], [2100, 7260], font_size=9.2)
add_heading(doc, "1.3 Protocol / plan approval", 2)
add_body(doc, "By approving this plan, signatories confirm that the scope, requirements, risks, test coverage, acceptance criteria and assigned responsibilities are adequate for execution.")
add_signature_table(doc, ["IT service owner", "Business process owner", "Validation lead", "Quality / Compliance (if applicable)"])

doc.add_page_break()
add_heading(doc, "2. Application and transition profile", 1)
add_heading(doc, "2.1 Current-state and target-state summary", 2)
add_key_value_table(doc, [
    ("Source / current-state application", "[Name, owner, hosting location, tenant / instance]"),
    ("Target newMBD application", "[Name, owner, hosting location, tenant / instance]"),
    ("Current Thermo Fisher dependency", "[Infrastructure, identity, network, data, license, support, integration or process dependency]"),
    ("Separation approach", "[Migration / clone / fresh build / carve-out / replacement / coexistence]"),
    ("Transition service agreement dependency", "[TSA service, end date and exit criterion, or N/A]"),
    ("Data in scope", "[Master, transactional, historical, attachment, audit trail and configuration data]"),
    ("Users / locations", "[Populations, approximate volumes, sites, countries and time zones]"),
    ("Critical business periods", "[Month-end, batch release, audits, customer operations, blackout windows]"),
    ("Support model after cutover", "[Service desk, L2/L3, vendor, hours, escalation route]"),
])
add_heading(doc, "2.2 Classification and applicability", 2)
add_table(doc, ("Assessment", "Selection", "Rationale / reference"), [
    ("Business criticality", "[ ] Critical  [ ] High  [ ] Medium  [ ] Low", PLACEHOLDER),
    ("GxP / regulated use", "[ ] Yes  [ ] No  [ ] To be determined", "[Assessment / policy reference]"),
    ("Electronic records / signatures", "[ ] Yes  [ ] No", "[21 CFR Part 11 / Annex 11 or other applicability, if relevant]"),
    ("Personal or sensitive data", "[ ] Yes  [ ] No", "[Privacy assessment / data classification reference]"),
    ("Cybersecurity tier", "[Tier / rating]", "[Risk assessment reference]"),
    ("Financial / SOX relevance", "[ ] Yes  [ ] No", "[Control / assessment reference]"),
    ("Records retention / legal hold", "[ ] Applicable  [ ] Not applicable", "[Retention schedule / legal reference]"),
    ("Validation approach", "[ ] Full IQ/OQ/PQ  [ ] Leveraged supplier evidence  [ ] Reduced / risk-based", "[Approved rationale]"),
], [2250, 2850, 4260], font_size=8.7)

add_heading(doc, "2.3 In-scope / out-of-scope boundaries", 2)
add_table(doc, ("Type", "Item / component / process", "Rationale / owner"), [
    ("In scope", PLACEHOLDER, PLACEHOLDER),
    ("In scope", PLACEHOLDER, PLACEHOLDER),
    ("Out of scope", PLACEHOLDER, "[Reason and approving owner]"),
    ("Out of scope", PLACEHOLDER, "[Reason and approving owner]"),
], [1300, 4300, 3760])

doc.add_page_break()
add_heading(doc, "3. Validation strategy, risk and acceptance criteria", 1)
add_heading(doc, "3.1 Validation strategy", 2)
add_key_value_table(doc, [
    ("Overall validation objective", PLACEHOLDER),
    ("Test environments", "[Environment names, versions, URLs / identifiers, data classification]"),
    ("Supplier evidence leveraged", "[Supplier validation, certifications, release notes, SOC report or N/A]"),
    ("Automation / tools", "[Test tool, scripts, data migration tooling, evidence repository]"),
    ("Test data approach", "[Synthetic / masked / production copy; approval and destruction arrangements]"),
    ("Segregation of duties", "[Tester / approver independence and any justified exceptions]"),
    ("Regression scope", "[Critical regression pack and selection rationale]"),
])
add_heading(doc, "3.2 Risk assessment and control coverage", 2)
add_table(doc, ("Risk ID", "Failure mode / risk", "Impact", "Likelihood", "Risk rating", "Control / test coverage", "Residual risk"), [
    ("R-01", "[Example: incomplete migration of open transactions]", "[H/M/L]", "[H/M/L]", "[H/M/L]", "[Requirement / test IDs]", "[H/M/L + rationale]"),
    ("R-02", PLACEHOLDER, "[H/M/L]", "[H/M/L]", "[H/M/L]", "[Requirement / test IDs]", PLACEHOLDER),
    ("R-03", PLACEHOLDER, "[H/M/L]", "[H/M/L]", "[H/M/L]", "[Requirement / test IDs]", PLACEHOLDER),
], [700, 2200, 750, 850, 850, 2200, 1810], font_size=8.0)
add_heading(doc, "3.3 Phase acceptance criteria", 2)
add_table(doc, ("Phase", "Acceptance criteria", "Approved exception route"), [
    ("IQ", "All mandatory installation/configuration checks pass; baselines are recorded; environment is controlled; no open critical/high defect prevents OQ.", "[Deviation authority / process]"),
    ("OQ", "All critical functions and controls pass; interfaces and failure handling meet requirements; defects are resolved or formally accepted; no open critical/high defect prevents PQ.", "[Deviation authority / process]"),
    ("PQ", "Representative end-to-end business scenarios pass; reconciliations meet thresholds; users and support are ready; residual risks are accepted.", "[Deviation authority / process]"),
    ("Release", "Required phase approvals are complete; cutover/rollback and support readiness are confirmed; release authority approves production use.", "[Release governance reference]"),
], [1100, 5860, 2400], font_size=8.8)

doc.add_page_break()
add_heading(doc, "4. Requirements and traceability", 1)
add_body(doc, "List business, functional, data, interface, security, compliance, operability and continuity requirements. Every applicable requirement should trace to risk and test evidence.")
add_table(doc, ("Req. ID", "Requirement / acceptance statement", "Type", "Risk ID", "Priority", "Verification phase / test ID", "Status"), [
    ("REQ-001", "[The system shall ...]", "[BUS/FUN/DATA/SEC/INT/OPS/REG]", "[R-xx]", "[Must/Should]", "[IQ/OQ/PQ - test ID]", "[Open/Passed]"),
    ("REQ-002", PLACEHOLDER, "[Type]", "[R-xx]", "[Priority]", "[Phase / test ID]", "[Status]"),
    ("REQ-003", PLACEHOLDER, "[Type]", "[R-xx]", "[Priority]", "[Phase / test ID]", "[Status]"),
    ("REQ-004", PLACEHOLDER, "[Type]", "[R-xx]", "[Priority]", "[Phase / test ID]", "[Status]"),
    ("REQ-005", PLACEHOLDER, "[Type]", "[R-xx]", "[Priority]", "[Phase / test ID]", "[Status]"),
], [800, 2900, 1100, 800, 1000, 1850, 910], font_size=7.9)
add_heading(doc, "4.1 Traceability completeness check", 2)
add_table(doc, ("Check", "Result", "Evidence / comment"), [
    ("All Must / critical requirements have an approved test", "[Pass / Fail]", PLACEHOLDER),
    ("All high risks have preventive/detective controls and test coverage", "[Pass / Fail]", PLACEHOLDER),
    ("All tests trace to an approved requirement or risk", "[Pass / Fail]", PLACEHOLDER),
    ("All failed tests trace to a deviation / defect and disposition", "[Pass / Fail]", PLACEHOLDER),
], [4400, 1300, 3660])

doc.add_page_break()
add_heading(doc, "5. Installation Qualification (IQ)", 1)
add_body(doc, "Objective: verify that the target application and supporting components are installed, configured, documented and controlled in the approved newMBD environment.")
add_heading(doc, "5.1 IQ prerequisites", 2)
add_table(doc, ("Prerequisite", "Status", "Evidence / reference"), [
    ("Approved architecture / design and environment specification", "[Ready / Not ready / N/A]", PLACEHOLDER),
    ("Approved build, deployment or configuration instructions", "[Ready / Not ready / N/A]", PLACEHOLDER),
    ("Licenses, certificates, keys and vendor agreements available", "[Ready / Not ready / N/A]", PLACEHOLDER),
    ("Network, identity, backup and monitoring dependencies available", "[Ready / Not ready / N/A]", PLACEHOLDER),
    ("Configuration and migration packages version-controlled", "[Ready / Not ready / N/A]", PLACEHOLDER),
], [4800, 1900, 2660], font_size=8.8)
add_heading(doc, "5.2 IQ verification record", 2)
add_test_case_table(doc, "IQ", 8)
add_heading(doc, "5.3 IQ minimum coverage prompts", 2)
add_table(doc, ("Area", "Verify / document"), [
    ("Infrastructure and platform", "Hosting, compute, storage, database, operating system, middleware, tenant/instance, regions, resilience and time synchronization."),
    ("Application and configuration", "Application version/build, modules, feature flags, configuration baseline, custom code/package versions and configuration comparison."),
    ("Connectivity and identity", "DNS, certificates, firewall/routes, SSO, service accounts, identity provider, MFA and privileged access."),
    ("Operations", "Backup/restore configuration, monitoring, alerting, logging, batch scheduler, support access, runbooks and job ownership."),
    ("Security baseline", "Hardening, encryption, secrets storage, vulnerability status, endpoint/security tooling and approved exceptions."),
    ("Documentation / inventory", "CMDB or application inventory, architecture, data flow, support contacts, licenses and technical ownership."),
], [2200, 7160], font_size=9.0)
add_heading(doc, "5.4 IQ deviations and completion", 2)
add_table(doc, ("Metric", "Result"), [
    ("IQ tests planned / executed / passed / failed / N/A", "[## / ## / ## / ## / ##]"),
    ("Open critical / high deviations", "[Count and IDs, or None]"),
    ("IQ conclusion", "[Pass / Pass with conditions / Fail]"),
    ("Conditions / rationale", PLACEHOLDER),
], [3900, 5460])
add_signature_table(doc, ["IQ executor / technical lead", "IT service owner", "Validation lead", "Quality / Compliance (if applicable)"])

doc.add_page_break()
add_heading(doc, "6. Operational Qualification (OQ)", 1)
add_body(doc, "Objective: demonstrate that the application operates as intended across approved functional ranges, controls, roles, interfaces and failure conditions.")
add_heading(doc, "6.1 OQ prerequisites", 2)
add_table(doc, ("Prerequisite", "Status", "Evidence / reference"), [
    ("IQ approved, or approved exception permits OQ", "[Ready / Not ready]", PLACEHOLDER),
    ("Requirements and OQ scripts approved", "[Ready / Not ready]", PLACEHOLDER),
    ("Test accounts, roles and representative data available", "[Ready / Not ready]", PLACEHOLDER),
    ("Interfaces / dependencies available or simulated under control", "[Ready / Not ready]", PLACEHOLDER),
    ("Defect and deviation process active", "[Ready / Not ready]", PLACEHOLDER),
], [4800, 1900, 2660], font_size=8.8)
add_heading(doc, "6.2 OQ functional and control tests", 2)
add_test_case_table(doc, "OQ", 10)
add_heading(doc, "6.3 OQ minimum coverage prompts", 2)
add_table(doc, ("Area", "Verify / document"), [
    ("Core functions", "Normal, boundary, negative and exception processing for critical requirements."),
    ("Access and segregation", "Role-based access, joiner/mover/leaver, privileged access, segregation of duties, authentication and session controls."),
    ("Interfaces and jobs", "Inbound/outbound interfaces, schedules, retries, duplicate prevention, reconciliation, error queues and alerting."),
    ("Records and auditability", "Audit trail, timestamps, attribution, retention, retrieval, e-signature controls and report accuracy where applicable."),
    ("Failure / recovery", "Application errors, interface outage, job failure, restart/recovery, backup restore and rollback behavior."),
    ("Security / privacy", "Encryption, logging, monitoring, vulnerability remediation, privacy controls, export/download and data minimization."),
], [2200, 7160], font_size=9.0)
add_heading(doc, "6.4 OQ completion", 2)
add_table(doc, ("Metric", "Result"), [
    ("OQ tests planned / executed / passed / failed / N/A", "[## / ## / ## / ## / ##]"),
    ("Open critical / high deviations", "[Count and IDs, or None]"),
    ("OQ conclusion", "[Pass / Pass with conditions / Fail]"),
    ("Conditions / rationale", PLACEHOLDER),
], [3900, 5460])
add_signature_table(doc, ["OQ test lead", "IT service owner", "Business process owner", "Validation lead", "Quality / Compliance (if applicable)"])

doc.add_page_break()
add_heading(doc, "7. Performance Qualification (PQ)", 1)
add_body(doc, "Objective: confirm that trained business users can execute representative, end-to-end newMBD processes with acceptable outcomes, performance and operational support under intended conditions.")
add_heading(doc, "7.1 PQ prerequisites", 2)
add_table(doc, ("Prerequisite", "Status", "Evidence / reference"), [
    ("OQ approved, or approved exception permits PQ", "[Ready / Not ready]", PLACEHOLDER),
    ("Business scenarios and acceptance criteria approved", "[Ready / Not ready]", PLACEHOLDER),
    ("Representative users trained and access approved", "[Ready / Not ready]", PLACEHOLDER),
    ("Representative data volumes and process variants available", "[Ready / Not ready]", PLACEHOLDER),
    ("Business procedures / work instructions available", "[Ready / Not ready]", PLACEHOLDER),
], [4800, 1900, 2660], font_size=8.8)
add_heading(doc, "7.2 PQ end-to-end scenarios", 2)
add_test_case_table(doc, "PQ", 8)
add_heading(doc, "7.3 Business acceptance and operability checks", 2)
add_table(doc, ("Check", "Result", "Evidence / business comment"), [
    ("Critical end-to-end processes complete successfully", "[Pass / Fail / N/A]", PLACEHOLDER),
    ("Expected transaction / batch volumes meet agreed service levels", "[Pass / Fail / N/A]", PLACEHOLDER),
    ("Reports, outputs and downstream handoffs are accurate and usable", "[Pass / Fail / N/A]", PLACEHOLDER),
    ("Users can execute exception and escalation procedures", "[Pass / Fail / N/A]", PLACEHOLDER),
    ("Training, work instructions and support routes are adequate", "[Pass / Fail / N/A]", PLACEHOLDER),
    ("Operational monitoring and business control ownership are accepted", "[Pass / Fail / N/A]", PLACEHOLDER),
], [4700, 1700, 2960], font_size=8.8)
add_heading(doc, "7.4 PQ completion", 2)
add_table(doc, ("Metric", "Result"), [
    ("PQ scenarios planned / executed / passed / failed / N/A", "[## / ## / ## / ## / ##]"),
    ("Open critical / high deviations", "[Count and IDs, or None]"),
    ("PQ conclusion", "[Pass / Pass with conditions / Fail]"),
    ("Conditions / rationale", PLACEHOLDER),
], [3900, 5460])
add_signature_table(doc, ["PQ test / business lead", "Business process owner", "IT service owner", "Validation lead", "Quality / Compliance (if applicable)"])

doc.add_page_break()
add_heading(doc, "8. Data migration and reconciliation", 1)
add_body(doc, "Complete this section when data is migrated, carved out, transformed, archived or otherwise transferred from Thermo Fisher-controlled sources to newMBD.")
add_heading(doc, "8.1 Migration population and rules", 2)
add_table(doc, ("Data object / population", "Source", "Target", "Selection / transformation rule", "Expected count / control total", "Owner"), [
    ("[Object / table / file]", "[System / extract]", "[System / object]", "[Rule / mapping reference]", "[Count / amount / hash]", "[Name]"),
    (PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
    (PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
], [1400, 1200, 1200, 2450, 1900, 1210], font_size=7.9)
add_heading(doc, "8.2 Reconciliation results", 2)
add_table(doc, ("Control / population", "Source result", "Target result", "Variance", "Tolerance", "Status", "Evidence / deviation"), [
    ("[Record count / financial total / sample / hash]", "[Value]", "[Value]", "[Value / %]", "[Approved tolerance]", "[Pass / Fail]", "[Evidence / deviation ID]"),
    (PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, "[Pass / Fail]", PLACEHOLDER),
    (PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, "[Pass / Fail]", PLACEHOLDER),
], [1750, 1200, 1200, 1000, 1250, 1000, 1960], font_size=7.8)
add_heading(doc, "8.3 Data disposition and privacy", 2)
add_key_value_table(doc, [
    ("Rejected / quarantined records", "[Count, location, owner and resolution]"),
    ("Historical data availability", "[Online / archive / read-only; retrieval method and SLA]"),
    ("Source data freeze / delta handling", "[Freeze window, delta process and final reconciliation]"),
    ("Thermo Fisher data segregation", "[How data not transferring to newMBD was excluded / protected]"),
    ("Temporary files / extracts", "[Secure storage, access, retention and destruction evidence]"),
    ("Migration conclusion", "[Pass / Pass with conditions / Fail, with rationale]"),
])

doc.add_page_break()
add_heading(doc, "9. Interfaces, security and operational readiness", 1)
add_heading(doc, "9.1 Interface and dependency register", 2)
add_table(doc, ("ID", "System / party", "Direction", "Data / service", "Frequency", "Owner", "Test / status"), [
    ("INT-01", "[System / vendor / customer]", "[In / Out / Bi]", "[Payload / service]", "[Real-time / schedule]", "[Name]", "[Test ID / status]"),
    ("INT-02", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
    ("INT-03", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
], [700, 1700, 900, 2050, 1300, 1200, 1510], font_size=7.9)
add_heading(doc, "9.2 Security, compliance and service readiness", 2)
add_table(doc, ("Readiness item", "Owner", "Status", "Evidence / exception"), [
    ("Security risk assessment and required remediation complete", "[Owner]", "[Ready / Conditional / Not ready / N/A]", PLACEHOLDER),
    ("Privacy assessment / data processing arrangements complete", "[Owner]", "[Ready / Conditional / Not ready / N/A]", PLACEHOLDER),
    ("Access roles, approvers, recertification and privileged access defined", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("Monitoring, alerting, logging and audit retention operational", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("Backup, restore and continuity / disaster recovery arrangements validated", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("Support model, service desk knowledge and escalation paths active", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("Vendor support, licensing, contracts and SLAs effective for newMBD", "[Owner]", "[Ready / Conditional / Not ready / N/A]", PLACEHOLDER),
    ("CMDB / application inventory, ownership and documentation updated", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
], [4000, 1200, 2200, 1960], font_size=8.3)

doc.add_page_break()
add_heading(doc, "10. Deviations, defects and residual risk", 1)
add_heading(doc, "10.1 Deviation / defect register", 2)
add_table(doc, ("ID", "Phase / test", "Description", "Severity", "Impact / risk assessment", "Disposition / retest", "Status"), [
    ("DEV-001", "[IQ/OQ/PQ - ID]", PLACEHOLDER, "[Critical/High/Medium/Low]", PLACEHOLDER, "[Fix / accept / defer; evidence]", "[Open/Closed]"),
    ("DEV-002", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
    ("DEV-003", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
], [750, 1100, 2050, 1100, 1700, 1850, 810], font_size=7.7)
add_heading(doc, "10.2 Open-item acceptance", 2)
add_table(doc, ("Item ID", "Residual risk / business impact", "Compensating control", "Owner", "Due date", "Approving authority"), [
    ("[DEV / RISK ID]", PLACEHOLDER, PLACEHOLDER, "[Name]", "[YYYY-MM-DD]", "[Name / role / approval ref.]"),
    (PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
], [1100, 2200, 2100, 1200, 1200, 1560], font_size=8.0)
add_callout(doc, "Release constraint", "Critical or high-severity unresolved items should block release unless the designated governance authority explicitly accepts the residual risk and documents the time-bound control and remediation plan.", fill=RED)

doc.add_page_break()
add_heading(doc, "11. Cutover, rollback and hypercare readiness", 1)
add_table(doc, ("Readiness check", "Owner", "Status", "Evidence / reference"), [
    ("Approved cutover plan, sequence, timing and decision checkpoints", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("Source freeze and final delta / reconciliation plan", "[Owner]", "[Ready / Conditional / Not ready / N/A]", PLACEHOLDER),
    ("Rollback criteria, authority, steps and recovery point defined", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("Business communications and user readiness confirmed", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("Command center / hypercare staffing and contact list confirmed", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("Monitoring dashboard and first-business-cycle checks ready", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
    ("TSA exit / Thermo Fisher dependency removal confirmed or controlled", "[Owner]", "[Ready / Conditional / Not ready / N/A]", PLACEHOLDER),
    ("Post-go-live review date and success measures agreed", "[Owner]", "[Ready / Conditional / Not ready]", PLACEHOLDER),
], [4050, 1200, 2100, 2010], font_size=8.3)
add_heading(doc, "11.1 Go / no-go recommendation", 2)
add_table(doc, ("Decision element", "Response"), [
    ("Recommended decision", "[ ] GO  [ ] GO WITH CONDITIONS  [ ] NO-GO"),
    ("Conditions / hold points", PLACEHOLDER),
    ("Rollback trigger(s)", PLACEHOLDER),
    ("Decision meeting / record reference", "[Meeting date, minutes / ticket / workflow ID]"),
], [2800, 6560])

doc.add_page_break()
add_heading(doc, "12. Final validation summary and release authorization", 1)
add_heading(doc, "12.1 Phase summary", 2)
add_table(doc, ("Phase / domain", "Conclusion", "Approval date", "Open conditions / references"), [
    ("IQ", "[Pass / Conditional / Fail / N/A]", "[YYYY-MM-DD]", PLACEHOLDER),
    ("OQ", "[Pass / Conditional / Fail / N/A]", "[YYYY-MM-DD]", PLACEHOLDER),
    ("PQ", "[Pass / Conditional / Fail / N/A]", "[YYYY-MM-DD]", PLACEHOLDER),
    ("Data migration / reconciliation", "[Pass / Conditional / Fail / N/A]", "[YYYY-MM-DD]", PLACEHOLDER),
    ("Security / privacy / compliance", "[Ready / Conditional / Not ready / N/A]", "[YYYY-MM-DD]", PLACEHOLDER),
    ("Operational / support readiness", "[Ready / Conditional / Not ready]", "[YYYY-MM-DD]", PLACEHOLDER),
], [2550, 2350, 1550, 2910], font_size=8.6)
add_heading(doc, "12.2 Validation conclusion", 2)
add_key_value_table(doc, [
    ("Requirements / tests summary", "[Requirements total/passed; tests total/executed/passed/failed/N/A]"),
    ("Deviations summary", "[Counts by severity and open/closed status]"),
    ("Residual risk conclusion", "[Acceptable / Not acceptable, with rationale]"),
    ("Validated intended use", "[State precisely the business processes, users, data and boundaries approved]"),
    ("Excluded / conditional use", "[Restrictions, workarounds, locations, modules or dates]"),
    ("Post-release commitments", "[Owner, action, due date and tracking reference]"),
])
add_heading(doc, "12.3 Approval statements", 2)
add_callout(doc, "IT approval statement", "I confirm that the newMBD application and its supporting technical services have been installed/configured, tested, secured, documented and transitioned to an operable support model within the approved scope. Any exceptions and residual technical risks are recorded above.")
add_callout(doc, "Business approval statement", "I confirm that the application supports the approved newMBD business processes and intended use, that representative business qualification has been completed, and that documented residual business risks and conditions are accepted.")
add_callout(doc, "Quality / Compliance statement (when applicable)", "I confirm that the validation record provides appropriate documented evidence for the applicable regulated scope and that deviations, traceability and approvals have been managed according to the governing quality procedures.")
add_heading(doc, "12.4 Final release authorization", 2)
add_signature_table(doc, ["Business owner", "IT service owner", "Validation lead", "Information Security / Privacy (if required)", "Quality / Compliance (if applicable)", "Program / release authority"])
add_table(doc, ("Release decision", "Effective date / time", "Conditions / release record"), [
    ("[ ] Authorized  [ ] Authorized with conditions  [ ] Not authorized", "[YYYY-MM-DD HH:MM time zone]", "[Conditions, change/release ticket, decision record]"),
], [3300, 2200, 3860], header_fill=GREEN, font_size=9.0)

doc.add_page_break()
add_heading(doc, "Appendix A. Evidence index", 1)
add_table(doc, ("Evidence ID", "Description", "Related phase / test / requirement", "Repository / immutable reference", "Owner", "Date"), [
    ("EVD-001", "[Screenshot, log, export, report, approval or reconciliation]", "[IQ/OQ/PQ; test / req. ID]", "[Controlled location / link / record ID]", "[Name]", "[YYYY-MM-DD]"),
    ("EVD-002", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
    ("EVD-003", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
    ("EVD-004", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, PLACEHOLDER),
], [1100, 2300, 1900, 2300, 1000, 760], font_size=7.7)

add_heading(doc, "Appendix B. Referenced documents", 1)
add_table(doc, ("Document / record", "Identifier / version", "Owner", "Location / reference", "Status"), [
    ("Business requirements / intended use", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, "[Approved]"),
    ("Architecture / design specification", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, "[Approved]"),
    ("Risk / security / privacy assessment", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, "[Approved / N/A]"),
    ("Migration / cutover / rollback plan", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, "[Approved]"),
    ("Test scripts / execution records", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, "[Approved]"),
    ("Procedures / training / support runbooks", PLACEHOLDER, PLACEHOLDER, PLACEHOLDER, "[Effective]"),
], [2600, 1700, 1300, 2600, 1160], font_size=8.3)

add_heading(doc, "Appendix C. Optional application-specific approvals", 1)
add_body(doc, "Add or remove roles according to the approved governance model. Examples include Data Owner, Process Control Owner, Records Management, Legal, Vendor Management, Infrastructure Owner, Integration Owner and Site Quality.")
add_signature_table(doc, ["[Additional approver role]", "[Additional approver role]", "[Additional approver role]"])

# Document properties and final layout safeguards.
doc.core_properties.title = "newMBD Application IQ/OQ/PQ Signoff Template"
doc.core_properties.subject = "Application migration and build-back validation for the newMBD divestiture program"
doc.core_properties.author = "newMBD Transition Program"
doc.core_properties.keywords = "newMBD, divestiture, application migration, build-back, IQ, OQ, PQ, validation, signoff"

for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        if run.font.name is None:
            set_font(run, size=10.5)

doc.save(OUT)
print(OUT)
