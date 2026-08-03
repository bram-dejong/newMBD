from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "Asset_Work_Order_Finder_BRD.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E0"
WHITE = "FFFFFF"
BORDER = "C9D2DC"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, bold=None, color=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    if keep_lines:
        node = OxmlElement("w:keepLines")
        p_pr.append(node)


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_char, instr, fld_sep, placeholder, fld_end):
        run._r.append(node)


def add_custom_numbering(doc, abstract_id, num_id, kind):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, ind, spacing))
    lvl.extend((start, num_fmt, lvl_text, lvl_jc, p_pr))
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def add_list_item(doc, text, numbered=False):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "42" if numbered else "41")
    num_pr.extend((ilvl, num_id))
    p_pr.append(num_pr)
    run = p.add_run(text)
    set_font(run, size=11, color=INK)
    return p


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, label in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT_GRAY)
        p = header.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(label)
        set_font(run, size=font_size, bold=True, color=DARK_BLUE)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(value))
            set_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, label, text, fill=PALE_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(TABLE_INDENT / 1440)
    p.paragraph_format.right_indent = Inches(TABLE_INDENT / 1440)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), BORDER)
        borders.append(border)
    p_pr.append(borders)
    r = p.add_run(f"{label}: ")
    set_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_keep(p, keep_next=True)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_font(first, size=11, bold=True, color=INK)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_font(run, size=11, color=INK)
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    left = hp.add_run("BUSINESS REQUIREMENTS DOCUMENT")
    set_font(left, size=8.5, bold=True, color=MUTED)
    hp.add_run("\t")
    right = hp.add_run("Salesforce Mobile Asset Work Order Finder")
    set_font(right, size=8.5, color=MUTED)
    tabs = hp.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_after = Pt(0)
    left = fp.add_run("Internal business review")
    set_font(left, size=8.5, color=MUTED)
    fp.add_run("\t")
    tabs = fp.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_number(fp)

    add_custom_numbering(doc, 41, 41, "bullet")
    add_custom_numbering(doc, 42, 42, "decimal")


def build():
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "Business Requirements Document - Salesforce Mobile Asset Work Order Finder"
    props.subject = "Onsite Asset identification and open Work Order access"
    props.author = "Service Applications"
    props.keywords = "Salesforce, mobile, asset, work order, QR, field service, BRD"

    # Memo masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("BUSINESS REQUIREMENTS DOCUMENT")
    set_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Salesforce Mobile Asset Work Order Finder")
    set_font(r, size=25, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Enable an onsite field service engineer to identify equipment and open the correct active service record")
    set_font(r, size=13, color=MUTED)

    add_table(
        doc,
        ["Document field", "Value"],
        [
            ("Version", "1.0"),
            ("Status", "Draft for business review"),
            ("Date", "3 August 2026"),
            ("Business owner", "Field Service Operations - to be confirmed"),
            ("Product owner", "Service Applications - to be confirmed"),
            ("Target platform", "Salesforce mobile app"),
            ("Current delivery status", "Implemented and deployed to the newMBD scratch org"),
        ],
        [2200, 7160],
        font_size=9.6,
    )

    add_callout(
        doc,
        "Business outcome",
        "An FSE arriving onsite can identify the physical Asset by QR code or verified nameplate text, see only related open Work Orders that the user is authorized to access, and open the selected Work Order without manually searching Salesforce.",
    )

    add_heading(doc, "1. Executive summary", 1)
    add_body(
        doc,
        "Field service engineers need a fast and reliable way to move from the physical equipment in front of them to the correct service work in Salesforce. Manual searching by customer, serial number, or Work Order number takes time and can lead to selection of the wrong Asset or a closed service record."
    )
    add_body(
        doc,
        "The proposed capability provides a mobile-first Asset Work Order Finder. The FSE scans a QR code or uses the camera to read a nameplate, verifies the identified Asset, reviews its open Work Orders, and opens the required record. The process is read-only and respects existing Salesforce record sharing and field permissions."
    )

    add_heading(doc, "2. Purpose and business need", 1)
    add_heading(doc, "2.1 Purpose", 2)
    add_body(
        doc,
        "This BRD defines the business outcomes, scope, user journey, requirements, controls, and acceptance criteria for onsite Asset identification and open Work Order retrieval in the Salesforce mobile app. It is intended to align Field Service Operations, Service Applications, Salesforce administration, security, data owners, and support teams before production rollout."
    )
    add_heading(doc, "2.2 Problem statement", 2)
    for item in [
        "An FSE may not know the Salesforce Asset name or record identifier when standing beside the equipment.",
        "Searching by customer, site, product, serial number, or Work Order can be slow and may return multiple records.",
        "The relationship between a physical Asset and its current open Work Orders is not immediately visible from the equipment itself.",
        "Selecting a closed or unrelated Work Order can create service delay, incorrect updates, and audit concerns.",
        "Camera-based identification must remain assistive: the FSE must verify the result before navigation.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "3. Objectives and success outcomes", 1)
    add_table(
        doc,
        ["Objective", "Expected business outcome"],
        [
            ("Reduce search effort", "Move from the physical equipment to the relevant Work Order in a small number of mobile actions."),
            ("Improve record accuracy", "Use a deterministic QR identifier or a user-confirmed serial match before showing service work."),
            ("Protect service data", "Show only Assets and Work Orders the signed-in user is authorized to read."),
            ("Support field conditions", "Provide QR, camera/OCR, and manual entry paths with clear recovery messages."),
            ("Improve adoption", "Make the finder available as a dedicated entry point in mobile app navigation."),
        ],
        [2550, 6810],
    )

    add_heading(doc, "4. Stakeholders and users", 1)
    add_table(
        doc,
        ["Stakeholder", "Interest / responsibility"],
        [
            ("Field Service Engineer (FSE)", "Primary user; identifies the onsite Asset and opens the correct active Work Order."),
            ("Field Service Operations", "Owns the business process, adoption expectations, and field success measures."),
            ("Dispatch / Service Coordination", "Requires accurate Asset-to-Work Order relationships and may support exception handling."),
            ("Installed Base / Asset Data Owner", "Maintains Asset identifiers, serial numbers, product, account, and location quality."),
            ("Salesforce Product Owner / Administrator", "Owns configuration, permission assignment, mobile placement, and support."),
            ("Information Security and Privacy", "Reviews camera use, data access, image handling, and approved recognition services."),
            ("Service Leadership", "Reviews productivity, data quality, adoption, and service outcome measures."),
        ],
        [2800, 6560],
    )

    add_heading(doc, "5. Scope", 1)
    add_heading(doc, "5.1 In scope", 2)
    for item in [
        "A dedicated Asset Work Order Finder entry in the Salesforce mobile app.",
        "Scanning a QR value from the equipment label.",
        "Matching the scanned value to the unique Salesforce Asset QR identifier.",
        "Using mobile camera/OCR to suggest a serial number from an Asset nameplate.",
        "Requiring the FSE to verify or correct nameplate text and explicitly confirm an Asset candidate.",
        "Displaying the confirmed Asset and its related Work Orders that are not closed.",
        "Opening the selected Asset or Work Order record in Salesforce mobile.",
        "Manual QR or serial entry when camera services are unavailable.",
        "Read-only access through a dedicated FSE permission set.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "5.2 Out of scope", 2)
    for item in [
        "Creating, editing, closing, assigning, or dispatching Work Orders from the finder.",
        "Automatically relocating, updating, or merging Asset records.",
        "Automatically selecting an Asset based only on visual similarity, model name, or low-confidence OCR.",
        "Generating or physically applying QR labels as part of this release.",
        "Guaranteed offline lookup when Salesforce data is unavailable on the device.",
        "Storing nameplate photos in Salesforce or establishing an enterprise image-retention process.",
        "Production integration with an external image-recognition provider beyond approved future enhancement work.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "6. Future-state user journey", 1)
    for item in [
        "The FSE arrives onsite, signs in to the Salesforce mobile app, and opens Asset Work Order Finder.",
        "The FSE scans the equipment QR code. If the QR cannot be read, the FSE may use the nameplate camera/OCR or enter an identifier manually.",
        "Salesforce resolves the identifier to an accessible Asset. Camera/OCR output is presented as a suggested serial number that the FSE verifies or corrects.",
        "When more than one Asset matches a serial number, the FSE selects and confirms the candidate using Asset, product, serial, and location context.",
        "Salesforce displays open Work Orders related to the confirmed Asset. Closed Work Orders are excluded.",
        "The FSE selects the appropriate Work Order and opens its Salesforce mobile record page.",
        "If no open Work Order exists, the app retains the Asset context and clearly explains that no open work was found.",
    ]:
        add_list_item(doc, item, numbered=True)

    add_heading(doc, "7. Business requirements", 1)
    requirements = [
        ("BR-001", "Must", "The FSE must be able to open the finder from Salesforce mobile navigation."),
        ("BR-002", "Must", "The FSE must be able to scan a QR code using the mobile device camera."),
        ("BR-003", "Must", "The scanned QR value must match the dedicated Asset QR identifier using a trimmed, case-insensitive exact comparison."),
        ("BR-004", "Must", "A QR value must resolve to no more than one accessible Asset; duplicate or ambiguous results must stop the process."),
        ("BR-005", "Must", "The FSE must receive a clear message when a QR value is blank, invalid, too long, unknown, or inaccessible."),
        ("BR-006", "Must", "The FSE must be able to use the device camera to read text from an Asset nameplate when the QR path is unavailable."),
        ("BR-007", "Must", "Camera/OCR output must be treated as a suggestion that the FSE can verify or correct before searching."),
        ("BR-008", "Must", "The FSE must explicitly confirm an Asset when a serial number returns one or more candidates."),
        ("BR-009", "Must", "The finder must display the identified Asset name, serial number, product, and location when available."),
        ("BR-010", "Must", "The finder must display only Work Orders related to the confirmed Asset where the Salesforce closed indicator is false."),
        ("BR-011", "Must", "When several open Work Orders exist, all accessible results up to the supported limit must be displayed for FSE selection."),
        ("BR-012", "Must", "The FSE must be able to open a selected Work Order in its Salesforce mobile record page."),
        ("BR-013", "Must", "The finder must clearly state when the confirmed Asset has no open Work Orders and must not substitute a closed Work Order."),
        ("BR-014", "Must", "The FSE must be able to enter a QR code or serial number manually when camera capabilities are unavailable."),
        ("BR-015", "Must", "All Asset and Work Order results must respect the user's Salesforce record sharing, object permissions, and field access."),
        ("BR-016", "Must", "The finder must not create, edit, or delete Asset or Work Order records."),
        ("BR-017", "Must", "Nameplate images must be processed transiently and must not be retained by this capability."),
        ("BR-018", "Should", "Open Work Orders should be presented in a deterministic sequence that favors scheduled work, followed by recently created work."),
        ("BR-019", "Should", "The finder should provide actionable messages for denied camera permission, scanner cancellation, unsupported devices, and connection errors."),
        ("BR-020", "Should", "The business should be able to measure adoption, lookup success, and key data-quality exceptions without retaining sensitive image content."),
    ]
    add_table(doc, ["ID", "Priority", "Requirement"], requirements, [1050, 1100, 7210], font_size=8.7)

    add_heading(doc, "8. Business rules", 1)
    rules = [
        ("BRU-01", "An Asset QR identifier is a stable, unique, case-insensitive value and must not contain credentials or personal data."),
        ("BRU-02", "A physical scan is evidence of the equipment in front of the FSE; it is not authority to change Asset master data."),
        ("BRU-03", "An open Work Order is defined as a Work Order whose Salesforce Is Closed indicator is false."),
        ("BRU-04", "Camera confidence or model similarity alone must never auto-select an Asset or Work Order."),
        ("BRU-05", "If a serial number matches multiple Assets, the user must select and confirm the correct candidate."),
        ("BRU-06", "If the FSE cannot read an Asset or Work Order under Salesforce sharing, the finder must not reveal that record's details."),
        ("BRU-07", "Closed Work Orders are not displayed as alternatives when no open work exists."),
        ("BRU-08", "The finder returns no more than 50 open Work Orders for one Asset and indicates when the limit is reached."),
    ]
    add_table(doc, ["Rule", "Definition"], rules, [1250, 8110], font_size=9)

    add_heading(doc, "9. Data requirements and governance", 1)
    add_table(
        doc,
        ["Data element", "Business expectation", "Owner"],
        [
            ("Asset QR identifier", "Unique, stable, populated for labelled equipment, and governed through controlled loading or integration.", "Installed Base / Data Owner"),
            ("Asset serial number", "Accurate enough to support exact candidate search; duplicate serials require product/location context.", "Installed Base / Data Owner"),
            ("Asset product, account, location", "Available where known to help the FSE confirm the physical equipment.", "Installed Base / Data Owner"),
            ("Work Order Asset relationship", "Each service record must reference the correct Asset for the finder to return it.", "Service Operations / Integration Owner"),
            ("Work Order closure status", "Status configuration must correctly drive the Salesforce Is Closed indicator.", "Service Process Owner"),
            ("Nameplate image", "Transient processing only; no persistence under the current scope.", "Security / Privacy"),
        ],
        [1900, 5100, 2360],
        font_size=8.6,
    )

    add_heading(doc, "10. Security, privacy, and compliance", 1)
    for item in [
        "Access is granted through the Asset Work Order Finder User permission set or an approved permission-set group.",
        "The permission grants read-only access to the required Asset, Product, Account, Location, and Work Order data and no View All or Modify All access.",
        "Server-side access must run with Salesforce sharing and user-mode field/object enforcement.",
        "QR and OCR values are untrusted input and must be length-limited, normalized, and used only as query values.",
        "QR payloads must not be executed or followed as URLs by the finder.",
        "Images must not be logged or retained. Any future external recognition provider requires security, privacy, retention, and vendor approval.",
        "Error messages must support recovery without disclosing inaccessible record existence or sensitive details.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "11. Non-functional requirements", 1)
    add_table(
        doc,
        ["Area", "Requirement"],
        [
            ("Usability", "Primary actions must be usable on supported phone form factors with touch-friendly controls and concise instructions."),
            ("Performance", "Under normal mobile connectivity, a lookup should return quickly enough to support onsite work; the production target is to be baselined during pilot."),
            ("Reliability", "Failures must preserve user context where practical and provide a retry or manual fallback."),
            ("Accessibility", "Labels, status messages, focus order, and actionable controls must support Salesforce mobile accessibility standards."),
            ("Compatibility", "The pilot must verify supported managed iOS and Android devices, camera permissions, and current Salesforce mobile versions."),
            ("Scalability", "A single lookup must be bounded; no more than 50 open Work Orders are returned for one Asset."),
            ("Supportability", "Support teams must be able to distinguish scanner, permission, data-quality, and Salesforce availability issues."),
        ],
        [1900, 7460],
        font_size=9,
    )

    add_heading(doc, "12. Exception handling", 1)
    add_table(
        doc,
        ["Scenario", "Expected outcome"],
        [
            ("QR is unreadable", "Offer camera/nameplate and manual-entry alternatives."),
            ("QR is unknown", "State that no accessible Asset matched; do not navigate."),
            ("Duplicate QR data", "Stop and direct the user to data support; do not choose either Asset."),
            ("OCR finds no clear serial", "Prompt the FSE to enter or correct the serial manually."),
            ("Serial has multiple candidates", "Display identifying context and require explicit Asset selection."),
            ("No open Work Orders", "Show the confirmed Asset and an explicit no-open-work message."),
            ("More than 50 open Work Orders", "Show the bounded result and indicate that the supported limit was reached."),
            ("Camera permission denied", "Explain how to enable permission and retain manual entry."),
            ("User lacks record access", "Return only accessible records and avoid revealing restricted details."),
            ("Poor or lost connectivity", "Show a recoverable error and allow retry when connectivity returns."),
        ],
        [2750, 6610],
        font_size=8.8,
    )

    add_heading(doc, "13. Acceptance criteria", 1)
    acceptance = [
        ("AC-01", "Given a valid unique Asset QR code, when the FSE scans it, then the correct accessible Asset and only its open Work Orders are displayed."),
        ("AC-02", "Given a QR value with whitespace or case variation, when it is submitted, then the same unique Asset is found."),
        ("AC-03", "Given an unknown or ambiguous QR value, when it is submitted, then no Asset is selected and a clear error is shown."),
        ("AC-04", "Given a supported mobile device, when the FSE scans a nameplate, then OCR suggests text and the FSE can verify or correct the serial before search."),
        ("AC-05", "Given multiple Assets with the same serial number, when candidates are returned, then no Asset is used until the FSE explicitly confirms one."),
        ("AC-06", "Given one or more open Work Orders, when the FSE selects one, then Salesforce mobile opens that Work Order record."),
        ("AC-07", "Given only closed Work Orders, when the Asset is identified, then the finder states that no open Work Order exists and displays no closed substitute."),
        ("AC-08", "Given a user without access to a related Work Order, when a lookup is performed, then the restricted Work Order is not returned."),
        ("AC-09", "Given camera services are unavailable, when the FSE uses manual entry, then the Asset and open-work lookup remains available."),
        ("AC-10", "Given a completed nameplate attempt, then the finder does not retain the image as a Salesforce file or business record."),
    ]
    add_table(doc, ["ID", "Acceptance criterion"], acceptance, [1050, 8310], font_size=8.8)

    add_heading(doc, "14. Dependencies, assumptions, and constraints", 1)
    add_heading(doc, "14.1 Dependencies", 2)
    for item in [
        "Assets are populated with a stable QR identifier and physically labelled where the QR route is expected.",
        "Work Orders reference the correct Asset through the standard Asset relationship.",
        "Salesforce status configuration correctly maintains the Work Order closed indicator.",
        "The Salesforce mobile app and device operating system permit camera use.",
        "FSE users receive the required permission set and appropriate record sharing.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "14.2 Assumptions and constraints", 2)
    for item in [
        "The first production release is online-first; a true offline data design is not assumed.",
        "QR codes encode an opaque Asset identifier, not credentials or a privileged URL.",
        "Serial-number identification uses exact matching after the FSE verifies the OCR suggestion.",
        "The scratch-org implementation demonstrates the requirement but does not constitute production approval.",
        "Production deployment remains subject to business acceptance, security review, release governance, and mobile pilot results.",
    ]:
        add_list_item(doc, item)

    add_heading(doc, "15. Risks and mitigations", 1)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ("Missing or incorrect QR values", "Asset cannot be identified or the wrong label is used.", "Preflight data, enforce uniqueness, govern label generation, and monitor exceptions."),
            ("Duplicate serial numbers", "Camera/nameplate search returns several Assets.", "Require FSE confirmation using product, location, and Asset context."),
            ("Incorrect Work Order to Asset relationship", "Relevant work is missing or unrelated work appears.", "Add data-quality reporting and ownership for relationship corrections."),
            ("Restricted record sharing", "FSE cannot see required work.", "Test representative territories and correct the sharing model through approved governance."),
            ("Camera/OCR limitations", "Nameplate text is not read accurately.", "Keep QR as primary, require verification, and retain manual entry."),
            ("Poor field connectivity", "Lookup is delayed or unavailable.", "Provide clear retry behavior and evaluate offline needs after pilot evidence."),
            ("Unapproved image processing", "Privacy or vendor risk.", "Use transient native processing; require formal review before external-provider adoption."),
        ],
        [2500, 2800, 4060],
        font_size=8.2,
    )

    add_heading(doc, "16. Rollout and change management", 1)
    for item in [
        "Business approval: confirm scope, ownership, open-work definition, data stewardship, and pilot measures.",
        "Data readiness: profile QR completeness, QR uniqueness, serial quality, Asset relationships, and Work Order closure behavior.",
        "Security readiness: approve permission-set assignment, sharing outcomes, camera use, and transient image handling.",
        "Pilot: assign a small FSE cohort across representative devices, sites, products, and service territories.",
        "Training: provide a short mobile guide covering QR first, nameplate verification, manual fallback, and error escalation.",
        "Measure and refine: review lookup success, time to open a Work Order, no-match rate, access issues, and user feedback.",
        "Production rollout: proceed through standard release governance after pilot acceptance and support readiness.",
    ]:
        add_list_item(doc, item, numbered=True)

    add_heading(doc, "17. Success measures", 1)
    add_table(
        doc,
        ["Measure", "Definition", "Initial target"],
        [
            ("Lookup success rate", "Percent of attempts that identify the intended accessible Asset.", "Baseline in pilot; target to be approved"),
            ("Time to open Work Order", "Median time from opening the finder to opening the selected Work Order.", "Baseline in pilot; target to be approved"),
            ("QR data coverage", "Percent of in-scope labelled Assets with a valid unique QR identifier.", "Target to be set by data owner"),
            ("No-open-work rate", "Percent of successful Asset identifications with no open Work Order.", "Monitor for process insight"),
            ("Access failure rate", "Percent of attempts blocked by missing permission or sharing.", "Trend toward zero after pilot"),
            ("FSE adoption", "Percent of pilot FSEs using the finder for eligible onsite visits.", "Target to be agreed with operations"),
        ],
        [2250, 4450, 2660],
        font_size=8.4,
    )

    add_heading(doc, "18. Decisions and open questions", 1)
    add_table(
        doc,
        ["Decision / question", "Owner", "Status"],
        [
            ("Confirm the business owner and product owner for production rollout.", "Field Service / Service Applications", "Open"),
            ("Confirm whether Is Closed = false is the final enterprise definition of open work.", "Service Process Owner", "Open"),
            ("Confirm the in-scope Asset population and QR labelling approach.", "Installed Base / Operations", "Open"),
            ("Define pilot success thresholds and reporting cadence.", "Field Service Leadership", "Open"),
            ("Confirm supported managed mobile devices and minimum Salesforce mobile version.", "Mobility / IT", "Open"),
            ("Decide whether production requires offline lookup capability.", "Field Service / Architecture", "Open"),
            ("Decide whether an approved external image-recognition provider is needed after native OCR evaluation.", "Security / Architecture / Business", "Open"),
        ],
        [4680, 2920, 1760],
        font_size=8.4,
    )

    add_heading(doc, "19. Validation status", 1)
    add_body(
        doc,
        "A working implementation has been deployed to the newMBD scratch org for validation. The deployment included the mobile component, server-side lookup, permission set, mobile App Page, navigation tab, and supporting nameplate-candidate behavior. Eleven relevant automated tests passed with no failures. This technical validation does not replace business acceptance testing or production release approval."
    )

    add_heading(doc, "20. Approval", 1)
    add_body(doc, "Approval confirms agreement with the business scope and requirements in this document. It does not by itself authorize production deployment.")
    add_table(
        doc,
        ["Role", "Name", "Decision", "Date"],
        [
            ("Business Owner", "", "Approve / Reject", ""),
            ("Field Service Operations", "", "Approve / Reject", ""),
            ("Salesforce Product Owner", "", "Approve / Reject", ""),
            ("Information Security / Privacy", "", "Approve / Reject", ""),
        ],
        [2400, 2400, 2560, 2000],
        font_size=9,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
