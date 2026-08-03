from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "PREVENTIVE_MAINTENANCE_AUTOMATION_BRD.md"
OUT = ROOT / "output" / "newMBD_Preventive_Maintenance_Automation_BRD.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
BORDER = "C9D2DC"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, size=None, bold=None, italic=None, color=INK, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_keep(paragraph, keep_next=False, keep_lines=True):
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT))
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def add_numbering(doc, abstract_id, num_id, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multilevel = OxmlElement("w:multiLevelType")
    multilevel.set(qn("w:val"), "singleLevel")
    abstract.append(multilevel)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, attr, value in (
        ("w:start", "w:val", "1"),
        ("w:numFmt", "w:val", "bullet" if bullet else "decimal"),
        ("w:lvlText", "w:val", "•" if bullet else "%1."),
        ("w:lvlJc", "w:val", "left"),
    ):
        node = OxmlElement(tag)
        node.set(qn(attr), value)
        level.append(node)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend((tabs, indent, spacing))
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    for tag, attr, value in (
        ("w:fldChar", "w:fldCharType", "begin"),
        ("w:instrText", "xml:space", "preserve"),
        ("w:fldChar", "w:fldCharType", "separate"),
        ("w:t", None, None),
        ("w:fldChar", "w:fldCharType", "end"),
    ):
        node = OxmlElement(tag)
        if attr:
            node.set(qn(attr), value)
        if tag == "w:instrText":
            node.text = " PAGE "
        if tag == "w:t":
            node.text = "1"
        run._r.append(node)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    hp = section.header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.add_run("BUSINESS REQUIREMENTS DOCUMENT"), size=8.5, bold=True, color=MUTED)
    hp.add_run("\t")
    set_font(hp.add_run("newMBD Preventive Maintenance Automation"), size=8.5, color=MUTED)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

    fp = section.footer.paragraphs[0]
    fp.paragraph_format.space_after = Pt(0)
    set_font(fp.add_run("Draft for business validation"), size=8.5, color=MUTED)
    fp.add_run("\t")
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_number(fp)
    add_numbering(doc, 71, 71, bullet=True)
    add_numbering(doc, 72, 72, bullet=False)


def add_inline(paragraph, text, size=11, color=INK):
    token = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    position = 0
    for match in token.finditer(text):
        if match.start() > position:
            set_font(paragraph.add_run(text[position:match.start()]), size=size, color=color)
        value = match.group(0)
        if value.startswith("`"):
            set_font(paragraph.add_run(value[1:-1]), size=size - 0.3, color=DARK_BLUE, name="Consolas")
        else:
            set_font(paragraph.add_run(value[2:-2]), size=size, bold=True, color=color)
        position = match.end()
    if position < len(text):
        set_font(paragraph.add_run(text[position:]), size=size, color=color)


def add_list(doc, text, numbered=False):
    paragraph = doc.add_paragraph()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), "72" if numbered else "71")
    numpr.extend((ilvl, numid))
    paragraph._p.get_or_add_pPr().append(numpr)
    add_inline(paragraph, text)
    return paragraph


def parse_table(lines):
    rows = []
    for line in lines:
        rows.append([cell.strip() for cell in line.strip().strip("|").split("|")])
    return [rows[0]] + rows[2:]


def table_widths(headers):
    count = len(headers)
    lowered = [header.lower() for header in headers]
    if count == 2:
        return [2200, 7160]
    if count == 3 and lowered[0] == "id":
        return [1250, 1250, 6860]
    if count == 3:
        return [2100, 2700, 4560]
    return [TABLE_WIDTH // count] * (count - 1) + [TABLE_WIDTH - (TABLE_WIDTH // count) * (count - 1)]


def add_table(doc, matrix):
    headers, rows = matrix[0], matrix[1:]
    widths = table_widths(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        shade_cell(table.rows[0].cells[index], LIGHT_GRAY)
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_inline(paragraph, header, size=9.1, color=DARK_BLUE)
        for run in paragraph.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(paragraph, value, size=8.8)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_callout(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(TABLE_INDENT / 1440)
    paragraph.paragraph_format.right_indent = Inches(TABLE_INDENT / 1440)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE_BLUE)
    ppr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), BORDER)
        borders.append(border)
    ppr.append(borders)
    add_inline(paragraph, text, size=10.5)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    doc = Document()
    configure(doc)
    props = doc.core_properties
    props.title = title
    props.subject = "ServiceMax-independent preventive-maintenance Work Order generation"
    props.author = "Service Applications"
    props.keywords = "Salesforce, newMBD, Asset, Entitlement, Service Contract, Work Order, preventive maintenance, BRD"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(14)
    kicker.paragraph_format.space_after = Pt(2)
    set_font(kicker.add_run("BUSINESS REQUIREMENTS DOCUMENT"), size=10, bold=True, color=BLUE)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(5)
    set_font(title_p.add_run("Preventive Maintenance Work Order Automation"), size=25, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("Replace ServiceMax PM generation with standard Asset, Entitlement, Service Contract, and Work Order automation"), size=13, color=MUTED)

    index = 1
    paragraph_buffer = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, " ".join(item.strip() for item in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        if not line.strip():
            flush_paragraph()
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|"):
            flush_paragraph()
            block = []
            while index < len(lines) and lines[index].startswith("|"):
                block.append(lines[index])
                index += 1
            add_table(doc, parse_table(block))
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = min(len(heading.group(1)) - 1, 3)
            paragraph = doc.add_paragraph(heading.group(2), style=f"Heading {level}")
            set_keep(paragraph, keep_next=True)
            index += 1
            continue
        if line.startswith("> "):
            flush_paragraph()
            add_callout(doc, line[2:].strip())
            index += 1
            continue
        if line.startswith("- "):
            flush_paragraph()
            add_list(doc, line[2:].strip())
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            add_list(doc, numbered.group(1), numbered=True)
            index += 1
            continue
        paragraph_buffer.append(line)
        index += 1
    flush_paragraph()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
